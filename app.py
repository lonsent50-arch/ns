#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Novel Studio - 专业小说写作工具后端
功能：项目管理、角色管理、情节大纲、章节管理、AI写作、导出
"""

import os
import json
import sqlite3

from dotenv import load_dotenv
load_dotenv()
import secrets
import hashlib
import time
import threading
import functools
from datetime import datetime
from pathlib import Path

import gzip
from flask import Flask, request, jsonify, send_from_directory, render_template_string, Response, g
from flask_cors import CORS
import markdown
from novel_analyzer import (analyze_cliffhanger_tail, realism_radar,
                             analyze_rhythm, scan_cliches)

# ── Supabase 数据仓库（多租户云端模式）──
try:
    from supabase_repo import init_supabase_for_request, get_repo, get_current_user_id, is_supabase_mode
    _HAS_SUPABASE_REPO = True
except ImportError:
    _HAS_SUPABASE_REPO = False
    def init_supabase_for_request(): pass
    def get_repo(): return None
    def get_current_user_id(): return ''
    def is_supabase_mode(): return False

# ── 启动时环境变量校验 ──
_REQUIRED_ENV_VARS = [
    'DEEPSEEK_API_KEY',
]
_OPTIONAL_ENV_VARS = [
    'CLAUDE_API_KEY',
    'GEMINI_API_KEY',
    'SUPABASE_URL',
    'SUPABASE_ANON_KEY',
    'SUPABASE_SERVICE_ROLE_KEY',
    'LICENSE_SECRET',
]
_missing = [v for v in _REQUIRED_ENV_VARS if not os.environ.get(v)]
_APP_DEGRADED = bool(_missing)
if _missing:
    print(f'[Novel Studio] ⚠️  以下必需环境变量未设置，AI 功能将不可用：{", ".join(_missing)}。请在 .env 文件或云平台环境变量中配置。')
_missing_opt = [v for v in _OPTIONAL_ENV_VARS if not os.environ.get(v)]
if _missing_opt:
    print(f'[Novel Studio] ⚠️  以下可选环境变量未设置（相应功能将不可用）：{", ".join(_missing_opt)}')
from docx import Document
from docx.shared import Pt, Inches
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

app = Flask(__name__, static_folder='static', template_folder='templates')

# ── CORS 跨域配置：生产环境锁定来源，开发环境允许 localhost ──
_ALLOWED_ORIGINS = os.environ.get('ALLOWED_ORIGINS', '')
if _ALLOWED_ORIGINS:
    _cors_origins = [o.strip() for o in _ALLOWED_ORIGINS.split(',') if o.strip()]
else:
    # 默认：仅允许本地开发 + 常见 Vercel 部署域
    _cors_origins = [
        'http://localhost:5050',
        'http://localhost:3000',
        'http://127.0.0.1:5050',
        'http://127.0.0.1:3000',
        'https://novel-studio.vercel.app',
    ]
CORS(app, resources={
    r"/api/*": {
        "origins": _cors_origins,
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization", "X-Admin-Token"],
    }
})

@app.before_request
def _init_supabase():
    """每个请求前初始化 Supabase 客户端（从 JWT 注入）"""
    if _HAS_SUPABASE_REPO and is_supabase_mode():
        try:
            init_supabase_for_request()
        except Exception as e:
            print(f'[Supabase] before_request 初始化失败: {e}')

@app.after_request
def _add_response_headers(response):
    """网络优化：gzip 压缩 + 安全头 + 静态资源缓存"""
    # 安全头
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'

    # 静态资源长缓存
    path = request.path
    if '/static/' in path or path.endswith('.js') or path.endswith('.css'):
        if path.endswith('.woff2') or path.endswith('.woff') or path.endswith('.ttf'):
            response.headers['Cache-Control'] = 'public, max-age=2592000, immutable'
        elif path.endswith('.png') or path.endswith('.jpg') or path.endswith('.svg') or path.endswith('.ico'):
            response.headers['Cache-Control'] = 'public, max-age=604800'
        else:
            response.headers['Cache-Control'] = 'public, max-age=86400'

    # gzip 压缩（跳过二进制类型和直接透传模式）
    content_type = response.content_type or ''
    skip_ct = any(t in content_type for t in ('image', 'video', 'audio', 'zip', 'pdf', 'octet-stream'))
    accept_encoding = request.headers.get('Accept-Encoding', '')
    if not skip_ct and 'gzip' in accept_encoding and response.content_length and response.content_length > 500:
        try:
            data = response.get_data()
            response.data = gzip.compress(data)
            response.headers['Content-Encoding'] = 'gzip'
            response.headers['Content-Length'] = str(len(response.data))
        except RuntimeError:
            pass  # direct passthrough mode (static files), skip gzip

    return response

# ===== 配置 =====
BASE_DIR = Path(__file__).parent
PROJECTS_DIR = BASE_DIR / 'projects'
PROJECTS_DIR.mkdir(exist_ok=True)
DEEPSEEK_API_KEY = os.environ.get('DEEPSEEK_API_KEY', '')
DEEPSEEK_API_URL = 'https://api.deepseek.com/chat/completions'

# AI 模型版本 —— 集中管理，后续升级只需改此处
DEEPSEEK_MODEL = 'deepseek-chat'
CLAUDE_MODEL = 'claude-sonnet-4-20250514'
GEMINI_MODEL = 'gemini-2.5-flash'
GEMINI_API_URL = f'https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent'

LICENSE_SECRET = os.environ.get('LICENSE_SECRET')
if not LICENSE_SECRET:
    raise RuntimeError('[Novel Studio] 启动失败：LICENSE_SECRET 环境变量未设置。请在 .env 或云平台环境变量中配置。')

def validate_project_id(project_id):
    """防止路径遍历攻击：确保 project_id 不包含危险字符且路径不越界"""
    if not project_id or '..' in project_id or '/' in project_id or '\\' in project_id:
        return None
    project_path = (PROJECTS_DIR / project_id).resolve()
    if not str(project_path).startswith(str(PROJECTS_DIR.resolve())):
        return None
    return project_id

# 列名白名单验证器，防止 SQL 注入
_SQL_COLUMN_PATTERN = __import__('re').compile(r'^[a-zA-Z_][a-zA-Z0-9_]*$')

def validate_sql_column(name, allowed=None):
    """确保列名只包含安全字符，可选白名单校验。"""
    if not _SQL_COLUMN_PATTERN.match(name):
        return False
    if allowed is not None and name not in allowed:
        return False
    return True

def safe_json_extract(text, default=None):
    """从 AI 返回文本中安全提取 JSON。
    自动处理 markdown 代码块包裹、尾部多余逗号等常见格式问题。
    """
    if not text or not isinstance(text, str):
        return default
    cleaned = text.strip()
    # 移除 markdown 代码块标记
    for prefix in ['```json\n', '```json', '```\n', '```']:
        if cleaned.startswith(prefix):
            cleaned = cleaned[len(prefix):]
            break
    for suffix in ['\n```', '```']:
        if cleaned.endswith(suffix):
            cleaned = cleaned[:-len(suffix)]
            break
    cleaned = cleaned.strip()
    if not cleaned:
        return default
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass
    # 尝试修复：去除尾部逗号（AI 常见错误）
    import re as _re
    cleaned = _re.sub(r',\s*}', '}', cleaned)
    cleaned = _re.sub(r',\s*\]', ']', cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return default

@app.before_request
def _validate_project_id():
    """对所有包含 project_id 的路由自动验证路径安全 + 项目存在性"""
    project_id = request.view_args.get('project_id') if request.view_args else None
    if project_id:
        if not validate_project_id(project_id):
            return jsonify({'error': '无效的项目ID'}), 400
        # Supabase 模式：跳过本地文件检查
        if _HAS_SUPABASE_REPO and is_supabase_mode():
            return
        project_path = PROJECTS_DIR / project_id
        if not project_path.is_dir() or not (project_path / 'novel.db').exists():
            return jsonify({'error': '项目不存在'}), 404

# ===== 数据库工具 =====
SQLITE_TIMEOUT = 10  # 秒，等待锁的最长时间

def _configure_db_conn(conn):
    """统一配置 SQLite 连接：WAL 模式 + 忙等超时"""
    conn.execute(f'PRAGMA journal_mode=WAL')
    conn.execute(f'PRAGMA busy_timeout={SQLITE_TIMEOUT * 1000}')
    conn.row_factory = sqlite3.Row

def get_db(project_id):
    """获取项目数据库连接（调用方须在完成后 conn.close()）"""
    db_path = PROJECTS_DIR / project_id / 'novel.db'
    conn = sqlite3.connect(str(db_path))
    _configure_db_conn(conn)
    return conn

from contextlib import contextmanager

@contextmanager
def get_db_safe(project_id):
    """安全数据库上下文管理器，自动关闭连接，异常安全"""
    conn = sqlite3.connect(str(PROJECTS_DIR / project_id / 'novel.db'))
    _configure_db_conn(conn)
    try:
        yield conn
    finally:
        conn.close()

# ===== 请求频率限制 =====
_rate_limits = {}  # {user_id: [(timestamp, endpoint), ...]}
_rate_lock = threading.Lock()
RATE_LIMIT_WINDOW = 60  # 秒
RATE_LIMIT_MAX = 30     # 每窗口最多请求数（AI 端点）
RATE_LIMIT_GENERAL = 60 # 普通端点

def check_rate_limit(user_id, endpoint='general'):
    """简单滑动窗口限流，防止 API 滥用。每 100 次调用自动清理过期条目。"""
    max_req = RATE_LIMIT_MAX if endpoint.startswith('ai') or 'ai/' in endpoint else RATE_LIMIT_GENERAL
    now = time.time()
    with _rate_lock:
        if user_id not in _rate_limits:
            _rate_limits[user_id] = []
        window = [t for t in _rate_limits[user_id] if now - t[0] < RATE_LIMIT_WINDOW]
        if len(window) >= max_req:
            return False
        window.append((now, endpoint))
        _rate_limits[user_id] = window

        # 每 100 次调用清理一次过期用户条目
        check_rate_limit.counter = getattr(check_rate_limit, 'counter', 0) + 1
        if check_rate_limit.counter % 100 == 0:
            expired_users = [uid for uid, entries in _rate_limits.items() if not entries]
            for uid in expired_users:
                del _rate_limits[uid]
        return True
check_rate_limit.counter = 0

def init_project_db(project_id):
    """初始化项目数据库"""
    db_path = PROJECTS_DIR / project_id / 'novel.db'
    conn = sqlite3.connect(str(db_path))
    c = conn.cursor()
    
    # 项目信息表
    c.execute('''
        CREATE TABLE IF NOT EXISTS project_info (
            key TEXT PRIMARY KEY,
            value TEXT
        )
    ''')
    
    # 角色表
    c.execute('''
        CREATE TABLE IF NOT EXISTS characters (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            gender TEXT,
            age TEXT,
            personality TEXT,
            background TEXT,
            goal TEXT,
            appearance TEXT,
            notes TEXT,
            created_at TEXT,
            updated_at TEXT
        )
    ''')
    
    # 大纲节点表
    c.execute('''
        CREATE TABLE IF NOT EXISTS outline (
            id TEXT PRIMARY KEY,
            parent_id TEXT,
            title TEXT NOT NULL,
            content TEXT,
            level INTEGER DEFAULT 0,
            sort_order INTEGER DEFAULT 0,
            created_at TEXT,
            updated_at TEXT,
            FOREIGN KEY (parent_id) REFERENCES outline(id)
        )
    ''')
    
    # 章节表
    c.execute('''
        CREATE TABLE IF NOT EXISTS chapters (
            id TEXT PRIMARY KEY,
            title TEXT NOT NULL,
            content TEXT DEFAULT '',
            sort_order INTEGER DEFAULT 0,
            created_at TEXT,
            updated_at TEXT,
            word_count INTEGER DEFAULT 0,
            source TEXT DEFAULT 'manual'
        )
    ''')
    
    # 写作统计表
    c.execute('''
        CREATE TABLE IF NOT EXISTS writing_stats (
            date TEXT PRIMARY KEY,
            chars_added INTEGER DEFAULT 0,
            chars_deleted INTEGER DEFAULT 0,
            time_spent INTEGER DEFAULT 0,
            sessions INTEGER DEFAULT 0
        )
    ''')
    
    # 写作目标表
    c.execute('''
        CREATE TABLE IF NOT EXISTS writing_goals (
            id TEXT PRIMARY KEY,
            goal_type TEXT,
            target_value INTEGER,
            current_value INTEGER DEFAULT 0,
            deadline TEXT,
            is_active INTEGER DEFAULT 1
        )
    ''')

    # 章节摘要表（长篇上下文管理）
    c.execute('''
        CREATE TABLE IF NOT EXISTS chapter_summaries (
            chapter_id TEXT PRIMARY KEY,
            summary TEXT NOT NULL,
            key_events TEXT DEFAULT '[]',
            character_states TEXT DEFAULT '{}',
            plot_threads TEXT DEFAULT '[]',
            generated_at TEXT,
            word_count INTEGER DEFAULT 0,
            FOREIGN KEY (chapter_id) REFERENCES chapters(id)
        )
    ''')

    # 关键事件表（跨章节追踪）
    c.execute('''
        CREATE TABLE IF NOT EXISTS key_events (
            id TEXT PRIMARY KEY,
            chapter_id TEXT NOT NULL,
            title TEXT NOT NULL,
            description TEXT DEFAULT '',
            event_type TEXT DEFAULT 'event',
            involved_characters TEXT DEFAULT '[]',
            sort_order REAL DEFAULT 0,
            created_at TEXT,
            FOREIGN KEY (chapter_id) REFERENCES chapters(id)
        )
    ''')

    # 角色状态快照表
    c.execute('''
        CREATE TABLE IF NOT EXISTS character_states (
            id TEXT PRIMARY KEY,
            character_id TEXT NOT NULL,
            chapter_id TEXT NOT NULL,
            location TEXT DEFAULT '',
            status TEXT DEFAULT 'alive',
            emotional_state TEXT DEFAULT '',
            knowledge_gained TEXT DEFAULT '',
            relationships TEXT DEFAULT '{}',
            snapshot_at TEXT,
            FOREIGN KEY (character_id) REFERENCES characters(id),
            FOREIGN KEY (chapter_id) REFERENCES chapters(id)
        )
    ''')

    # 角色知识账本表（角色信息隔离墙）
    c.execute('''
        CREATE TABLE IF NOT EXISTS character_knowledge (
            id TEXT PRIMARY KEY,
            character_id TEXT NOT NULL,
            chapter_id TEXT NOT NULL,
            known_names TEXT DEFAULT '[]',
            known_items TEXT DEFAULT '[]',
            known_events TEXT DEFAULT '[]',
            snapshot_at TEXT,
            FOREIGN KEY (character_id) REFERENCES characters(id),
            FOREIGN KEY (chapter_id) REFERENCES chapters(id)
        )
    ''')

    # 情节线程表
    c.execute('''
        CREATE TABLE IF NOT EXISTS plot_threads (
            id TEXT PRIMARY KEY,
            title TEXT NOT NULL,
            description TEXT DEFAULT '',
            thread_type TEXT DEFAULT 'subplot',
            status TEXT DEFAULT 'active',
            start_chapter_id TEXT,
            end_chapter_id TEXT,
            created_at TEXT,
            updated_at TEXT
        )
    ''')

    # AI 对话归类表
    c.execute('''
        CREATE TABLE IF NOT EXISTS ai_conversations (
            id TEXT PRIMARY KEY,
            category TEXT DEFAULT 'general',
            topic TEXT DEFAULT '',
            source_tab TEXT DEFAULT 'chat',
            message_count INTEGER DEFAULT 0,
            last_message_at TEXT,
            created_at TEXT
        )
    ''')

    c.execute('''
        CREATE TABLE IF NOT EXISTS ai_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id TEXT NOT NULL,
            role TEXT NOT NULL,
            content TEXT NOT NULL,
            timestamp TEXT,
            FOREIGN KEY (conversation_id) REFERENCES ai_conversations(id)
        )
    ''')

    # 章节版本历史表
    c.execute('''
        CREATE TABLE IF NOT EXISTS chapter_snapshots (
            id TEXT PRIMARY KEY,
            chapter_id TEXT NOT NULL,
            version INTEGER NOT NULL,
            title TEXT,
            content TEXT,
            word_count INTEGER DEFAULT 0,
            snapshot_at TEXT NOT NULL,
            FOREIGN KEY (chapter_id) REFERENCES chapters(id)
        )
    ''')
    c.execute('CREATE INDEX IF NOT EXISTS idx_snapshots_chapter ON chapter_snapshots(chapter_id, version)')

    # 世界观构建器表
    c.execute('''
        CREATE TABLE IF NOT EXISTS worldbuilding (
            id TEXT PRIMARY KEY,
            category TEXT NOT NULL,
            name TEXT NOT NULL,
            description TEXT DEFAULT '',
            details TEXT DEFAULT '{}',
            sort_order REAL DEFAULT 0,
            created_at TEXT,
            updated_at TEXT
        )
    ''')
    c.execute('CREATE INDEX IF NOT EXISTS idx_worldbuilding_cat ON worldbuilding(category)')

    # RLS 占位：为后续多租户隔离预留 user_id 列
    try:
        c.execute('ALTER TABLE project_info ADD COLUMN user_id TEXT')
    except sqlite3.OperationalError:
        pass  # 列已存在

    conn.commit()
    conn.close()

# ===== 项目管理 API =====
@app.route('/api/health')
def api_health():
    status = 'degraded' if _APP_DEGRADED else 'ok'
    resp = {'status': status, 'version': '1.0', 'timestamp': datetime.utcnow().isoformat()}
    if _APP_DEGRADED:
        resp['missing_env'] = _REQUIRED_ENV_VARS
    return jsonify(resp)

@app.route('/api/projects', methods=['GET'])
def list_projects():
    """列出所有项目（Supabase 模式自动多租户过滤）"""
    # Supabase 模式：RLS 自动按 user_id 过滤
    repo = get_repo()
    if repo:
        books = repo.list_books()
        projects = []
        for b in (books or []):
            projects.append({
                'id': b['id'],
                'name': b.get('title', '未命名项目'),
                'description': b.get('description', ''),
                'genre': b.get('genre', ''),
                'created_at': b.get('created_at', ''),
                'updated_at': b.get('updated_at', '')
            })
        return jsonify(projects)

    # SQLite 模式（降级）
    projects = []
    for pdir in PROJECTS_DIR.iterdir():
        if pdir.is_dir() and (pdir / 'novel.db').exists():
            conn = get_db(pdir.name)
            info = conn.execute('SELECT key, value FROM project_info').fetchall()
            conn.close()
            info_dict = {row['key']: row['value'] for row in info}
            projects.append({
                'id': pdir.name,
                'name': info_dict.get('name', '未命名项目'),
                'created_at': info_dict.get('created_at', ''),
                'updated_at': info_dict.get('updated_at', '')
            })
    return jsonify(projects)

@app.route('/api/projects', methods=['POST'])
def create_project():
    """创建新项目"""
    # Feature Gating: 检查项目数限制
    if not check_feature('create_project'):
        tier = get_user_id()  # 复用 get_user_id 的 Supabase 逻辑
        limit = 50 if check_feature('premium') else 1
        return jsonify({'error': f'免费版最多创建{limit}个项目，请升级会员'}), 403

    data = request.json
    name = data.get('name', '未命名项目')
    desc = data.get('description', '')
    genre = data.get('genre', '')

    # Supabase 模式
    repo = get_repo()
    if repo:
        project_id = repo.create_book(title=name, description=desc, genre=genre)
        if project_id:
            return jsonify({'id': project_id, 'name': name})
        return jsonify({'error': '创建项目失败'}), 500

    # SQLite 模式（降级）
    project_id = secrets.token_hex(8)
    (PROJECTS_DIR / project_id).mkdir(exist_ok=True)

    init_project_db(project_id)
    now = datetime.now().isoformat()

    conn = get_db(project_id)
    conn.execute('INSERT INTO project_info VALUES (?, ?)', ('name', name))
    if desc:
        conn.execute('INSERT INTO project_info VALUES (?, ?)', ('description', desc))
    if genre:
        conn.execute('INSERT INTO project_info VALUES (?, ?)', ('genre', genre))
    conn.execute('INSERT INTO project_info VALUES (?, ?)', ('created_at', now))
    conn.execute('INSERT INTO project_info VALUES (?, ?)', ('updated_at', now))
    conn.commit()
    conn.close()

    return jsonify({'id': project_id, 'name': name})

@app.route('/api/projects/<project_id>', methods=['DELETE'])
def delete_project(project_id):
    """删除项目"""
    # Supabase 模式：RLS 确保只能删除自己的项目
    repo = get_repo()
    if repo:
        _, err = repo.delete_book(project_id)
        if err:
            return jsonify({'error': '删除失败'}), 500
        return jsonify({'success': True})

    # SQLite 模式（降级）
    import shutil
    project_path = PROJECTS_DIR / project_id
    if project_path.exists():
        shutil.rmtree(project_path)
    return jsonify({'success': True})

@app.route('/api/projects/<project_id>/info', methods=['GET'])
def get_project_info(project_id):
    """获取项目信息（含基础统计）"""
    repo = get_repo()
    if repo:
        book = repo.get_book(project_id)
        stats = repo.get_book_stats(project_id)
        if not book:
            return jsonify({'error': '项目不存在'}), 404
        return jsonify({
            'name': book.get('title', '未命名项目'),
            'description': book.get('description', ''),
            'genre': book.get('genre', ''),
            'created_at': book.get('created_at', ''),
            'updated_at': book.get('updated_at', ''),
            'total_words': stats.get('total_words', 0),
            'chapter_count': stats.get('chapter_count', 0),
        })

    conn = get_db(project_id)
    info = conn.execute('SELECT key, value FROM project_info').fetchall()
    total_words = conn.execute('SELECT COALESCE(SUM(word_count), 0) as total FROM chapters').fetchone()['total']
    chapter_count = conn.execute('SELECT COUNT(*) as cnt FROM chapters').fetchone()['cnt']
    conn.close()
    info_dict = {row['key']: row['value'] for row in info}
    info_dict['total_words'] = total_words
    info_dict['chapter_count'] = chapter_count
    return jsonify(info_dict)

@app.route('/api/projects/<project_id>/info', methods=['PUT'])
def update_project_info(project_id):
    """更新项目信息"""
    data = request.json

    repo = get_repo()
    if repo:
        update_data = {}
        if 'name' in data:
            update_data['title'] = data['name']
        if 'description' in data:
            update_data['description'] = data['description']
        if 'genre' in data:
            update_data['genre'] = data['genre']
        if update_data:
            repo.update_book(project_id, update_data)
        return jsonify({'success': True})

    conn = get_db(project_id)
    for key, value in data.items():
        existing = conn.execute('SELECT key FROM project_info WHERE key = ?', (key,)).fetchone()
        if existing:
            conn.execute('UPDATE project_info SET value = ? WHERE key = ?', (str(value), key))
        else:
            conn.execute('INSERT INTO project_info VALUES (?, ?)', (key, str(value)))
    conn.execute('UPDATE project_info SET value = ? WHERE key = ?',
                 (datetime.now().isoformat(), 'updated_at'))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

# ===== 角色管理 API =====
@app.route('/api/projects/<project_id>/characters', methods=['GET'])
def list_characters(project_id):
    repo = get_repo()
    if repo:
        return jsonify(repo.list_characters(project_id))
    conn = get_db(project_id)
    rows = conn.execute('SELECT * FROM characters ORDER BY created_at').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/projects/<project_id>/characters', methods=['POST'])
def create_character(project_id):
    data = request.json
    repo = get_repo()
    if repo:
        char_id = repo.create_character(project_id, data)
        if char_id:
            return jsonify({'id': char_id})
        return jsonify({'error': '创建失败'}), 500

    now = datetime.now().isoformat()
    char_id = secrets.token_hex(8)
    conn = get_db(project_id)
    conn.execute('''
        INSERT INTO characters (id, name, gender, age, personality, background, goal, appearance, notes, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        char_id, data.get('name', ''), data.get('gender', ''), data.get('age', ''),
        data.get('personality', ''), data.get('background', ''), data.get('goal', ''),
        data.get('appearance', ''), data.get('notes', ''), now, now
    ))
    conn.commit()
    conn.close()
    return jsonify({'id': char_id})

@app.route('/api/projects/<project_id>/characters/<char_id>', methods=['PUT'])
def update_character(project_id, char_id):
    data = request.json
    repo = get_repo()
    if repo:
        _, err = repo.update_character(char_id, data)
        if err:
            return jsonify({'error': '更新失败'}), 500
        return jsonify({'success': True})

    conn = get_db(project_id)
    conn.execute('''
        UPDATE characters SET name=?, gender=?, age=?, personality=?, background=?, 
        goal=?, appearance=?, notes=?, updated_at=?
        WHERE id=?
    ''', (
        data.get('name', ''), data.get('gender', ''), data.get('age', ''),
        data.get('personality', ''), data.get('background', ''), data.get('goal', ''),
        data.get('appearance', ''), data.get('notes', ''), datetime.now().isoformat(), char_id
    ))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

@app.route('/api/projects/<project_id>/characters/<char_id>', methods=['DELETE'])
def delete_character(project_id, char_id):
    repo = get_repo()
    if repo:
        _, err = repo.delete_character(char_id)
        if err:
            return jsonify({'error': '删除失败'}), 500
        return jsonify({'success': True})

    conn = get_db(project_id)
    conn.execute('DELETE FROM characters WHERE id = ?', (char_id,))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

# ===== 角色知识账本 API =====

@app.route('/api/projects/<project_id>/knowledge/<char_id>', methods=['GET'])
def get_character_knowledge(project_id, char_id):
    """获取某角色在所有章节的知识记录"""
    conn = get_db(project_id)
    rows = conn.execute('''
        SELECT ck.*, ch.title as chapter_title, ch.sort_order
        FROM character_knowledge ck
        JOIN chapters ch ON ck.chapter_id = ch.id
        WHERE ck.character_id = ?
        ORDER BY ch.sort_order
    ''', (char_id,)).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/projects/<project_id>/knowledge/<char_id>', methods=['POST'])
def save_character_knowledge(project_id, char_id):
    """保存/更新某角色在某章节的知识记录（upsert by character_id + chapter_id）"""
    data = request.json
    chapter_id = data.get('chapter_id', '')
    if not chapter_id:
        return jsonify({'error': 'chapter_id is required'}), 400

    conn = get_db(project_id)
    import uuid, datetime
    # 查找是否已有该角色在该章的记录
    existing = conn.execute(
        'SELECT id FROM character_knowledge WHERE character_id = ? AND chapter_id = ?',
        (char_id, chapter_id)
    ).fetchone()

    known_names = json.dumps(data.get('known_names', []), ensure_ascii=False)
    known_items = json.dumps(data.get('known_items', []), ensure_ascii=False)
    known_events = json.dumps(data.get('known_events', []), ensure_ascii=False)
    now = datetime.datetime.now().isoformat()

    if existing:
        conn.execute('''
            UPDATE character_knowledge
            SET known_names = ?, known_items = ?, known_events = ?, snapshot_at = ?
            WHERE id = ?
        ''', (known_names, known_items, known_events, now, existing['id']))
    else:
        new_id = str(uuid.uuid4())
        conn.execute('''
            INSERT INTO character_knowledge (id, character_id, chapter_id, known_names, known_items, known_events, snapshot_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (new_id, char_id, chapter_id, known_names, known_items, known_events, now))

    conn.commit()
    conn.close()
    return jsonify({'success': True})

# ===== 大纲管理 API =====
@app.route('/api/projects/<project_id>/outline', methods=['GET'])
def list_outline(project_id):
    repo = get_repo()
    if repo:
        return jsonify(repo.list_outline(project_id))
    conn = get_db(project_id)
    rows = conn.execute('SELECT * FROM outline ORDER BY sort_order, created_at').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/projects/<project_id>/outline', methods=['POST'])
def create_outline_node(project_id):
    data = request.json
    repo = get_repo()
    if repo:
        node_id = repo.create_outline_node(project_id, data)
        if node_id:
            return jsonify({'id': node_id})
        return jsonify({'error': '创建失败'}), 500

    now = datetime.now().isoformat()
    node_id = secrets.token_hex(8)
    conn = get_db(project_id)
    max_sort = conn.execute('SELECT COALESCE(MAX(sort_order), 0) as mx FROM outline WHERE parent_id = ?',
                           (data.get('parent_id'),)).fetchone()['mx']
    conn.execute('''
        INSERT INTO outline (id, parent_id, title, content, level, sort_order, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    ''', (node_id, data.get('parent_id'), data.get('title', '新节点'),
          data.get('content', ''), data.get('level', 0), max_sort + 1, now, now))
    conn.commit()
    conn.close()
    return jsonify({'id': node_id})

@app.route('/api/projects/<project_id>/outline/<node_id>', methods=['PUT'])
def update_outline_node(project_id, node_id):
    data = request.json
    repo = get_repo()
    if repo:
        _, err = repo.update_outline_node(node_id, data)
        if err:
            return jsonify({'error': '更新失败'}), 500
        return jsonify({'success': True})

    conn = get_db(project_id)
    fields = []
    values = []
    for key in ['title', 'content', 'parent_id', 'level', 'sort_order']:
        if key in data:
            fields.append(f'{key} = ?')
            values.append(data[key])
    fields.append('updated_at = ?')
    values.append(datetime.now().isoformat())
    values.append(node_id)
    conn.execute(f'UPDATE outline SET {", ".join(fields)} WHERE id = ?', values)
    conn.commit()
    conn.close()
    return jsonify({'success': True})

@app.route('/api/projects/<project_id>/outline/<node_id>', methods=['DELETE'])
def delete_outline_node(project_id, node_id):
    """递归删除大纲节点"""
    repo = get_repo()
    if repo:
        _, err = repo.delete_outline_node(node_id)
        if err:
            return jsonify({'error': '删除失败'}), 500
        return jsonify({'success': True})

    conn = get_db(project_id)
    def delete_recursive(nid):
        children = conn.execute('SELECT id FROM outline WHERE parent_id = ?', (nid,)).fetchall()
        for child in children:
            delete_recursive(child['id'])
        conn.execute('DELETE FROM outline WHERE id = ?', (nid,))
    delete_recursive(node_id)
    conn.commit()
    conn.close()
    return jsonify({'success': True})

# ===== 章节管理 API =====
@app.route('/api/projects/<project_id>/chapters', methods=['GET'])
def list_chapters(project_id):
    repo = get_repo()
    if repo:
        return jsonify(repo.list_chapters(project_id))
    conn = get_db(project_id)
    rows = conn.execute('SELECT * FROM chapters ORDER BY sort_order, created_at').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/projects/<project_id>/chapters', methods=['POST'])
def create_chapter(project_id):
    data = request.json
    repo = get_repo()
    if repo:
        chapter_id = repo.create_chapter(project_id, data.get('title', '新章节'))
        if chapter_id:
            return jsonify({'id': chapter_id})
        return jsonify({'error': '创建失败'}), 500

    now = datetime.now().isoformat()
    chapter_id = secrets.token_hex(8)
    conn = get_db(project_id)
    max_sort = conn.execute('SELECT COALESCE(MAX(sort_order), 0) as mx FROM chapters').fetchone()['mx']
    conn.execute('''
        INSERT INTO chapters (id, title, content, sort_order, created_at, updated_at, word_count)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (chapter_id, data.get('title', '新章节'), '', max_sort + 1, now, now, 0))
    conn.commit()
    conn.close()
    return jsonify({'id': chapter_id})

@app.route('/api/projects/<project_id>/chapters/<chapter_id>', methods=['PUT'])
def update_chapter(project_id, chapter_id):
    data = request.json
    content = data.get('content', '')
    word_count = len(content.replace(' ', '').replace('\n', ''))

    repo = get_repo()
    if repo:
        _, err = repo.update_chapter(chapter_id, {'title': data.get('title'), 'content': content})
        if err:
            return jsonify({'error': '更新失败'}), 500
        return jsonify({'success': True, 'word_count': word_count})

    conn = get_db(project_id)

    # 自动保存版本快照（仅当内容变化超过50字时）
    old = conn.execute('SELECT title, content, word_count FROM chapters WHERE id = ?', (chapter_id,)).fetchone()
    if old and content:
        old_content = old['content'] or ''
        if abs(len(content) - len(old_content)) > 50:
            max_ver = conn.execute(
                'SELECT COALESCE(MAX(version), 0) FROM chapter_snapshots WHERE chapter_id = ?',
                (chapter_id,)
            ).fetchone()[0]
            now = datetime.now().isoformat()
            snap_id = secrets.token_hex(8)
            conn.execute('''
                INSERT INTO chapter_snapshots (id, chapter_id, version, title, content, word_count, snapshot_at)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (snap_id, chapter_id, max_ver + 1, old['title'], old_content,
                  old['word_count'] or 0, now))

    fields = []
    values = []
    for key in ['title', 'content']:
        if key in data:
            fields.append(f'{key} = ?')
            values.append(data[key])
    fields.append('word_count = ?')
    values.append(word_count)
    fields.append('updated_at = ?')
    values.append(datetime.now().isoformat())
    values.append(chapter_id)
    conn.execute(f'UPDATE chapters SET {", ".join(fields)} WHERE id = ?', values)
    conn.commit()
    conn.close()
    return jsonify({'success': True, 'word_count': word_count})

@app.route('/api/projects/<project_id>/chapters/<chapter_id>/snapshots', methods=['GET'])
def list_chapter_snapshots(project_id, chapter_id):
    repo = get_repo()
    if repo:
        return jsonify(repo.list_chapter_snapshots(chapter_id))

    conn = get_db(project_id)
    snaps = conn.execute('''
        SELECT id, version, title, word_count, snapshot_at
        FROM chapter_snapshots WHERE chapter_id = ?
        ORDER BY version DESC LIMIT 30
    ''', (chapter_id,)).fetchall()
    conn.close()
    return jsonify([dict(s) for s in snaps])

@app.route('/api/projects/<project_id>/chapters/<chapter_id>/snapshots/<snap_id>', methods=['GET'])
def get_chapter_snapshot(project_id, chapter_id, snap_id):
    repo = get_repo()
    if repo:
        snap = repo.get_snapshot(snap_id)
        if not snap:
            return jsonify({'error': '快照不存在'}), 404
        return jsonify(snap)

    conn = get_db(project_id)
    snap = conn.execute(
        'SELECT * FROM chapter_snapshots WHERE id = ? AND chapter_id = ?',
        (snap_id, chapter_id)
    ).fetchone()
    conn.close()
    if not snap:
        return jsonify({'error': '快照不存在'}), 404
    return jsonify(dict(snap))

@app.route('/api/projects/<project_id>/chapters/<chapter_id>/snapshots/<snap_id>/revert', methods=['POST'])
def revert_chapter_snapshot(project_id, chapter_id, snap_id):
    """回退到指定版本（创建当前版本快照后恢复）"""
    repo = get_repo()
    if repo:
        snap = repo.get_snapshot(snap_id)
        if not snap:
            return jsonify({'error': '快照不存在'}), 404
        repo.update_chapter(chapter_id, {
            'title': snap['title'],
            'content': snap['content']
        })
        return jsonify({'success': True, 'title': snap['title'],
                        'word_count': snap.get('word_count', 0),
                        'version': snap.get('version', 0)})

    conn = get_db(project_id)
    snap = conn.execute(
        'SELECT * FROM chapter_snapshots WHERE id = ? AND chapter_id = ?',
        (snap_id, chapter_id)
    ).fetchone()
    if not snap:
        conn.close()
        return jsonify({'error': '快照不存在'}), 404
    current = conn.execute('SELECT title, content, word_count FROM chapters WHERE id = ?', (chapter_id,)).fetchone()
    if current:
        max_ver = conn.execute(
            'SELECT COALESCE(MAX(version), 0) FROM chapter_snapshots WHERE chapter_id = ?',
            (chapter_id,)
        ).fetchone()[0]
        now = datetime.now().isoformat()
        conn.execute('''
            INSERT INTO chapter_snapshots (id, chapter_id, version, title, content, word_count, snapshot_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (secrets.token_hex(8), chapter_id, max_ver + 1, current['title'],
              current['content'] or '', current['word_count'] or 0, now))
    conn.execute('''
        UPDATE chapters SET title = ?, content = ?, word_count = ?, updated_at = ?
        WHERE id = ?
    ''', (snap['title'], snap['content'], snap['word_count'], datetime.now().isoformat(), chapter_id))
    conn.commit()
    conn.close()
    return jsonify({'success': True, 'title': snap['title'], 'word_count': snap['word_count'],
                    'version': snap['version']})

@app.route('/api/projects/<project_id>/chapters/<chapter_id>', methods=['DELETE'])
def delete_chapter(project_id, chapter_id):
    repo = get_repo()
    if repo:
        _, err = repo.delete_chapter(chapter_id)
        if err:
            return jsonify({'error': '删除失败'}), 500
        return jsonify({'success': True})

    conn = get_db(project_id)
    conn.execute('DELETE FROM chapters WHERE id = ?', (chapter_id,))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

@app.route('/api/projects/<project_id>/chapters/reorder', methods=['POST'])
def reorder_chapters(project_id):
    """重新排序章节"""
    data = request.json
    repo = get_repo()
    if repo:
        repo.reorder_chapters(data.get('chapter_ids', []))
        return jsonify({'success': True})

    conn = get_db(project_id)
    for idx, chapter_id in enumerate(data.get('chapter_ids', [])):
        conn.execute('UPDATE chapters SET sort_order = ? WHERE id = ?', (idx, chapter_id))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

# ===== 写作统计 API =====
@app.route('/api/projects/<project_id>/stats', methods=['GET'])
def get_writing_stats(project_id):
    date = request.args.get('date', datetime.now().strftime('%Y-%m-%d'))
    repo = get_repo()
    if repo:
        stats = repo.get_writing_stats(project_id, date)
        if not stats:
            stats = {'date': date, 'chars_added': 0, 'chars_deleted': 0, 'time_spent': 0, 'sessions': 0}
        book_stats = repo.get_book_stats(project_id)
        return jsonify({**stats, 'total_words': book_stats.get('total_words', 0),
                        'chapter_count': book_stats.get('chapter_count', 0)})

    conn = get_db(project_id)
    stats = conn.execute('SELECT * FROM writing_stats WHERE date = ?', (date,)).fetchone()
    if not stats:
        stats = {'date': date, 'chars_added': 0, 'chars_deleted': 0, 'time_spent': 0, 'sessions': 0}
    else:
        stats = dict(stats)
    total_words = conn.execute('SELECT COALESCE(SUM(word_count), 0) as total FROM chapters').fetchone()['total']
    chapter_count = conn.execute('SELECT COUNT(*) as cnt FROM chapters').fetchone()['cnt']
    conn.close()
    return jsonify({**stats, 'total_words': total_words, 'chapter_count': chapter_count})

@app.route('/api/projects/<project_id>/stats', methods=['POST'])
def update_writing_stats(project_id):
    data = request.json
    date = data.get('date', datetime.now().strftime('%Y-%m-%d'))

    repo = get_repo()
    if repo:
        _, err = repo.upsert_writing_stats(project_id, {
            'date': date,
            'chars_added': data.get('chars_added', 0),
            'chars_deleted': data.get('chars_deleted', 0),
            'time_spent': data.get('time_spent', 0),
            'sessions': data.get('sessions', 0),
        })
        if err:
            return jsonify({'error': '更新失败'}), 500
        return jsonify({'success': True})

    conn = get_db(project_id)
    existing = conn.execute('SELECT date FROM writing_stats WHERE date = ?', (date,)).fetchone()
    if existing:
        for key in ['chars_added', 'chars_deleted', 'time_spent', 'sessions']:
            if key in data:
                conn.execute(f'UPDATE writing_stats SET {key} = {key} + ? WHERE date = ?', (data[key], date))
    else:
        conn.execute('''
            INSERT INTO writing_stats (date, chars_added, chars_deleted, time_spent, sessions)
            VALUES (?, ?, ?, ?, ?)
        ''', (date, data.get('chars_added', 0), data.get('chars_deleted', 0),
              data.get('time_spent', 0), data.get('sessions', 0)))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

# ===== 写作目标 API =====
@app.route('/api/projects/<project_id>/goals', methods=['GET'])
def list_goals(project_id):
    repo = get_repo()
    if repo:
        return jsonify(repo.list_goals(project_id))

    conn = get_db(project_id)
    rows = conn.execute('SELECT * FROM writing_goals WHERE is_active = 1').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/projects/<project_id>/goals', methods=['POST'])
def create_goal(project_id):
    data = request.json

    repo = get_repo()
    if repo:
        goal_id = repo.create_goal(project_id, data)
        if goal_id:
            return jsonify({'id': goal_id})
        return jsonify({'error': '创建失败'}), 500

    goal_id = secrets.token_hex(8)
    conn = get_db(project_id)
    conn.execute('''
        INSERT INTO writing_goals (id, goal_type, target_value, current_value, deadline, is_active)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', (goal_id, data.get('goal_type', 'daily'), data.get('target_value', 500),
          0, data.get('deadline', ''), 1))
    conn.commit()
    conn.close()
    return jsonify({'id': goal_id})

@app.route('/api/projects/<project_id>/goals/<goal_id>', methods=['PUT'])
def update_goal(project_id, goal_id):
    data = request.json

    repo = get_repo()
    if repo:
        _, err = repo.update_goal(goal_id, data)
        if err:
            return jsonify({'error': '更新失败'}), 500
        return jsonify({'success': True})

    conn = get_db(project_id)
    for key in ['goal_type', 'target_value', 'current_value', 'deadline', 'is_active']:
        if key in data:
            conn.execute(f'UPDATE writing_goals SET {key} = ? WHERE id = ?', (data[key], goal_id))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

# ===== 长篇上下文管理（RAG 设定记忆系统） =====

def _extract_keywords(content, max_keywords=15):
    """从当前内容中提取关键词，用于 RAG 检索。
    基于中文分词简单策略：按常见分隔符切分，提取2-4字词，去重去停用词。
    """
    import re
    # 停用词（高频无意义词）
    stop_words = {'的是','不是','一个','可以','没有','自己','他们','我们','什么','怎么',
                  '这个','那个','因为','所以','但是','如果','虽然','已经','还是','只是',
                  '不过','就是','的话','或者','而且','然后','之后','之前','时候','地方',
                  '一点','一下','有些','这里','那里','一种','觉得','知道','看到','听到',
                  '说道','突然','立刻','一定','一样','这样','那样','出来','起来','过来',
                  '过去','下来','上去','进','出','来','去','在','是','有','的','了','着',
                  '我','你','他','她','它','都','也','很','就','要','会','能','把','被',
                  '让','给','向','从','对','和','与','但','却','只','还','更','最','又'}
    # 从中文文本中提取 2-4 字的词
    cleaned = re.sub(r'[^一-鿿]', ' ', content)
    words = []
    for seg in cleaned.split():
        seg = seg.strip()
        if len(seg) >= 2 and seg not in stop_words:
            words.append(seg)
        # 也尝试 2-3 字滑动窗口
        for wlen in [2, 3]:
            for i in range(len(seg) - wlen + 1):
                w = seg[i:i+wlen]
                if w not in stop_words and len(w) >= 2:
                    words.append(w)
    # 按频率排序取 top N
    from collections import Counter
    word_counts = Counter(words)
    top_words = [w for w, _ in word_counts.most_common(max_keywords)]
    return top_words

def _search_rag_context(project_id, keywords, conn=None):
    """基于关键词从 SQLite 中检索相关上下文。
    检索范围：章节摘要、角色、情节线程、关键事件。
    Supabase 模式：简化为返回角色档案和近期摘要（不做关键词精准检索）。
    返回：结构化的上下文字符串。
    """
    # Supabase 模式：简化检索
    repo = get_repo()
    if repo:
        if not keywords:
            return ''
        parts = []
        # 获取所有角色概要
        all_chars = repo.list_characters(project_id)
        if all_chars:
            # 按关键词过滤角色
            matched_chars = []
            for c in all_chars:
                name = c.get('name', '')
                personality = c.get('personality', '') or ''
                background = c.get('background', '') or ''
                goal = c.get('goal', '') or ''
                combined = f'{name} {personality} {background} {goal}'
                if any(kw in combined for kw in keywords[:5]):
                    matched_chars.append(c)
            if matched_chars:
                parts.append('\n【相关角色】')
                for c in matched_chars[:5]:
                    detail = c.get('name', '')
                    if c.get('personality'):
                        detail += f'（{c["personality"][:60]}）'
                    if c.get('goal'):
                        detail += f' 目标：{c["goal"][:50]}'
                    parts.append(f'- {detail}')
            elif all_chars:
                parts.append('\n【角色档案】')
                for c in all_chars[:10]:
                    detail = c.get('name', '')
                    if c.get('personality'):
                        detail += f'（{c["personality"][:60]}）'
                    if c.get('goal'):
                        detail += f' 目标：{c["goal"][:50]}'
                    parts.append(f'- {detail}')

        # 获取近期摘要
        summaries = repo.list_chapter_summaries(project_id, limit=3)
        chapters_map = {}
        if summaries:
            chapters_list = repo.list_chapters(project_id)
            chapters_map = {c['id']: c for c in chapters_list} if chapters_list else {}
            parts.append('【相关前文摘要】')
            for s in reversed(summaries):
                ch = chapters_map.get(s.get('chapter_id', ''), {})
                title = ch.get('title', f'章节{s.get("chapter_id", "")}')
                parts.append(f'《{title}》：{s.get("summary", "")[:200]}')

        return '\n'.join(parts) if parts else ''

    own_conn = conn is None
    if own_conn:
        conn = get_db(project_id)

    if not keywords:
        if own_conn: conn.close()
        return ''

    parts = []

    # 构建 LIKE 查询条件（每个关键词用 OR 连接）
    like_clauses = ' OR '.join(['(summary LIKE ? OR key_events LIKE ? OR character_states LIKE ?)'
                                for _ in keywords[:8]])  # 最多 8 个关键词避免查询过长
    like_params = []
    for kw in keywords[:8]:
        like_params.extend([f'%{kw}%', f'%{kw}%', f'%{kw}%'])

    # 1. 检索相关章节摘要
    if like_clauses:
        summaries = conn.execute(f'''
            SELECT cs.summary, cs.key_events, cs.character_states,
                   c.title, c.sort_order
            FROM chapter_summaries cs JOIN chapters c ON cs.chapter_id = c.id
            WHERE {like_clauses}
            ORDER BY c.sort_order DESC LIMIT 6
        ''', like_params).fetchall()

        if summaries:
            parts.append('【相关前文摘要】')
            for s in reversed(summaries):
                parts.append(f'《{s["title"]}》：{s["summary"][:200]}')

    # 2. 检索相关角色（通过关键词匹配角色名/性格/背景）
    char_like = ' OR '.join(['(name LIKE ? OR personality LIKE ? OR background LIKE ?)'
                              for _ in keywords[:5]])
    char_params = []
    for kw in keywords[:5]:
        char_params.extend([f'%{kw}%', f'%{kw}%', f'%{kw}%'])
    chars = conn.execute(f'''
        SELECT DISTINCT c.*, cs.status, cs.location, cs.emotional_state, cs.knowledge_gained
        FROM characters c
        LEFT JOIN character_states cs ON c.id = cs.character_id
            AND cs.id = (SELECT id FROM character_states cs2 WHERE cs2.character_id = c.id ORDER BY cs2.snapshot_at DESC LIMIT 1)
        WHERE {char_like}
    ''', char_params).fetchall() if keywords else []

    # 也获取所有角色的概要（始终包含）
    all_chars = conn.execute('SELECT id, name, personality, background, goal FROM characters').fetchall()
    if all_chars and not chars:
        parts.append('\n【角色档案】')
        for c in all_chars[:10]:
            detail = f'{c["name"]}'
            if c['personality']: detail += f'（{c["personality"][:60]}）'
            if c['goal']: detail += f' 目标：{c["goal"][:50]}'
            parts.append(f'- {detail}')
    elif chars:
        parts.append('\n【相关角色】')
        for c in chars[:5]:
            detail = f'{c["name"]}'
            if c['personality']: detail += f'（{c["personality"][:60]}）'
            if 'emotional_state' in c.keys() and c['emotional_state']:
                detail += f' 情绪：{c["emotional_state"]}'
            parts.append(f'- {detail}')

    # 3. 检索相关情节线程
    if like_clauses:
        thread_like = ' OR '.join(['(title LIKE ? OR description LIKE ?)' for _ in keywords[:5]])
        thread_params = []
        for kw in keywords[:5]:
            thread_params.extend([f'%{kw}%', f'%{kw}%'])
        threads = conn.execute(f'''
            SELECT * FROM plot_threads WHERE status = 'active' AND ({thread_like})
        ''', thread_params).fetchall()
        if threads:
            parts.append('\n【相关情节线程】')
            for t in threads:
                parts.append(f'- {t["title"]}（{t["thread_type"]}）：{t["description"][:120]}')

    # 4. 最近关键事件（始终包含最新 5 个）
    events = conn.execute('''
        SELECT ke.*, c.title as chapter_title
        FROM key_events ke JOIN chapters c ON ke.chapter_id = c.id
        ORDER BY ke.sort_order DESC LIMIT 5
    ''').fetchall()
    if events:
        parts.append('\n【近期关键事件】')
        for ev in reversed(events):
            parts.append(f'- [{ev["chapter_title"]}] {ev["title"]}：{ev["description"][:100]}')

    if own_conn:
        conn.close()

    return '\n'.join(parts) if parts else ''

# ===== 剧情记忆与大纲锚定引擎 =====
def build_condensed_context(project_id, current_chapter_id=None):
    """构建精简上下文包，替代全量 build_context_package。

    数据包结构（严格控制在 ~1500 字以内）：
    1. [当前卷全局核心大纲] — 当前卷的 level-0/1 大纲节点
    2. [前一章极致缩略梗概] — ≤500 字
    3. [当前章节细化线索] — 当前章的 level-2/3 大纲子节点
    4. [本章登场人物极简标签] — 仅核心角色，每人 ≤60 字
    """
    # Supabase 模式
    repo = get_repo()
    if repo:
        parts = []
        current_chapter_title = None
        current_chapter_content = ''
        current_chapter_sort = 0

        if current_chapter_id:
            ch = repo.get_chapter(current_chapter_id)
            if ch:
                current_chapter_title = ch.get('title', '')
                current_chapter_content = ch.get('content') or ''
                current_chapter_sort = ch.get('sort_order', 0)

        # 2. 当前卷核心大纲
        outline_nodes = repo.list_outline(project_id)
        volumes = [n for n in outline_nodes if n.get('level') == 0]
        chapters_outline = [n for n in outline_nodes if n.get('level') == 1]
        sub_nodes_map = {}
        for n in outline_nodes:
            if n.get('level', 0) >= 2 and n.get('parent_id'):
                sub_nodes_map.setdefault(n['parent_id'], []).append(n)

        if volumes:
            active_volume = None
            for v in volumes:
                v_chapters = [c for c in chapters_outline if c.get('parent_id') == v.get('id')]
                for vc in v_chapters:
                    if current_chapter_title and (
                        current_chapter_title in vc.get('title', '')
                        or vc.get('title', '') in (current_chapter_title or '')
                    ):
                        active_volume = v
                        break
                if active_volume:
                    break
            if not active_volume:
                active_volume = volumes[0] if volumes else None

            if active_volume:
                v_chapters = [c for c in chapters_outline if c.get('parent_id') == active_volume.get('id')]
                outline_lines = [f'【当前卷核心大纲】{active_volume.get("title", "")}']
                for vc in v_chapters:
                    marker = ' ← 当前' if (current_chapter_title and (
                        current_chapter_title in vc.get('title', '')
                        or vc.get('title', '') in (current_chapter_title or '')
                    )) else ''
                    outline_lines.append(f'  {vc.get("title", "")}{marker}')
                parts.append('\n'.join(outline_lines))

        # 3. 前一章梗概
        if current_chapter_id and current_chapter_sort > 1:
            chapters_list = repo.list_chapters(project_id)
            prev_ch = None
            for c in chapters_list:
                if c.get('sort_order', 0) == current_chapter_sort - 1:
                    prev_ch = c
                    break
            if prev_ch:
                summary_data = repo.get_chapter_summary(prev_ch['id'])
                summary = summary_data.get('summary', '') if summary_data else ''
                if not summary:
                    # 即时生成摘要
                    content = prev_ch.get('content') or ''
                    if content:
                        cleaned = content.replace('\n', ' ').replace('\r', ' ').strip()
                        summary = cleaned[:500] if len(cleaned) <= 500 else f'{cleaned[:200]}……{cleaned[-200:]}'
                    else:
                        summary = '（本章暂无内容）'
                parts.append(f'【前一章缩略梗概】{prev_ch.get("title", "")}\n{summary[:500]}')

        # 4. 当前章节线索
        if current_chapter_title:
            outline_node = None
            for n in chapters_outline:
                if n.get('title') and (
                    current_chapter_title in n.get('title', '')
                    or n.get('title', '') in current_chapter_title
                ):
                    outline_node = n
                    break
            if outline_node:
                sub_nodes = sub_nodes_map.get(outline_node['id'], [])
                if sub_nodes:
                    clue_lines = [f'【当前章节线索】{outline_node.get("title", "")}']
                    for sn in sub_nodes:
                        prefix = '  ◦' if sn.get('level') == 3 else '•'
                        clue_lines.append(f'  {prefix} {sn.get("title", "")}')
                        if sn.get('content'):
                            clue_lines.append(f'    {sn["content"][:120]}')
                    parts.append('\n'.join(clue_lines))
                elif outline_node.get('content'):
                    parts.append(f'【当前章节线索】{outline_node.get("title", "")}\n  {outline_node["content"][:300]}')

        # 5. 登场人物标签
        char_tags = _build_character_tags(project_id, current_chapter_id)
        if char_tags:
            parts.append(f'【登场人物设定——必须严格遵守】\n{char_tags}')

        return '\n\n'.join(parts) if parts else ''

    conn = get_db(project_id)
    parts = []

    # 1. 确定当前章节所在的卷
    current_volume_title = None
    current_chapter_title = None
    current_chapter_content = ''
    if current_chapter_id:
        ch = conn.execute('SELECT title, content, sort_order FROM chapters WHERE id = ?',
                          (current_chapter_id,)).fetchone()
        if ch:
            current_chapter_title = ch['title']
            current_chapter_content = ch['content'] or ''

    # 2. [当前卷全局核心大纲] — level-0 卷节点 + level-1 章节点
    volumes = conn.execute(
        'SELECT * FROM outline WHERE level = 0 ORDER BY sort_order'
    ).fetchall()

    if volumes:
        # 找到当前章节所属的卷（通过标题/排序推断）
        active_volume = None
        for v in volumes:
            v_chapters = conn.execute(
                'SELECT * FROM outline WHERE parent_id = ? AND level = 1 ORDER BY sort_order',
                (v['id'],)
            ).fetchall()
            for vc in v_chapters:
                if current_chapter_title and (
                    current_chapter_title in vc['title']
                    or vc['title'] in (current_chapter_title or '')
                ):
                    active_volume = v
                    break
            if active_volume:
                break

        if not active_volume and volumes:
            # 默认取第一个卷
            active_volume = volumes[0]

        if active_volume:
            vol_chapters = conn.execute(
                'SELECT * FROM outline WHERE parent_id = ? AND level = 1 ORDER BY sort_order',
                (active_volume['id'],)
            ).fetchall()

            outline_lines = [f'【当前卷核心大纲】{active_volume["title"]}']
            for vc in vol_chapters:
                marker = ' ← 当前' if (current_chapter_title and (
                    current_chapter_title in vc['title']
                    or vc['title'] in (current_chapter_title or '')
                )) else ''
                outline_lines.append(f'  {vc["title"]}{marker}')
            parts.append('\n'.join(outline_lines))

    # 3. [前一章极致缩略梗概] — 取紧邻的前一章
    if current_chapter_id:
        prev_ch = conn.execute('''
            SELECT id, title, content FROM chapters
            WHERE sort_order < (SELECT sort_order FROM chapters WHERE id = ?)
            ORDER BY sort_order DESC LIMIT 1
        ''', (current_chapter_id,)).fetchone()

        if prev_ch:
            summary = _get_or_generate_summary(conn, prev_ch)
            parts.append(f'【前一章缩略梗概】{prev_ch["title"]}\n{summary[:500]}')

    # 4. [当前章节细化线索] — 匹配大纲中的子节点
    if current_chapter_title:
        # 查找当前章节对应的大纲节点
        outline_node = conn.execute(
            'SELECT id, title FROM outline WHERE level = 1 AND ? LIKE \'%\' || title || \'%\'',
            (current_chapter_title,)
        ).fetchone()
        if not outline_node:
            outline_node = conn.execute(
                'SELECT id, title FROM outline WHERE level = 1 AND title LIKE \'%\' || ? || \'%\'',
                (current_chapter_title,)
            ).fetchone()

        if outline_node:
            sub_nodes = conn.execute('''
                SELECT * FROM outline WHERE parent_id = ? AND level >= 2 ORDER BY sort_order
            ''', (outline_node['id'],)).fetchall()

            if sub_nodes:
                clue_lines = [f'【当前章节线索】{outline_node["title"]}']
                for sn in sub_nodes:
                    prefix = '  ◦' if sn['level'] == 3 else '•'
                    clue_lines.append(f'  {prefix} {sn["title"]}')
                    if sn.get('content'):
                        clue_lines.append(f'    {sn["content"][:120]}')
                parts.append('\n'.join(clue_lines))
            else:
                # 没有子节点时，用大纲节点自身的内容
                if outline_node.get('content'):
                    parts.append(f'【当前章节线索】{outline_node["title"]}\n  {outline_node["content"][:300]}')

    # 5. [本章登场人物极简标签]
    char_tags = _build_character_tags(project_id, current_chapter_id, conn)
    if char_tags:
        parts.append(f'【登场人物设定——必须严格遵守】\n{char_tags}')

    conn.close()
    return '\n\n'.join(parts) if parts else ''


def _get_or_generate_summary(conn, chapter_row, max_chars=500):
    """获取已有摘要或即时生成缩略梗概"""
    existing = conn.execute(
        'SELECT summary FROM chapter_summaries WHERE chapter_id = ?',
        (chapter_row['id'],)
    ).fetchone()
    if existing and existing['summary']:
        text = existing['summary']
        return text[:max_chars] + ('…' if len(text) > max_chars else '')

    # 即时生成：取章节前 200 字 + 后 200 字 + 中间关键句
    content = chapter_row['content'] or ''
    if not content.strip():
        return '（本章暂无内容）'

    cleaned = content.replace('\n', ' ').replace('\r', ' ').strip()
    if len(cleaned) <= 500:
        return cleaned

    # 提取式摘要：开头 + 结尾
    head = cleaned[:200]
    tail = cleaned[-200:] if len(cleaned) > 400 else ''
    return f'{head}……{tail}' if tail else head


def _build_character_tags(project_id, current_chapter_id, conn=None):
    """构建登场人物上下文标签：每人姓名 + 性格 + 外貌 + 背景 + 目标，确保AI不写偏角色"""
    repo = get_repo()
    if repo:
        chars = repo.list_characters(project_id)
        if not chars:
            return ''
        tags = []
        for c in chars[:6]:
            parts = [f'{c.get("name", "")}']
            if c.get('personality'):
                parts.append(f'性格：{c["personality"].replace(chr(10), " ").strip()[:80]}')
            if c.get('appearance'):
                parts.append(f'外貌：{c["appearance"].replace(chr(10), " ").strip()[:60]}')
            if c.get('background'):
                parts.append(f'背景：{c["background"].replace(chr(10), " ").strip()[:80]}')
            if c.get('goal'):
                parts.append(f'目标：{c["goal"].replace(chr(10), " ").strip()[:50]}')
            tags.append('- ' + ' | '.join(parts))
        return '\n'.join(tags)

    own_conn = False
    if conn is None:
        conn = get_db(project_id)
        own_conn = True

    chars = conn.execute(
        'SELECT name, personality, goal, background, appearance FROM characters ORDER BY created_at'
    ).fetchall()

    if not chars:
        if own_conn:
            conn.close()
        return ''

    tags = []
    for c in chars[:6]:
        parts = [f'{c["name"]}']
        if c['personality']:
            parts.append(f'性格：{c["personality"].replace(chr(10), " ").strip()[:80]}')
        if c['appearance']:
            parts.append(f'外貌：{c["appearance"].replace(chr(10), " ").strip()[:60]}')
        if c['background']:
            parts.append(f'背景：{c["background"].replace(chr(10), " ").strip()[:80]}')
        if c['goal']:
            parts.append(f'目标：{c["goal"].replace(chr(10), " ").strip()[:50]}')
        tags.append('- ' + ' | '.join(parts))

    if own_conn:
        conn.close()
    return '\n'.join(tags)


def _extract_outline_keywords(project_id):
    """从大纲中提取关键词列表，用于偏离度检测"""
    # Supabase 模式
    repo = get_repo()
    if repo:
        nodes = repo.list_outline(project_id)
        keywords = set()
        for n in nodes:
            if n.get('level', 0) < 1:
                continue
            text = (n.get('title', '') + ' ' + (n.get('content') or ''))
            import re
            words = re.split(r'[，。、；：！？\s,\.;:!?\n]+', text)
            for w in words:
                w = w.strip()
                if 2 <= len(w) <= 12 and w not in ('本章', '章节', '情节', '故事', '一个', '这个', '那个'):
                    keywords.add(w)
        return list(keywords)

    conn = get_db(project_id)
    nodes = conn.execute(
        'SELECT title, content FROM outline WHERE level >= 1 ORDER BY sort_order'
    ).fetchall()
    conn.close()

    keywords = set()
    for n in nodes:
        # 从标题和内容中提取名词/关键词
        text = (n['title'] + ' ' + (n.get('content') or ''))
        # 简单分词：按常见分隔符拆分，过滤短词
        import re
        words = re.split(r'[，。、；：！？\s,\.;:!?\n]+', text)
        for w in words:
            w = w.strip()
            if 2 <= len(w) <= 12 and w not in ('本章', '章节', '情节', '故事', '一个', '这个', '那个'):
                keywords.add(w)
    return list(keywords)


def _quick_quality_scan(content, project_id, chapter_id=None):
    """AI 生成后轻量级质量扫描，检查角色/大纲/逻辑一致性。
    返回警告列表，无问题则返回空列表。
    """
    warnings = []
    if not content or not content.strip():
        return warnings

    repo = get_repo()
    if repo:
        # 1. 角色名检查：大纲中提及的主要角色是否在生成内容中出现
        chars = repo.list_characters(project_id)
        if chars:
            core_chars = [c for c in chars if c.get('is_core')][:3] or chars[:3]
            for c in core_chars:
                name = c.get('name', '')
                if name and len(name) >= 2 and name not in content:
                    warnings.append(f'核心角色「{name}」未在续写内容中出现')

        # 2. 大纲关键词匹配检查
        if chapter_id:
            keywords = _extract_outline_keywords(project_id)
            if keywords:
                hit_count = sum(1 for kw in keywords if kw in content)
                if hit_count < max(1, len(keywords) * 0.3):
                    warnings.append(f'大纲关键词命中率低（{hit_count}/{len(keywords)}），可能偏离主线')

    return warnings


@app.route('/api/projects/<project_id>/ai/check-plot-deviation', methods=['POST'])
def api_check_plot_deviation(project_id):
    """剧情偏离度检测：对比当前正文与大纲关键词匹配度"""
    # Token 精算：QPS 频率限制
    if not check_rate_limit(get_user_id(), 'ai/dashboard'):
        return jsonify({'error': '请求过于频繁，请稍后再试'}), 429
    # SaaS 会员阶梯门控
    if not check_feature('plot_deviation'):
        return jsonify({'error': '此功能需要升级会员', 'code': 'tier_required'}), 403

    data = request.json
    content = data.get('content', '')
    chapter_id = data.get('chapter_id', '')

    if not content.strip():
        return jsonify({'deviation': 0, 'status': 'idle', 'message': '无内容可检测'})

    # 获取大纲关键词
    outline_keywords = _extract_outline_keywords(project_id)
    if not outline_keywords:
        return jsonify({'deviation': 0, 'status': 'idle', 'message': '无大纲关键词，请先生成大纲'})

    # 匹配度计算
    content_lower = content.lower()
    matched = [kw for kw in outline_keywords if kw.lower() in content_lower]
    match_rate = len(matched) / len(outline_keywords) if outline_keywords else 0

    # 偏离度 = 1 - 匹配率（越高越危险）
    deviation = round(max(0, 1 - match_rate) * 100)

    # 分级预警
    if deviation <= 20:
        status = 'safe'
        message = '✓ 内容与大纲高度吻合'
        color = 'var(--green)'
    elif deviation <= 50:
        status = 'watch'
        message = '⚠ 剧情有轻微偏离，建议回顾大纲'
        color = 'var(--orange)'
    else:
        status = 'danger'
        message = '🚨 剧情偏离主线！请检查大纲锚点'
        color = 'var(--red)'

    return jsonify({
        'deviation': deviation,
        'status': status,
        'message': message,
        'color': color,
        'matched_keywords': matched[:10],
        'total_keywords': len(outline_keywords),
        'match_rate': round(match_rate * 100)
    })


# ===== 有限视角硬核巡逻员 =====
# POV 心理活动词库 — 任何角色使用这些词意味着"进入该角色内心"
POV_MIND_VERBS = [
    '暗想', '心想', '心道', '心说', '暗忖', '寻思', '暗自', '心底',
    '心中一惊', '心中暗道', '心头一凛', '心里暗暗', '心中暗',
    '心头涌起', '心念一动', '心中冷笑', '心中嗤笑',
    '默默想着', '心里盘算', '暗喜', '窃喜', '暗怒',
]
POV_GOD_VIEW_MARKERS = [
    '殊不知', '却不知', '并未察觉', '没注意到', '浑然不觉',
    '尚不知', '并不知道', '完全没发现', '丝毫没有察觉',
]
POV_INNER_STATE_WORDS = [
    '冷笑', '心中', '心底', '内心深处', '暗自', '暗地',
]

def audit_pov_text(content, pov_character, all_characters):
    """硬核视角审计：逐行扫描正文，检测非 POV 角色的心理越界。

    Args:
        content: 正文文本
        pov_character: 当前 POV 角色名
        all_characters: 所有角色名列表

    Returns:
        dict: {violations: [{line, text, reason, severity}], total, safe}
    """
    if not content or not pov_character:
        return {'violations': [], 'total': 0, 'safe': True}

    lines = content.split('\n')
    violations = []
    pov_lower = pov_character.lower()
    # 构建非 POV 角色集合
    other_chars = [c for c in all_characters if c.lower() != pov_lower]

    for line_idx, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        line_lower = stripped.lower()

        # 1. 检查上帝视角标记（无论主语是谁，都是越界）
        for marker in POV_GOD_VIEW_MARKERS:
            if marker in stripped:
                violations.append({
                    'line': line_idx + 1,
                    'text': stripped[:80] + ('…' if len(stripped) > 80 else ''),
                    'reason': f'上帝视角：「{marker}」透露出非主角所知信息',
                    'severity': 'high',
                })
                break  # 一行只报一次

        # 2. 检查非 POV 角色 + 心理动词组合
        for other in other_chars:
            if other not in stripped:
                continue
            # 找到该角色名在行中的位置
            char_pos = stripped.find(other)
            after_char = stripped[char_pos + len(other):char_pos + len(other) + 15]

            for verb in POV_MIND_VERBS:
                if verb in after_char or verb in stripped[max(0, char_pos - 5):char_pos]:
                    violations.append({
                        'line': line_idx + 1,
                        'text': stripped[:80] + ('…' if len(stripped) > 80 else ''),
                        'reason': f'配角「{other}」心理越界：「{verb}」进入非 POV 角色内心',
                        'severity': 'high',
                    })
                    break
            else:
                continue
            break  # 一行只报一次

        # 3. 检查孤立心理状态词（无明确主语时的上帝视角判定）
        if not any(v for v in violations if v['line'] == line_idx + 1):
            # 如果行中有心理状态词且没有明确表示是 POV 角色观察到的
            for word in ['殊不知', '却不知', '并未察觉']:
                if word in stripped:
                    # 已经在上面的 god view check 中处理了
                    break
            else:
                # 检查"冷笑"等表情词 — 如果非 POV 角色是主语且用了冷笑
                for other in other_chars:
                    if other in stripped and '冷笑' in stripped:
                        char_pos = stripped.find(other)
                        # 检查"冷笑"是否在角色名附近
                        sneer_pos = stripped.find('冷笑')
                        if abs(sneer_pos - char_pos) < 30:
                            violations.append({
                                'line': line_idx + 1,
                                'text': stripped[:80] + ('…' if len(stripped) > 80 else ''),
                                'reason': f'配角「{other}」视角嫌疑：「冷笑」可能泄露内心活动（建议改为外部动作描写）',
                                'severity': 'medium',
                            })
                            break

    return {
        'violations': violations,
        'total': len(violations),
        'safe': len(violations) == 0,
    }


def audit_knowledge_leaks(content, chapter_id, all_characters, conn):
    """角色信息隔离墙审计：检测正文中是否存在认知越界。

    对每个出场角色，查询其到当前章之前的所有已知信息，
    若该角色在文中说出了不在已知集内的实体 → 认知越界。

    Args:
        content: 正文文本
        chapter_id: 当前章节 ID
        all_characters: 所有角色名列表
        conn: 数据库连接（调用方管理生命周期）

    Returns:
        dict: {leaks: [{character, mentioned, mentioned_type, line, text, reason, severity}]}
    """
    if not content or not chapter_id or len(all_characters) < 2:
        return {'leaks': [], 'safe': True}

    import re
    lines = content.split('\n')
    leaks = []
    char_set = set(all_characters)

    # 获取当前章节的 sort_order
    cur_ch = conn.execute(
        'SELECT sort_order FROM chapters WHERE id = ?', (chapter_id,)
    ).fetchone()
    if not cur_ch:
        return {'leaks': [], 'safe': True}
    cur_order = cur_ch['sort_order']

    # 找出文中实际出场的所有角色
    appearing_chars = set()
    for line in lines:
        for name in char_set:
            if name in line:
                appearing_chars.add(name)

    # 对每个出场角色进行认知审计
    for char_name in appearing_chars:
        # 查找该角色 ID
        char_row = conn.execute(
            'SELECT id FROM characters WHERE name = ?', (char_name,)
        ).fetchone()
        if not char_row:
            continue
        char_id = char_row['id']

        # 累积该角色在当前章之前所有章节的已知信息
        prev_rows = conn.execute('''
            SELECT ck.known_names, ck.known_items
            FROM character_knowledge ck
            JOIN chapters ch ON ck.chapter_id = ch.id
            WHERE ck.character_id = ? AND ch.sort_order < ?
        ''', (char_id, cur_order)).fetchall()

        known_names = set()
        known_items = set()
        for row in prev_rows:
            try:
                for n in json.loads(row['known_names'] or '[]'):
                    known_names.add(n)
            except (json.JSONDecodeError, TypeError):
                pass
            try:
                for it in json.loads(row['known_items'] or '[]'):
                    known_items.add(it)
            except (json.JSONDecodeError, TypeError):
                pass

        # 如果没有知识记录，跳过（不作审计）
        if not known_names and not known_items:
            continue

        # 扫描文中该角色的对话行，检测认知越界
        dialogue_pattern = re.compile(
            re.escape(char_name) + r'.{0,6}[""](.+?)[""]'
        )

        for line_idx, line in enumerate(lines):
            match = dialogue_pattern.search(line)
            if not match:
                continue
            dialogue = match.group(1)

            # 检查对话中是否提到了未知角色名
            other_chars = [c for c in all_characters if c != char_name]
            for other in other_chars:
                if other in dialogue and other not in known_names:
                    leaks.append({
                        'character': char_name,
                        'mentioned': other,
                        'mentioned_type': 'character',
                        'line': line_idx + 1,
                        'text': line.strip()[:80] + ('…' if len(line.strip()) > 80 else ''),
                        'reason': f'角色「{char_name}」认知越界：此时不该知晓角色「{other}」',
                        'severity': 'high',
                    })

            # 检查对话中是否提到了未知物品
            for item in known_items:
                if item in dialogue:
                    break  # 已知物品不报
            else:
                # 提取专有名词（2-4 字汉字组合）与 known_items 比对
                proper_nouns = set(re.findall(r'[一-鿿]{2,4}', dialogue))
                for noun in proper_nouns:
                    if noun in known_items:
                        continue
                    # 只报有意义的专有名词（过滤掉常见虚词）
                    if noun in char_set or len(noun) < 2:
                        continue
                    # 不在 known_items 中 → 可能违规
                    leaks.append({
                        'character': char_name,
                        'mentioned': noun,
                        'mentioned_type': 'item',
                        'line': line_idx + 1,
                        'text': line.strip()[:80] + ('…' if len(line.strip()) > 80 else ''),
                        'reason': f'角色「{char_name}」认知越界：此时不该知晓物品「{noun}」',
                        'severity': 'medium',
                    })

    return {
        'leaks': leaks,
        'safe': len(leaks) == 0,
    }


@app.route('/api/projects/<project_id>/pov/audit', methods=['POST'])
def api_pov_audit(project_id):
    """硬核视角审计：检测正文中的 POV 越界"""
    # Token 精算：QPS 频率限制
    if not check_rate_limit(get_user_id(), 'ai/dashboard'):
        return jsonify({'error': '请求过于频繁，请稍后再试'}), 429

    data = request.json
    content = data.get('content', '')
    pov_char = data.get('pov_character', '')
    chapter_id = data.get('chapter_id', '')

    # 获取所有角色名
    repo = get_repo()
    conn = None
    if repo:
        chars_data = repo.list_characters(project_id)
        all_chars = [c.get('name', '') for c in chars_data]
    else:
        conn = get_db(project_id)
        chars = conn.execute('SELECT name FROM characters ORDER BY created_at').fetchall()
        all_chars = [c['name'] for c in chars]

    # 如果没有指定 POV 角色，默认取第一个
    if not pov_char and all_chars:
        pov_char = all_chars[0]

    result = audit_pov_text(content, pov_char, all_chars)
    result['pov_character'] = pov_char
    result['all_characters'] = all_chars

    # 知识泄露审计（角色信息隔离墙）
    # Supabase 模式：跳过知识泄露审计（需要复杂 SQL 查询）
    if repo:
        result['knowledge_leaks'] = []
        result['knowledge_safe'] = True
    elif chapter_id and conn:
        knowledge_result = audit_knowledge_leaks(content, chapter_id, all_chars, conn)
        result['knowledge_leaks'] = knowledge_result.get('leaks', [])
        result['knowledge_safe'] = knowledge_result.get('safe', True)
    else:
        result['knowledge_leaks'] = []
        result['knowledge_safe'] = True

    if conn:
        conn.close()
    return jsonify(result)




@app.route('/api/projects/<project_id>/ai/analyze-rhythm', methods=['POST'])
def api_analyze_rhythm(project_id):
    """实时分析正文的冲突/悬念/爽点/信息密度"""
    # Token 精算：QPS 频率限制
    if not check_rate_limit(get_user_id(), 'ai/dashboard'):
        return jsonify({'error': '请求过于频繁，请稍后再试'}), 429

    data = request.json
    content = data.get('content', '')
    result = analyze_rhythm(content)
    return jsonify(result)


@app.route('/api/projects/<project_id>/ai/realism-radar', methods=['POST'])
def api_realism_radar(project_id):
    """生存逻辑与写实度雷达：检测文本是否滑向无脑爽文"""
    if not check_rate_limit(get_user_id(), 'ai/dashboard'):
        return jsonify({'error': '请求过于频繁，请稍后再试'}), 429

    data = request.json
    content = data.get('content', '')
    result = realism_radar(content)
    return jsonify(result)


@app.route('/api/projects/<project_id>/ai/analyze-cliffhanger', methods=['POST'])
def api_analyze_cliffhanger(project_id):
    """断章钩子密度分析器：分析章节尾部文本的悬念切断质量"""
    if not check_rate_limit(get_user_id(), 'ai/dashboard'):
        return jsonify({'error': '请求过于频繁，请稍后再试'}), 429

    data = request.json
    tail_text = data.get('tail_text', '')
    result = analyze_cliffhanger_tail(tail_text)
    return jsonify(result)


@app.route('/api/projects/<project_id>/ai/predict-retention', methods=['POST'])
def api_predict_retention(project_id):
    """AI 验证期留存预测 — 唯一真正消耗云端 Token 的仪表盘功能"""
    # SaaS 会员阶梯门控
    if not check_feature('predict_retention'):
        return jsonify({'error': '此功能需要升级会员', 'code': 'tier_required'}), 403

    data = request.json
    content = data.get('content', '')

    if not content or len(content.strip()) < 50:
        return jsonify({
            'ch1_read_rate': None, 'ch3_retention': None,
            'd7_retention': None, 'd8_read_rate': None,
            'tip': '正文内容不足，无法预测', 'platform': None
        })

    sample = content[:3000]

    prompt = f"""你是一位资深的网文平台数据分析师。请根据以下小说章节内容，预测该作品在番茄/起点等主流平台的留存数据。

分析维度：
1. ch1_read_rate：第1章读完率（0-100%），评估开篇吸引力
2. ch3_retention：第3章留存率（0-100%），评估持续阅读意愿
3. d7_retention：7日留存率（0-100%），评估追读粘性
4. d8_read_rate：第8章读完率（0-100%），评估中期内容质量
5. platform：最适合发布的平台（"番茄小说"、"起点中文网"、"晋江文学城"、"QQ阅读" 之一）
6. tip：一句优化建议（不超过30字）

请严格按以下 JSON 格式返回，不要输出其他内容：
{{"ch1_read_rate": 数字, "ch3_retention": 数字, "d7_retention": 数字, "d8_read_rate": 数字, "platform": "平台名", "tip": "建议"}}

章节内容：
{sample}"""

    messages = [
        {'role': 'system', 'content': '你是一个专业的网文数据分析师，擅长预测阅读留存数据。请严格按照要求的 JSON 格式回复。'},
        {'role': 'user', 'content': prompt}
    ]

    result = call_ai(messages, temperature=0.3, max_tokens=500, endpoint='ai/predict-retention')

    if isinstance(result, tuple):
        return result

    try:
        import re
        json_match = re.search(r'\{[^}]+\}', str(result))
        if json_match:
            parsed = json.loads(json_match.group())
            billing = _get_billing()
            resp = {
                'ch1_read_rate': parsed.get('ch1_read_rate'),
                'ch3_retention': parsed.get('ch3_retention'),
                'd7_retention': parsed.get('d7_retention'),
                'd8_read_rate': parsed.get('d8_read_rate'),
                'platform': parsed.get('platform'),
                'tip': parsed.get('tip', '')
            }
            if billing:
                resp['billing'] = billing
            # 递增月度使用计数（Pro 层级配额）
            _increment_feature_usage(get_user_id(), 'predict_retention')
            return jsonify(resp)
        else:
            return jsonify({'error': 'AI 返回格式异常，请稍后重试'})
    except Exception as e:
        return jsonify({'error': f'预测失败: {str(e)}'})




@app.route('/api/projects/<project_id>/ai/deai-check', methods=['POST'])
def api_deai_check(project_id):
    """扫描正文中的 AI 味套话，返回分类命中结果"""
    # Token 精算：QPS 频率限制
    if not check_rate_limit(get_user_id(), 'ai/dashboard'):
        return jsonify({'error': '请求过于频繁，请稍后再试'}), 429

    data = request.json
    content = data.get('content', '')
    result = scan_cliches(content)
    return jsonify(result)


# 白描硬化前缀 — 注入所有 AI 写作 Prompt 的底层死命令
BAIMIAO_HARDENING = '''
【底层死命令 — 白描写作铁律】
你正在执行严格的白描写作模式。以下规则凌驾于所有其他指导之上：

1. 严格禁止：任何抒情语句、感叹句（"啊""多么""何等"）、全知视角的总结性发言（"这就是命运的转折""从那一刻起一切都变了"）
2. 严格禁止：对角色内心感受的直接描述（"他感到悲伤""她心中充满希望"）。内心只能通过外部动作暗示
3. 必须采用：冷峻、克制的白描手法。只描写客观可见的动作、环境细节、物品的具体参数
4. 必须描写：人物之间的认知信息差 — A 知道的 B 不知道，B 看到的 A 没看到。不要替任何角色做全知解释
5. 对话保持：每个角色只说自己在那个时刻应该说的话，不为读者解释剧情
6. 拒绝：任何 AI 高频词（"仿佛""然而""此外""史诗""那一刻"等），拒绝任何网文模板句式
7. 用：短句。多断句。节奏不匀。像人在喘气。
'''

SURVIVAL_REALISM_HARDENING = '''
【生存常识审计红线 — 凌驾于所有创作指令之上】

1. 物理与生理铁律：
   - 极度饥饿/脱水/失血/重伤状态下，人物动作必须严重受限
   - 受伤后战斗力呈指数衰减，绝不允许"身受重伤仍秒杀全场"
   - 普通人连续战斗不超过3-5分钟即力竭；修真者能量消耗需有明确代价
   - 环境因素（气温/地形/昼夜）对人物状态有直接、可感知的影响

2. 心理学与阶层认知铁律：
   - 配角绝不无理由倒头便拜/纳头便拜 — 每个人的臣服、背叛、结盟必须有基于自身利益的动机
   - 反派的每一个决策必须符合：a) 自身利益最大化 b) 阶层认知局限 c) 可得信息的限制
   - 禁止"为了作死而作死"的降智反派 — 反派的失败应来自信息差/资源差距/主角智取，而非莫名其妙犯蠢
   - 小人物（路人/店小二/小贩）的行为逻辑符合其社会阶层和生存压力

3. 信息差与认知鸿沟铁律：
   - 每个角色只知道自己亲身经历/被告知/可推断的信息
   - 禁止角色"未卜先知" — 即使读者知道，角色不能知道
   - 因情报不对等产生的误判和错误决策是正常且必要的叙事手段
   - 战斗/谈判场景中，双方基于不完整信息做决策，不依赖上帝视角

4. 环境残酷度铁律：
   - 乱世/末日/荒野场景中，生存压力（食物/水源/安全/疾病）持续存在
   - 舒适感和安全感必须被剧情"挣得"，不能自然而然拥有
   - 资源消耗（金钱/药品/弹药/法力）有明确代价和限制
   - 禁止"凭空出现救兵/资源/机缘" — 每个助力都应有前文铺垫
'''

def _build_rag_context(project_id, current_content='', max_context_chars=3000):
    """RAG 设定记忆引擎：结合关键词检索 + 结构化上下文，为 AI 调用提供精准前文信息。

    策略：
    1. 从当前写作内容中提取关键词
    2. 用关键词检索相关摘要/角色/情节
    3. 叠加最新的结构化上下文（角色状态、近期事件）
    4. 截断到 max_context_chars 以内，避免撑爆 prompt

    Args:
        project_id: 项目 ID
        current_content: 当前正在处理的内容（用于提取关键词）
        max_context_chars: 上下文最大字符数（防止过长）
    Returns:
        str: 格式化的 RAG 上下文
    """
    if not check_feature('premium'):
        return build_context_package(project_id)

    conn = get_db(project_id)

    # 1. 提取关键词
    keywords = _extract_keywords(current_content[:3000], max_keywords=12) if current_content else []

    # 2. 关键词检索上下文
    rag_text = _search_rag_context(project_id, keywords, conn)

    # 3. 叠加世界观上下文
    world_ctx = _build_world_context(project_id)
    if world_ctx:
        rag_text = world_ctx + '\n\n' + rag_text

    # 4. 叠加基础结构化上下文
    base_context = build_context_package(project_id)

    # 合并，RAG 结果优先
    combined = rag_text
    if base_context and base_context not in combined:
        combined += '\n' + base_context

    # 4. 截断
    if len(combined) > max_context_chars:
        # 按段落分割，保留最重要的部分
        paragraphs = combined.split('\n\n')
        result = ''
        for p in paragraphs:
            if len(result) + len(p) > max_context_chars:
                break
            result += p + '\n\n'
        combined = result.strip()

    conn.close()
    return combined

# 保留原有函数，作为基础上下文提供者
def build_context_package(project_id):
    """构建结构化上下文包，用于注入 AI prompt"""
    # Supabase 模式：从数据仓库获取
    repo = get_repo()
    if repo:
        # 获取摘要（含章节信息）
        summaries = repo.list_chapter_summaries(project_id, limit=5)
        # 为每个摘要附加章节标题（从 chapters 表获取）
        chapters_map = {}
        if summaries:
            chapters_list = repo.list_chapters(project_id)
            chapters_map = {c['id']: c for c in chapters_list} if chapters_list else {}

        threads = repo.list_plot_threads(project_id)
        char_states = repo.list_character_states(project_id)
        events = repo.list_key_events(project_id, limit=10)
        all_chars = repo.list_characters(project_id)

        # 格式化输出
        parts = []

        if summaries:
            parts.append('【前文摘要】')
            for s in reversed(summaries):
                ch = chapters_map.get(s.get('chapter_id', ''), {})
                title = ch.get('title', f'章节{s.get("chapter_id", "")}')
                parts.append(f'《{title}》：{s.get("summary", "")}')

        if threads:
            parts.append('\n【活跃情节线程】')
            for t in threads:
                parts.append(f'- {t.get("title", "")}（{t.get("thread_type", "")}）：{t.get("description", "")}')

        if char_states:
            parts.append('\n【角色当前状态】')
            char_name_map = {c['id']: c.get('name', '') for c in all_chars} if all_chars else {}
            for cs in char_states:
                name = char_name_map.get(cs.get('character_id', ''), cs.get('character_id', '未知'))
                detail = f'{name}：{cs.get("status", "")}'
                if cs.get('location'):
                    detail += f'，位置：{cs.get("location")}'
                if cs.get('emotional_state'):
                    detail += f'，情绪：{cs.get("emotional_state")}'
                if cs.get('knowledge_gained'):
                    detail += f'，已知：{cs.get("knowledge_gained")}'
                parts.append(f'- {detail}')

        if events:
            parts.append('\n【近期关键事件】')
            for e in events:
                involved = e.get('involved_characters', '')
                if isinstance(involved, str) and involved.strip():
                    try:
                        involved = json.loads(involved)
                        involved = '、'.join(involved) if isinstance(involved, list) else involved
                    except json.JSONDecodeError:
                        pass
                parts.append(f'- {e.get("title", "")}（{e.get("event_type", "")}）：{e.get("description", "")}'
                              f'{f" [涉及：{involved}]" if involved else ""}')

        return '\n'.join(parts)

    conn = get_db(project_id)

    # 最近章节摘要（最新5章）
    summaries = conn.execute('''
        SELECT cs.summary, cs.key_events, cs.character_states, cs.plot_threads,
               c.title, c.sort_order, c.id as chapter_id
        FROM chapter_summaries cs
        JOIN chapters c ON cs.chapter_id = c.id
        ORDER BY c.sort_order DESC LIMIT 5
    ''').fetchall()

    # 活跃情节线程
    threads = conn.execute(
        "SELECT * FROM plot_threads WHERE status = 'active' ORDER BY created_at"
    ).fetchall()

    # 所有角色最新状态
    char_states = conn.execute('''
        SELECT DISTINCT cs.character_id, cs.status, cs.location, cs.emotional_state,
               cs.knowledge_gained, cs.relationships, ch.name
        FROM character_states cs
        JOIN characters ch ON cs.character_id = ch.id
        WHERE cs.id IN (
            SELECT id FROM character_states cs2
            WHERE cs2.character_id = cs.character_id
            ORDER BY cs2.snapshot_at DESC LIMIT 1
        )
    ''').fetchall()

    # 最近关键事件（最近10个）
    events = conn.execute('''
        SELECT ke.*, c.title as chapter_title
        FROM key_events ke
        JOIN chapters c ON ke.chapter_id = c.id
        ORDER BY ke.sort_order DESC LIMIT 10
    ''').fetchall()

    conn.close()

    # 格式化输出
    parts = []

    if summaries:
        parts.append('【前文摘要】')
        for s in reversed(summaries):
            parts.append(f'《{s["title"]}》：{s["summary"]}')

    if threads:
        parts.append('\n【活跃情节线程】')
        for t in threads:
            parts.append(f'- {t["title"]}（{t["thread_type"]}）：{t["description"]}')

    if char_states:
        parts.append('\n【角色当前状态】')
        for cs in char_states:
            detail = f'{cs["name"]}：{cs["status"]}'
            if cs['location']:
                detail += f'，位置：{cs["location"]}'
            if cs['emotional_state']:
                detail += f'，情绪：{cs["emotional_state"]}'
            if cs['knowledge_gained']:
                detail += f'，已知：{cs["knowledge_gained"]}'
            parts.append(f'- {detail}')

    if events:
        parts.append('\n【近期关键事件】')
        for ev in events:
            parts.append(f'- [{ev["chapter_title"]}] {ev["title"]}：{ev["description"]}')

    return '\n'.join(parts) if parts else ''

def _parse_json_field(val, default=None):
    """安全解析 JSON 字段"""
    if not val:
        return default if default is not None else []
    try:
        return json.loads(val)
    except (json.JSONDecodeError, TypeError):
        return default if default is not None else []

# ===== 世界观构建器 API =====
WORLDBUILDING_CATEGORIES = {
    'geography': '🏔️ 地理区域',
    'faction': '⚔️ 势力组织',
    'magic': '✨ 能力体系',
    'history': '📜 历史事件',
    'culture': '🎭 文化风俗',
    'economy': '💰 经济贸易',
    'religion': '🙏 宗教神话',
    'other': '📌 其他设定',
}

@app.route('/api/projects/<project_id>/worldbuilding', methods=['GET'])
def api_list_worldbuilding(project_id):
    """列出世界观条目，支持按分类过滤"""
    cat = request.args.get('category', '')

    repo = get_repo()
    if repo:
        items = repo.list_worldbuilding(project_id, cat if cat else None)
        return jsonify({'categories': WORLDBUILDING_CATEGORIES, 'items': items})

    conn = get_db(project_id)
    if cat and cat in WORLDBUILDING_CATEGORIES:
        rows = conn.execute(
            'SELECT * FROM worldbuilding WHERE category = ? ORDER BY sort_order, created_at', (cat,)
        ).fetchall()
    else:
        rows = conn.execute('SELECT * FROM worldbuilding ORDER BY category, sort_order, created_at').fetchall()
    conn.close()
    return jsonify({'categories': WORLDBUILDING_CATEGORIES, 'items': [dict(r) for r in rows]})

@app.route('/api/projects/<project_id>/worldbuilding', methods=['POST'])
def api_create_worldbuilding(project_id):
    """创建世界观条目"""
    data = request.json
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'error': '名称不能为空'}), 400
    category = data.get('category', 'other')
    if category not in WORLDBUILDING_CATEGORIES:
        category = 'other'

    repo = get_repo()
    if repo:
        wid = repo.create_worldbuilding_entry(project_id, {
            'category': category, 'name': name,
            'description': data.get('description', ''),
            'details': data.get('details', {}),
            'sort_order': 0,
        })
        if wid:
            return jsonify({'id': wid})
        return jsonify({'error': '创建失败'}), 500

    now = datetime.now().isoformat()
    wid = secrets.token_hex(8)
    conn = get_db(project_id)
    max_sort = conn.execute('SELECT COALESCE(MAX(sort_order), 0) FROM worldbuilding WHERE category = ?', (category,)).fetchone()[0]
    conn.execute('''
        INSERT INTO worldbuilding (id, category, name, description, details, sort_order, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    ''', (wid, category, name, data.get('description', ''), json.dumps(data.get('details', {}), ensure_ascii=False),
          max_sort + 1, now, now))
    conn.commit()
    conn.close()
    return jsonify({'id': wid})

@app.route('/api/projects/<project_id>/worldbuilding/<item_id>', methods=['PUT'])
def api_update_worldbuilding(project_id, item_id):
    """更新世界观条目"""
    data = request.json

    repo = get_repo()
    if repo:
        update_data = {}
        for key in ['name', 'description', 'details', 'category']:
            if key in data:
                update_data[key] = data[key]
        if update_data:
            _, err = repo.update_worldbuilding_entry(item_id, update_data)
            if err:
                return jsonify({'error': '更新失败'}), 500
        return jsonify({'success': True})

    conn = get_db(project_id)
    fields = []
    values = []
    for key in ['name', 'description', 'details', 'category']:
        if key in data:
            if key == 'details' and isinstance(data[key], dict):
                fields.append(f'{key} = ?')
                values.append(json.dumps(data[key], ensure_ascii=False))
            else:
                fields.append(f'{key} = ?')
                values.append(data[key])
    fields.append('updated_at = ?')
    values.append(datetime.now().isoformat())
    values.append(item_id)
    conn.execute(f'UPDATE worldbuilding SET {", ".join(fields)} WHERE id = ?', values)
    conn.commit()
    conn.close()
    return jsonify({'success': True})

@app.route('/api/projects/<project_id>/worldbuilding/<item_id>', methods=['DELETE'])
def api_delete_worldbuilding(project_id, item_id):
    """删除世界观条目"""
    repo = get_repo()
    if repo:
        _, err = repo.delete_worldbuilding_entry(item_id)
        if err:
            return jsonify({'error': '删除失败'}), 500
        return jsonify({'success': True})

    conn = get_db(project_id)
    conn.execute('DELETE FROM worldbuilding WHERE id = ?', (item_id,))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

def _build_world_context(project_id):
    """构建世界观上下文，用于注入 AI prompt"""
    # Supabase 模式：从数据仓库获取
    repo = get_repo()
    if repo:
        items = repo.list_worldbuilding(project_id)
        if not items:
            return ''
        parts = ['【世界观设定】请严格遵守以下世界观设定：']
        current_cat = None
        for item in items:
            cat = item.get('category', '')
            if cat != current_cat:
                current_cat = cat
                parts.append(f'\n{WORLDBUILDING_CATEGORIES.get(cat, cat)}：')
            desc = item.get('description') or ''
            parts.append(f'- {item.get("name", "")}：{desc[:200]}')
        return '\n'.join(parts)

    conn = get_db(project_id)
    items = conn.execute('SELECT * FROM worldbuilding ORDER BY category, sort_order').fetchall()
    conn.close()
    if not items:
        return ''
    parts = ['【世界观设定】请严格遵守以下世界观设定：']
    current_cat = None
    for item in items:
        cat = item['category']
        if cat != current_cat:
            current_cat = cat
            parts.append(f'\n{WORLDBUILDING_CATEGORIES.get(cat, cat)}：')
        desc = item['description'] or ''
        parts.append(f'- {item["name"]}：{desc[:200]}')
    return '\n'.join(parts)

@app.route('/api/projects/<project_id>/ai/check-setting-consistency', methods=['POST'])
def api_check_setting_consistency(project_id):
    """检查当前内容是否与世界观设定一致"""
    if not check_feature('ai_call'):
        return jsonify({'error': '余额不足，请先充值'}), 403

    data = request.json
    content = data.get('content', '')
    if not content.strip():
        return jsonify({'error': '请提供要检查的内容'}), 400

    model = data.get('model', 'deepseek')

    # 获取所有世界观设定
    world_ctx = _build_world_context(project_id)
    if not world_ctx:
        return jsonify({'content': '暂无世界观设定条目。请先在「世界观」页面创建地理、势力、能力体系等设定，再进行一致性检查。'})

    # 获取角色信息
    repo = get_repo()
    char_ctx = ''
    if repo:
        chars_data = repo.list_characters(project_id)
        if chars_data:
            char_ctx = '\n'.join(
                f'{c.get("name", "")}（{c.get("gender") or ""}，{c.get("age") or ""}岁，'
                f'{c.get("personality") or ""}，{c.get("background") or ""}）'
                for c in chars_data
            )
    else:
        conn = get_db(project_id)
        chars = conn.execute('SELECT name, gender, age, personality, background FROM characters').fetchall()
        conn.close()
        if chars:
            char_ctx = '\n'.join(
                f'{c["name"]}（{c["gender"] or ""}，{c["age"] or ""}岁，'
                f'{c["personality"] or ""}，{c["background"] or ""}）'
                for c in chars
            )

    sys_prompt = (
        '你是一位严格的小说设定审核官。请将以下正文内容与世界观设定逐条比对，找出所有矛盾或不一致之处。\n\n'
        + world_ctx
        + ('\n\n【角色设定】\n' + char_ctx if char_ctx else '')
        + '\n\n请以 JSON 格式输出检查结果：\n'
        '{\n'
        '  "score": <0-100 设定吻合度分数>,\n'
        '  "verdict": "<一句话总结>",\n'
        '  "issues": [\n'
        '    {"severity": "high/medium/low", "setting": "<冲突的设定条目>", "problem": "<正文中的矛盾描述>", "suggestion": "<修改建议>"}\n'
        '  ],\n'
        '  "summary": "<整体评价>"\n'
        '}\n'
        '只输出 JSON，不要 markdown 代码块包裹。如果没有发现问题，issues 为空数组，score 为 100。'
    )

    user_content = f'请检查以下正文内容是否与世界观设定存在矛盾：\n\n{content[:4000]}'

    messages = [
        {'role': 'system', 'content': sys_prompt},
        {'role': 'user', 'content': user_content}
    ]
    try:
        result = call_ai(messages, model=model, temperature=0.3, max_tokens=2000)
        # Clean markdown code blocks
        cleaned = result.strip()
        if cleaned.startswith('```'):
            cleaned = cleaned.split('\n', 1)[-1] if '\n' in cleaned else cleaned[3:]
        if cleaned.endswith('```'):
            cleaned = cleaned[:-3].strip()
        try:
            json.loads(cleaned)
        except json.JSONDecodeError:
            pass  # 不强制要求有效 JSON，前端会容错
        return jsonify({'content': cleaned})
    except Exception as e:
        return jsonify({'error': f'设定检查失败: {str(e)}'}), 500


@app.route('/api/projects/<project_id>/chapters/<chapter_id>/summarize', methods=['POST'])
def api_summarize_chapter(project_id, chapter_id):
    """AI 生成章节摘要、关键事件、角色状态"""
    if not check_feature('premium'):
        return jsonify({'error': '章节摘要为高级功能，请升级会员'}), 403

    repo = get_repo()

    # Supabase 模式
    if repo:
        chapter = repo.get_chapter(chapter_id)
        if not chapter:
            return jsonify({'error': '章节不存在'}), 404
        content = chapter.get('content') or ''
        if len(content.strip()) < 100:
            return jsonify({'error': '章节内容太少，无法生成摘要'}), 400

        chars = repo.list_characters(project_id)
        char_list = '\n'.join(f'{c.get("name", "")}（{c.get("personality") or ""}）' for c in chars) if chars else '暂无角色'

        threads = repo.list_plot_threads(project_id)
        thread_list = '\n'.join(
            f'- {t.get("title", "")}（{t.get("thread_type", "")}）：{t.get("description", "")}'
            for t in threads if t.get('status') == 'active'
        ) if threads else '暂无'

        sample = content[:5000]

        messages = [{
            'role': 'system',
            'content': '''你是小说结构分析师。请分析章节内容并输出 JSON：

{
  "summary": "200字以内的章节摘要",
  "key_events": [
    {"title": "事件标题", "description": "描述", "event_type": "revelation|conflict|turning_point|character_moment|setup|payoff", "involved_characters": ["角色名"]}
  ],
  "character_states": {
    "角色名": {"status": "alive|injured|missing|dead", "location": "地点", "emotional_state": "情绪", "knowledge_gained": "新获得的信息"}
  },
  "plot_threads_affected": ["受影响的情节线程标题"]
}

直接输出 JSON，不含其他内容。'''
        }, {
            'role': 'user',
            'content': f'角色列表：\n{char_list}\n\n已有情节线程：\n{thread_list}\n\n章节标题：{chapter.get("title", "")}\n\n章节内容：\n{sample}'
        }]

        result = call_ai(messages, temperature=0.3, max_tokens=2000)
        parsed = safe_json_extract(result)
        if parsed is None:
            return jsonify({'content': result, 'parsed': False})

        now = datetime.now().isoformat()
        summary = parsed.get('summary', result[:300])
        key_events = parsed.get('key_events', [])
        char_states = parsed.get('character_states', {})

        # 存储摘要
        repo.upsert_chapter_summary(project_id, chapter_id, {
            'summary': summary,
            'key_events': json.dumps(key_events, ensure_ascii=False),
            'character_states': json.dumps(char_states, ensure_ascii=False),
            'plot_threads': json.dumps(parsed.get('plot_threads_affected', []), ensure_ascii=False),
            'word_count': chapter.get('word_count', 0),
        })

        # 存储关键事件
        for ev in key_events:
            repo.create_key_event(project_id, {
                'chapter_id': chapter_id,
                'title': ev.get('title', ''),
                'description': ev.get('description', ''),
                'event_type': ev.get('event_type', 'event'),
                'involved_characters': json.dumps(ev.get('involved_characters', []), ensure_ascii=False),
            })

        # 存储角色状态
        for char_name, state in char_states.items():
            # 查找角色 ID
            char_id = None
            for c in chars:
                if c.get('name') == char_name:
                    char_id = c.get('id')
                    break
            if char_id:
                repo.upsert_character_state(project_id, {
                    'character_id': char_id,
                    'chapter_id': chapter_id,
                    'location': state.get('location', ''),
                    'status': state.get('status', 'alive'),
                    'emotional_state': state.get('emotional_state', ''),
                    'knowledge_gained': state.get('knowledge_gained', ''),
                    'relationships': json.dumps(state.get('relationships', {}), ensure_ascii=False),
                })

        return jsonify({'success': True, 'summary': summary, 'event_count': len(key_events), 'parsed': True})

    conn = get_db(project_id)
    chapter = conn.execute('SELECT * FROM chapters WHERE id = ?', (chapter_id,)).fetchone()
    if not chapter:
        conn.close()
        return jsonify({'error': '章节不存在'}), 404

    content = chapter['content'] or ''
    if len(content.strip()) < 100:
        conn.close()
        return jsonify({'error': '章节内容太少，无法生成摘要'}), 400

    # 获取角色列表
    chars = conn.execute('SELECT id, name, personality FROM characters').fetchall()
    char_list = '\n'.join(f'{c["name"]}（{c["personality"] or ""}）' for c in chars) if chars else '暂无角色'

    # 获取已有情节线程
    threads = conn.execute("SELECT * FROM plot_threads WHERE status = 'active'").fetchall()
    thread_list = '\n'.join(f'- {t["title"]}（{t["thread_type"]}）：{t["description"]}' for t in threads) if threads else '暂无'

    conn.close()

    sample = content[:5000]

    messages = [{
        'role': 'system',
        'content': '''你是小说结构分析师。请分析章节内容并输出 JSON：

{
  "summary": "200字以内的章节摘要",
  "key_events": [
    {"title": "事件标题", "description": "描述", "event_type": "revelation|conflict|turning_point|character_moment|setup|payoff", "involved_characters": ["角色名"]}
  ],
  "character_states": {
    "角色名": {"status": "alive|injured|missing|dead", "location": "地点", "emotional_state": "情绪", "knowledge_gained": "新获得的信息"}
  },
  "plot_threads_affected": ["受影响的情节线程标题"]
}

直接输出 JSON，不含其他内容。'''
    }, {
        'role': 'user',
        'content': f'角色列表：\n{char_list}\n\n已有情节线程：\n{thread_list}\n\n章节标题：{chapter["title"]}\n\n章节内容：\n{sample}'
    }]

    result = call_ai(messages, temperature=0.3, max_tokens=2000)

    # 解析 AI 输出
    parsed = safe_json_extract(result)
    if parsed is None:
        return jsonify({'content': result, 'parsed': False})

    now = datetime.now().isoformat()

    # 存储摘要、关键事件、角色状态（使用安全上下文管理器）
    summary = parsed.get('summary', result[:300])
    key_events_json = json.dumps(parsed.get('key_events', []), ensure_ascii=False)
    char_states_json = json.dumps(parsed.get('character_states', {}), ensure_ascii=False)
    threads_json = json.dumps(parsed.get('plot_threads_affected', []), ensure_ascii=False)
    key_events = parsed.get('key_events', [])
    char_states = parsed.get('character_states', {})

    with get_db_safe(project_id) as conn:
        conn.execute('''
            INSERT OR REPLACE INTO chapter_summaries (chapter_id, summary, key_events, character_states, plot_threads, generated_at, word_count)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (chapter_id, summary, key_events_json, char_states_json, threads_json, now, chapter['word_count']))

        max_sort = conn.execute('SELECT COALESCE(MAX(sort_order), 0) as mx FROM key_events').fetchone()['mx']
        for i, ev in enumerate(key_events):
            ev_id = secrets.token_hex(8)
            conn.execute('''
                INSERT INTO key_events (id, chapter_id, title, description, event_type, involved_characters, sort_order, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (ev_id, chapter_id, ev.get('title', ''), ev.get('description', ''),
                  ev.get('event_type', 'event'), json.dumps(ev.get('involved_characters', []), ensure_ascii=False),
                  max_sort + i + 1, now))

        for char_name, state in char_states.items():
            char_row = conn.execute('SELECT id FROM characters WHERE name = ?', (char_name,)).fetchone()
            if char_row:
                cs_id = secrets.token_hex(8)
                conn.execute('''
                    INSERT INTO character_states (id, character_id, chapter_id, location, status, emotional_state, knowledge_gained, relationships, snapshot_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (cs_id, char_row['id'], chapter_id,
                      state.get('location', ''), state.get('status', 'alive'),
                      state.get('emotional_state', ''), state.get('knowledge_gained', ''),
                      json.dumps(state.get('relationships', {}), ensure_ascii=False), now))

        conn.commit()

    return jsonify({'success': True, 'summary': summary, 'event_count': len(key_events), 'parsed': True})

@app.route('/api/projects/<project_id>/context', methods=['GET'])
def api_get_context(project_id):
    """获取当前上下文包（RAG 增强版）"""
    if not check_feature('premium'):
        return jsonify({'error': '上下文管理为高级功能，请升级会员'}), 403
    # 从 query string 获取当前内容用于关键词提取
    current_text = request.args.get('current_text', '')
    context = _build_rag_context(project_id, current_text) if current_text else build_context_package(project_id)

    # Supabase 模式：上下文已由 build_* 函数构建，额外数据从 repo 获取
    repo = get_repo()
    if repo:
        summaries = repo.list_chapter_summaries(project_id, limit=50)
        # 为摘要附加章节标题
        chapters_list = repo.list_chapters(project_id)
        chapters_map = {c['id']: c for c in chapters_list} if chapters_list else {}
        for s in summaries:
            ch = chapters_map.get(s.get('chapter_id', ''), {})
            s['chapter_title'] = ch.get('title', '')
            s['sort_order'] = ch.get('sort_order', 0)

        threads = repo.list_plot_threads(project_id)
        events = repo.list_key_events(project_id, limit=50)
        # 为事件附加章节标题
        for e in events:
            ch = chapters_map.get(e.get('chapter_id', ''), {})
            e['chapter_title'] = ch.get('title', '')

        return jsonify({
            'context_text': context,
            'summaries': summaries,
            'plot_threads': threads,
            'key_events': events
        })

    conn = get_db(project_id)
    summaries = conn.execute('''
        SELECT cs.*, c.title as chapter_title, c.sort_order
        FROM chapter_summaries cs JOIN chapters c ON cs.chapter_id = c.id
        ORDER BY c.sort_order
    ''').fetchall()
    threads = conn.execute("SELECT * FROM plot_threads ORDER BY created_at").fetchall()
    events = conn.execute('''
        SELECT ke.*, c.title as chapter_title
        FROM key_events ke JOIN chapters c ON ke.chapter_id = c.id
        ORDER BY ke.sort_order
    ''').fetchall()
    conn.close()

    return jsonify({
        'context_text': context,
        'summaries': [dict(s) for s in summaries],
        'plot_threads': [dict(t) for t in threads],
        'key_events': [dict(e) for e in events]
    })

@app.route('/api/projects/<project_id>/plot-threads', methods=['GET'])
def api_list_plot_threads(project_id):
    if not check_feature('premium'):
        return jsonify({'error': '情节线程为高级功能，请升级会员'}), 403

    repo = get_repo()
    if repo:
        return jsonify(repo.list_plot_threads(project_id))

    conn = get_db(project_id)
    rows = conn.execute('SELECT * FROM plot_threads ORDER BY created_at').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/projects/<project_id>/plot-threads', methods=['POST'])
def api_create_plot_thread(project_id):
    if not check_feature('premium'):
        return jsonify({'error': '情节线程为高级功能，请升级会员'}), 403
    data = request.json

    repo = get_repo()
    if repo:
        tid = repo.create_plot_thread(project_id, data)
        if tid:
            return jsonify({'id': tid})
        return jsonify({'error': '创建失败'}), 500

    now = datetime.now().isoformat()
    tid = secrets.token_hex(8)
    conn = get_db(project_id)
    conn.execute('''
        INSERT INTO plot_threads (id, title, description, thread_type, status, start_chapter_id, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    ''', (tid, data.get('title', ''), data.get('description', ''),
          data.get('thread_type', 'subplot'), data.get('status', 'active'),
          data.get('start_chapter_id'), now, now))
    conn.commit()
    conn.close()
    return jsonify({'id': tid})

@app.route('/api/projects/<project_id>/plot-threads/<thread_id>', methods=['PUT'])
def api_update_plot_thread(project_id, thread_id):
    if not check_feature('premium'):
        return jsonify({'error': '情节线程为高级功能，请升级会员'}), 403
    data = request.json

    repo = get_repo()
    if repo:
        _, err = repo.update_plot_thread(thread_id, data)
        if err:
            return jsonify({'error': '更新失败'}), 500
        return jsonify({'success': True})

    conn = get_db(project_id)
    fields = []
    values = []
    for key in ['title', 'description', 'thread_type', 'status', 'end_chapter_id']:
        if key in data:
            fields.append(f'{key} = ?')
            values.append(data[key])
    fields.append('updated_at = ?')
    values.append(datetime.now().isoformat())
    values.append(thread_id)
    conn.execute(f'UPDATE plot_threads SET {", ".join(fields)} WHERE id = ?', values)
    conn.commit()
    conn.close()
    return jsonify({'success': True})

@app.route('/api/projects/<project_id>/key-events', methods=['GET'])
def api_list_key_events(project_id):
    if not check_feature('premium'):
        return jsonify({'error': '关键事件为高级功能，请升级会员'}), 403
    chapter_id = request.args.get('chapter_id', '')
    conn = get_db(project_id)
    if chapter_id:
        rows = conn.execute(
            'SELECT ke.*, c.title as chapter_title FROM key_events ke JOIN chapters c ON ke.chapter_id = c.id WHERE ke.chapter_id = ? ORDER BY ke.sort_order',
            (chapter_id,)
        ).fetchall()
    else:
        rows = conn.execute(
            'SELECT ke.*, c.title as chapter_title FROM key_events ke JOIN chapters c ON ke.chapter_id = c.id ORDER BY ke.sort_order'
        ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

# ===== AI 对话归类 =====

@app.route('/api/projects/<project_id>/ai/conversations', methods=['GET'])
def api_list_conversations(project_id):
    repo = get_repo()
    if repo:
        return jsonify(repo.list_conversations(project_id))

    category = request.args.get('category', '')
    conn = get_db(project_id)
    if category:
        rows = conn.execute('SELECT * FROM ai_conversations WHERE category = ? ORDER BY last_message_at DESC', (category,)).fetchall()
    else:
        rows = conn.execute('SELECT * FROM ai_conversations ORDER BY last_message_at DESC').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/projects/<project_id>/ai/conversations', methods=['POST'])
def api_create_conversation(project_id):
    data = request.json

    repo = get_repo()
    if repo:
        conv_id = repo.create_conversation(
            project_id,
            category=data.get('category', 'general'),
            topic=data.get('topic', ''),
            source_tab=data.get('source_tab', 'chat')
        )
        if conv_id:
            return jsonify({'id': conv_id})
        return jsonify({'error': '创建失败'}), 500

    now = datetime.now().isoformat()
    conv_id = secrets.token_hex(8)
    conn = get_db(project_id)
    conn.execute('''
        INSERT INTO ai_conversations (id, category, topic, source_tab, message_count, last_message_at, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (conv_id, data.get('category', 'general'), data.get('topic', ''),
          data.get('source_tab', 'chat'), 0, now, now))
    conn.commit()
    conn.close()
    return jsonify({'id': conv_id})

@app.route('/api/projects/<project_id>/ai/conversations/<conv_id>/messages', methods=['GET'])
def api_get_conversation_messages(project_id, conv_id):
    repo = get_repo()
    if repo:
        return jsonify(repo.list_messages(conv_id))

    conn = get_db(project_id)
    rows = conn.execute('SELECT * FROM ai_messages WHERE conversation_id = ? ORDER BY id', (conv_id,)).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/projects/<project_id>/ai/conversations/<conv_id>/messages', methods=['POST'])
def api_add_conversation_message(project_id, conv_id):
    data = request.json
    role = data.get('role', 'user')
    content = data.get('content', '')

    repo = get_repo()
    if repo:
        ok = repo.add_message(project_id, conv_id, role, content)
        if ok:
            return jsonify({'success': True})
        return jsonify({'error': '添加消息失败'}), 500

    now = datetime.now().isoformat()
    conn = get_db(project_id)
    conn.execute('''
        INSERT INTO ai_messages (conversation_id, role, content, timestamp)
        VALUES (?, ?, ?, ?)
    ''', (conv_id, role, content, now))

    msg_count = conn.execute('SELECT COUNT(*) as cnt FROM ai_messages WHERE conversation_id = ?', (conv_id,)).fetchone()['cnt']

    if role == 'user' and msg_count <= 4:
        recent = conn.execute(
            'SELECT content FROM ai_messages WHERE conversation_id = ? ORDER BY id DESC LIMIT 4', (conv_id,)
        ).fetchall()
        sample = ' | '.join(r['content'][:200] for r in reversed(recent))

        categories = {
            '角色': 'character', '人物': 'character', '性格': 'character',
            '情节': 'plot', '剧情': 'plot', '故事': 'plot', '冲突': 'plot',
            '世界观': 'worldbuilding', '设定': 'worldbuilding', '世界': 'worldbuilding',
            '文笔': 'writing_style', '风格': 'writing_style', '写作': 'writing_style',
            '修改': 'revision', '改': 'revision', '润色': 'revision', '优化': 'revision',
        }
        detected = 'general'
        for keyword, cat in categories.items():
            if keyword in sample:
                detected = cat
                break

        if msg_count % 6 == 0:
            try:
                classify_prompt = f'将以下小说写作对话归类为一个类别：plot(情节), character(角色), worldbuilding(世界观), writing_style(文笔), revision(修改), general(通用)。只输出类别英文名。\n对话片段：{sample[:300]}'
                ai_cat = call_ai([{'role': 'user', 'content': classify_prompt}], temperature=0.1, max_tokens=20)
                ai_cat = ai_cat.strip().lower()
                valid_cats = ['plot', 'character', 'worldbuilding', 'writing_style', 'revision', 'general']
                if ai_cat in valid_cats:
                    detected = ai_cat
                topic_prompt = f'为这段写作对话生成一个简短主题标签（最多10个字）：{sample[:200]}'
                topic = call_ai([{'role': 'user', 'content': topic_prompt}], temperature=0.3, max_tokens=30)
                topic = topic.strip()[:15]
                conn.execute('UPDATE ai_conversations SET topic = ? WHERE id = ?', (topic, conv_id))
            except Exception:
                pass

        conn.execute('UPDATE ai_conversations SET category = ? WHERE id = ?', (detected, conv_id))

    conn.execute('''
        UPDATE ai_conversations SET message_count = ?, last_message_at = ? WHERE id = ?
    ''', (msg_count, now, conv_id))

    conn.commit()
    conn.close()
    return jsonify({'success': True, 'message_count': msg_count})

# ===== 设置系统 =====

def get_setting(key, default=None):
    """读取设置值"""
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    row = conn.execute('SELECT value FROM settings WHERE key = ?', (key,)).fetchone()
    conn.close()
    return row['value'] if row else default

def save_setting(key, value):
    """保存设置值"""
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.execute('INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)', (key, value))
    conn.commit()
    conn.close()

def mask_key(key):
    """掩码 API Key，只显示首尾"""
    if not key or len(key) < 12:
        return key or ''
    return key[:6] + '...' + key[-4:]

@app.route('/api/settings', methods=['GET'])
def api_get_settings():
    """获取用户可见设置：套餐信息 + 可用模型 + 默认偏好（API Key 仅管理员可见）"""
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    rows = conn.execute('SELECT key, value FROM settings').fetchall()
    conn.close()

    user_id = get_user_id()
    tier = _get_user_tier(user_id)
    tier_info = dict(MODEL_TIERS[tier])
    available_models = tier_info['models']
    settings = {
        'tier': tier,
        'tier_name': tier_info['name'],
        'tier_desc': tier_info['desc'],
        'default_model': get_setting('default_model', 'deepseek'),
        'available_models': [dict(MODEL_META[m], id=m) for m in available_models],
        'is_admin': _is_admin()
    }

    # API Key 统一从环境变量读取，不再通过 settings 暴露
    return jsonify(settings)

@app.route('/api/settings', methods=['POST'])
def api_save_settings():
    """保存设置：普通用户只能改默认模型；管理员可改 API Key"""
    data = request.json
    if not data:
        return jsonify({'error': '请提供要保存的设置'}), 400

    # 普通用户只能修改 default_model
    if 'default_model' in data:
        model = data['default_model']
        user_id = get_user_id()
        available = _get_available_models(user_id)
        if model not in available:
            return jsonify({'error': f'当前套餐不支持模型 {model}，请升级套餐'}), 403
        save_setting('default_model', model)

    # API Key 统一从环境变量配置，不再通过 API 修改
    return jsonify({'success': True})

# ===== AI 功能 API =====
import urllib.request
import json as json_mod
import ssl

# 使用默认 SSL 证书验证
def _get_ssl_context():
    """获取 SSL 上下文，优先使用 certifi 证书"""
    try:
        import certifi
        return ssl.create_default_context(cafile=certifi.where())
    except ImportError:
        return ssl.create_default_context()
ANTHROPIC_API_URL = 'https://api.anthropic.com/v1/messages'

def _get_api_key(model):
    """从环境变量读取 API Key（绝不从数据库读取，确保安全性）"""
    key_map = {
        'deepseek': 'DEEPSEEK_API_KEY',
        'claude': 'CLAUDE_API_KEY',
        'gemini': 'GEMINI_API_KEY',
    }
    env_var = key_map.get(model)
    if not env_var:
        return ''
    key = os.environ.get(env_var, '')
    if not key:
        raise RuntimeError(
            f'{env_var} 环境变量未设置。请在 .env 文件中配置 {env_var}=your-key'
        )
    return key

# 当前请求的 AI 调用计费信息（请求级共享状态）
_billing_state = threading.local()

def _get_billing():
    if not hasattr(_billing_state, 'data'):
        _billing_state.data = None
    return _billing_state.data

def _set_billing(data):
    _billing_state.data = data

def _deduct_balance(user_id, model, est_tokens=2000):
    """检查并预估扣费。返回 (allowed, error_msg)"""
    price_per_1k = _get_model_price(model)
    est_cost = est_tokens / 1000.0 * price_per_1k
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    row = conn.execute('SELECT * FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()
    if not row:
        conn.close()
        _ensure_balance_record(user_id)
        conn = sqlite3.connect(str(LICENSE_DB))
        conn.row_factory = sqlite3.Row
        row = conn.execute('SELECT * FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()
    # 重置每日免费额度
    today = datetime.now().strftime('%Y-%m-%d')
    free_date = row['free_daily_date'] or ''
    if free_date != today:
        conn.execute('UPDATE user_balances SET free_daily_used = 0.0, free_daily_date = ? WHERE user_id = ?', (today, user_id))
        conn.commit()
        free_remaining = row['free_daily_credits'] or 0.02
    else:
        free_remaining = max(0, (row['free_daily_credits'] or 0.02) - (row['free_daily_used'] or 0.0))
    balance = row['balance'] or 0.0
    total_available = balance + free_remaining
    if total_available < est_cost:
        conn.close()
        return False, f'__INSUFFICIENT_BALANCE__余额不足，当前余额 ¥{balance:.2f}（含免费额度 ¥{free_remaining:.2f}），本次预估 ¥{est_cost:.4f}。请充值后使用。'
    conn.close()
    return True, None

def _record_transaction(user_id, model, input_tokens, output_tokens, endpoint='ai'):
    """记录 AI 调用交易并实际扣费"""
    price_per_1k = _get_model_price(model)
    total_tokens = input_tokens + output_tokens
    cost = total_tokens / 1000.0 * price_per_1k
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    row = conn.execute('SELECT * FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()
    # 重置每日免费额度
    today = datetime.now().strftime('%Y-%m-%d')
    free_date = row['free_daily_date'] or ''
    if free_date != today:
        free_remaining = row['free_daily_credits'] or 0.02
    else:
        free_remaining = max(0, (row['free_daily_credits'] or 0.02) - (row['free_daily_used'] or 0.0))
    balance_before = (row['balance'] or 0.0) + free_remaining
    # 先扣免费额度
    free_used = min(cost, free_remaining)
    balance_cost = cost - free_used
    new_balance = (row['balance'] or 0.0) - balance_cost
    new_free_used = (row['free_daily_used'] or 0.0) + free_used
    now = datetime.now().isoformat()
    conn.execute('''
        UPDATE user_balances SET balance = ?, free_daily_used = ?, free_daily_date = ?,
        total_cost = total_cost + ?, updated_at = ? WHERE user_id = ?
    ''', (max(0, new_balance), new_free_used, today, cost, now, user_id))
    conn.commit()
    # 记录交易
    row2 = conn.execute('SELECT balance FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()
    balance_after = (row2['balance'] or 0.0) + max(0, free_remaining - new_free_used)
    txn_id = secrets.token_hex(10)
    conn.execute('''
        INSERT INTO ai_transactions (id, user_id, model, input_tokens, output_tokens, total_tokens, cost, balance_before, balance_after, free_credit_used, endpoint, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (txn_id, user_id, model, input_tokens, output_tokens, total_tokens, round(cost, 6),
          round(balance_before, 4), round(balance_after, 4), round(free_used, 6), endpoint, now))
    conn.commit()
    conn.close()
    # 存到共享状态供 ai_result 读取
    _set_billing({
        'model': model,
        'model_price_per_1k': price_per_1k,
        'input_tokens': input_tokens,
        'output_tokens': output_tokens,
        'total_tokens': total_tokens,
        'cost': round(cost, 6),
        'free_credit_used': round(free_used, 6),
        'balance_remaining': round(balance_after, 4),
        'balance': round(row2['balance'] or 0.0, 4),
        'free_remaining': round(max(0, free_remaining - free_used), 4)
    })

def _consume_ai_call():
    """[已废弃] 旧版配额检查，改为调用 _deduct_balance"""
    user_id = get_user_id()
    model = get_setting('default_model', 'deepseek')
    allowed, err = _deduct_balance(user_id, model, 2000)
    if not allowed:
        return err
    return None

def call_deepseek(messages, temperature=0.7, max_tokens=2000):
    """调用 DeepSeek API，返回 (text, usage_dict)"""
    api_key = _get_api_key('deepseek')
    if not api_key:
        return '错误：请先在设置中配置 DeepSeek API Key', None
    payload = {
        'model': DEEPSEEK_MODEL,
        'messages': messages,
        'temperature': temperature,
        'max_tokens': max_tokens
    }
    data = json_mod.dumps(payload).encode('utf-8')
    req = urllib.request.Request(DEEPSEEK_API_URL, data=data, headers={
        'Content-Type': 'application/json',
        'Authorization': f'Bearer {api_key}'
    })
    try:
        with urllib.request.urlopen(req, context=_get_ssl_context(), timeout=60) as resp:
            result = json_mod.loads(resp.read().decode('utf-8'))
            usage = result.get('usage', {})
            return result['choices'][0]['message']['content'], {
                'input': usage.get('prompt_tokens', 0),
                'output': usage.get('completion_tokens', 0)
            }
    except Exception as e:
        return f'DeepSeek调用失败: {str(e)}', None

def call_claude(messages, temperature=0.7, max_tokens=2000):
    """调用 Claude API，返回 (text, usage_dict)"""
    api_key = _get_api_key('claude')
    if not api_key:
        return '错误：请先在设置中配置 Claude API Key', None

    system_content = ''
    claude_messages = []
    for m in messages:
        if m['role'] == 'system':
            system_content = m['content']
        else:
            claude_messages.append({'role': m['role'], 'content': m['content']})

    payload = {
        'model': CLAUDE_MODEL,
        'max_tokens': max_tokens,
        'temperature': temperature,
        'messages': claude_messages
    }
    if system_content:
        payload['system'] = system_content

    data = json_mod.dumps(payload).encode('utf-8')
    req = urllib.request.Request(ANTHROPIC_API_URL, data=data, headers={
        'Content-Type': 'application/json',
        'x-api-key': api_key,
        'anthropic-version': '2023-06-01'
    })
    try:
        with urllib.request.urlopen(req, context=_get_ssl_context(), timeout=90) as resp:
            result = json_mod.loads(resp.read().decode('utf-8'))
            usage = result.get('usage', {})
            return result['content'][0]['text'], {
                'input': usage.get('input_tokens', 0),
                'output': usage.get('output_tokens', 0)
            }
    except Exception as e:
        return f'Claude调用失败: {str(e)}', None

def call_gemini(messages, temperature=0.7, max_tokens=2000):
    """调用 Gemini API，返回 (text, usage_dict)"""
    api_key = _get_api_key('gemini')
    if not api_key:
        return '错误：请先在设置中配置 Gemini API Key', None

    gemini_contents = []
    system_instruction = ''
    for m in messages:
        if m['role'] == 'system':
            system_instruction = m['content']
        elif m['role'] == 'user':
            gemini_contents.append({'role': 'user', 'parts': [{'text': m['content']}]})
        elif m['role'] == 'assistant':
            gemini_contents.append({'role': 'model', 'parts': [{'text': m['content']}]})

    payload = {
        'contents': gemini_contents,
        'generationConfig': {
            'temperature': temperature,
            'maxOutputTokens': max_tokens,
        }
    }
    if system_instruction:
        payload['systemInstruction'] = {'parts': [{'text': system_instruction}]}

    url = GEMINI_API_URL
    data = json_mod.dumps(payload).encode('utf-8')
    req = urllib.request.Request(url, data=data, headers={
        'Content-Type': 'application/json',
        'x-goog-api-key': api_key,
    })
    try:
        with urllib.request.urlopen(req, context=_get_ssl_context(), timeout=90) as resp:
            result = json_mod.loads(resp.read().decode('utf-8'))
            usage = result.get('usageMetadata', {})
            return result['candidates'][0]['content']['parts'][0]['text'], {
                'input': usage.get('promptTokenCount', 0),
                'output': usage.get('candidatesTokenCount', 0)
            }
    except Exception as e:
        return f'Gemini调用失败: {str(e)}', None

def ai_result(result, model=None, endpoint='ai'):
    """统一处理 AI 返回值：区分正常内容、限流、余额不足、API错误。附带计费信息。"""
    if isinstance(result, str) and result.startswith('__RATE_LIMIT__'):
        return jsonify({'error': result[14:]}), 429
    if isinstance(result, str) and result.startswith('__INSUFFICIENT_BALANCE__'):
        return jsonify({'error': result[22:], 'code': 'insufficient_balance'}), 402

    if isinstance(result, dict):
        content = result.get('content', '')
        quality_warnings = result.get('quality_warnings', [])
    else:
        content = result
        quality_warnings = []

    billing = _get_billing()
    resp = {
        'content': content,
        'model': model or get_setting('default_model', 'deepseek'),
    }
    if billing:
        resp['billing'] = billing
    if quality_warnings:
        resp['quality_warnings'] = quality_warnings
    return jsonify(resp)

def call_ai(messages, model=None, temperature=0.7, max_tokens=2000, endpoint='ai'):
    """统一 AI 调用入口，路由到对应 API 并自动计费（含重试+限流）"""
    if not model:
        model = get_setting('default_model', 'deepseek')
    user_id = get_user_id()

    # 检查 API Key 是否已配置
    if _APP_DEGRADED and model == 'deepseek':
        return 'AI 服务未配置：缺少 DEEPSEEK_API_KEY，请在管理后台配置。'
    if not GEMINI_API_KEY and model == 'gemini':
        return 'Gemini API 未配置，请在管理后台添加 GEMINI_API_KEY。'
    if not CLAUDE_API_KEY and model == 'claude':
        return 'Claude API 未配置，请在管理后台添加 CLAUDE_API_KEY。'

    # 限流检查
    if not check_rate_limit(user_id, endpoint):
        return jsonify({'error': '请求过于频繁，请稍后再试'}), 429

    # 预估扣费检查
    est_tokens = max_tokens
    for m in messages:
        est_tokens += len(m.get('content', '')) * 2
    allowed, err = _deduct_balance(user_id, model, est_tokens)
    if not allowed:
        return err

    # 调用模型 API，带重试（最多 3 次，指数退避）
    text, usage = None, None
    for attempt in range(3):
        try:
            if model == 'claude':
                text, usage = call_claude(messages, temperature, max_tokens)
            elif model == 'gemini':
                text, usage = call_gemini(messages, temperature, max_tokens)
            else:
                text, usage = call_deepseek(messages, temperature, max_tokens)

            if text and not (isinstance(text, str) and text.startswith(('错误', 'DeepSeek', 'Claude', 'Gemini'))):
                break
            if attempt < 2:
                time.sleep(2 ** attempt)  # 1s, 2s 退避
        except Exception as e:
            if attempt < 2:
                time.sleep(2 ** attempt)
            else:
                text = f'AI 调用失败（已重试3次）: {str(e)}'
                usage = None

    # 记录交易
    if usage and not (isinstance(text, str) and text.startswith(('错误', 'AI 调用失败'))):
        _record_transaction(user_id, model, usage['input'], usage['output'], endpoint)
    else:
        _set_billing(None)
    return text

@app.route('/api/ai/chat', methods=['POST'])
def ai_chat():
    """AI 对话"""
    data = request.json
    messages = data.get('messages', [])
    result = call_ai(messages, model=data.get('model'), temperature=0.8, max_tokens=3000)
    return ai_result(result, data.get('model'))

@app.route('/api/projects/<project_id>/ai/continue', methods=['POST'])
def ai_continue_writing(project_id):
    """AI 续写"""
    data = request.json
    content = data.get('content', '')
    prompt = data.get('prompt', '请继续创作，保持文风一致，自然衔接')
    chapter_id = data.get('chapter_id', '')

    # 使用精简上下文锚定大纲
    condensed = build_condensed_context(project_id, chapter_id)
    recent_content = content[-2000:] if len(content) > 2000 else content

    context_block = ''
    if condensed:
        context_block = f'【大纲锚定信息】请严格参照以下信息续写，确保不偏离主线：\n{condensed}\n\n'

    messages = [
        {'role': 'system', 'content': SURVIVAL_REALISM_HARDENING + '\n' + BAIMIAO_HARDENING + '\n' +
            '你是一位优秀的小说作家。保持原有文风和人物性格，自然流畅地继续创作。\n\n'
            '【人物一致性铁律】\n'
            '1. 每个角色的性格、说话方式、行为动机必须与设定严格一致，绝不越界\n'
            '2. 角色不能突然改变立场或能力级别，除非有明确的剧情转折铺垫\n'
            '3. 对话必须体现角色的独特性格和出身背景，不同角色语气应有明显差异\n'
            '4. 已死亡/已离场的角色不能再次出现\n\n'
            '【情节一致性铁律】\n'
            '1. 严格对照大纲锚定信息，续写内容必须推进大纲指定的主线\n'
            '2. 不要引入大纲中不存在的新势力/新角色/新设定\n'
            '3. 前文埋下的伏笔和线索需要延续，不能无故消失\n'
            '4. 每个情节转折都应该是新鲜的、不可预测的\n\n'
            '【写作质量铁律】\n'
            '1. 不使用网文常见套路（如"嘴角微扬""眼中闪过一丝"等模板句式）\n'
            '2. 避免AI式的过度工整排比，句子长短要有变化\n'
            '3. 用真实的感官细节（触觉/听觉/嗅觉/视觉）替代抽象形容\n'
            '4. 直接输出续写内容，不要加任何解释和标注'},
        {'role': 'user', 'content': f'{context_block}当前章节末尾内容：\n{recent_content}\n\n{prompt}'}
    ]
    result = call_ai(messages, model=data.get('model'), temperature=0.75, max_tokens=2000)

    # Post-write quality scan
    if isinstance(result, str) and not result.startswith(('错误', '__')):
        quality_warnings = _quick_quality_scan(result, project_id, current_chapter_id)
        if quality_warnings:
            result = {'content': result, 'quality_warnings': quality_warnings}

    return ai_result(result, data.get('model'))

@app.route('/api/projects/<project_id>/ai/polish', methods=['POST'])
def ai_polish(project_id):
    """AI 润色"""
    data = request.json
    content = data.get('content', '')
    instruction = data.get('instruction', '请润色这段文字，让文笔更优美流畅')
    chapter_id = data.get('chapter_id', '')

    text = content[-1500:] if len(content) > 1500 else content
    condensed = build_condensed_context(project_id, chapter_id)
    context_block = ''
    if condensed:
        context_block = f'【上下文参考】以下信息用于保持润色时角色语气和设定一致：\n{condensed}\n\n'

    messages = [
        {'role': 'system', 'content': SURVIVAL_REALISM_HARDENING + '\n' + BAIMIAO_HARDENING + '\n' + '你是一位专业的文学编辑。请润色文本让文笔更优美流畅，但你绝对不能：使用任何网文陈词滥调（如"白衣胜雪""倾国倾城"等）、套用常见模板句式、让人物对话变得千篇一律。保留原作者的独特风格和语气，只做润色而非重写。每个角色的语言必须保持个性差异。直接输出润色后的文本。'},
        {'role': 'user', 'content': f'{context_block}{instruction}：\n{text}'}
    ]
    result = call_ai(messages, model=data.get('model'), temperature=0.5, max_tokens=2000)
    return ai_result(result, data.get('model'))

@app.route('/api/projects/<project_id>/ai/deai', methods=['POST'])
def ai_deai(project_id):
    """去 AI 味"""
    data = request.json
    content = data.get('content', '')
    chapter_id = data.get('chapter_id', '')

    text = content[-1500:] if len(content) > 1500 else content
    condensed = build_condensed_context(project_id, chapter_id)
    rag_block = f'\n\n【设定参考】保持以下设定一致：\n{condensed}' if condensed else ''

    messages = [
        {'role': 'system', 'content': SURVIVAL_REALISM_HARDENING + '\n' + BAIMIAO_HARDENING + '\n' + '你是一位文学作家。将文本改写为更自然、更有"人味"的文字。核心原则：打破AI式过于完美的结构，引入自然的短句和断句；消除"然而""此外""总之"等AI高频词；让每个角色的对话语气独特；用具体感官细节（触觉/听觉/嗅觉）替代抽象形容；节奏要有变化，不要均匀铺陈。绝对避免"白衣少年""绝美面容"等网文套话。直接输出改写后的文本。'},
        {'role': 'user', 'content': f'请去AI味，让这段文字更像真人写的：\n{text}{rag_block}'}
    ]
    result = call_ai(messages, model=data.get('model'), temperature=0.85, max_tokens=2000)
    return ai_result(result, data.get('model'))

@app.route('/api/projects/<project_id>/ai/brainstorm', methods=['POST'])
def ai_brainstorm(project_id):
    """AI 头脑风暴 - 基于角色/大纲生成创作建议"""
    data = request.json
    topic = data.get('topic', '')
    context = data.get('context', '')

    structured_context = _build_rag_context(project_id, context if context else topic)
    context_block = ''
    if structured_context:
        context_block = f'\n\n【当前创作状态】请确保建议与以下已有内容协调：\n{structured_context}'

    messages = [
        {'role': 'system', 'content': '你是一位富有创意的小说策划师。提供具体、有创意、可执行的建议。所有建议必须避免网文常见套路（如废柴逆袭、退婚流、系统流等），鼓励原创设定和新鲜的人物关系。避免推荐已被用烂的桥段。每个建议应该有独特的核心卖点。'},
        {'role': 'user', 'content': f'创作主题：{topic}\n\n背景信息：{context}{context_block}\n\n请提供创作建议：'}
    ]
    result = call_ai(messages, model=data.get('model'), temperature=0.9, max_tokens=2000)
    return ai_result(result, data.get('model'))

# ===== AI 上下文一致性检查 =====
@app.route('/api/projects/<project_id>/ai/check-consistency', methods=['POST'])
def ai_check_consistency(project_id):
    """AI 自查上下文一致性：角色、时间线、情节连贯性"""
    data = request.json or {}
    conn = get_db(project_id)
    chapters = conn.execute('SELECT title, content, word_count FROM chapters ORDER BY sort_order').fetchall()
    chars = conn.execute('SELECT name, personality, background, goal FROM characters').fetchall()
    conn.close()
    
    # 构建角色摘要
    char_summary = ''
    if chars:
        char_summary = '角色列表：\n' + '\n'.join(
            f'- {c["name"]}：{c.get("personality","")[:50]}，{c.get("background","")[:50]}' 
            for c in chars
        )
    
    # 构建章节摘要（取每章前200字）
    chapter_summaries = []
    for ch in chapters:
        preview = (ch['content'] or '')[:200]
        chapter_summaries.append(f'《{ch["title"]}》({ch["word_count"]}字): {preview}...')
    
    messages = [
        {'role': 'system', 'content': '''你是一位严格的小说审稿编辑。请检查以下小说内容的一致性问题，包括：
1. 角色一致性：角色名称、性格、背景是否前后一致
2. 时间线：情节时间线是否有矛盾
3. 情节连贯性：剧情衔接是否自然
4. 设定矛盾：世界观设定是否前后冲突

请用中文简洁报告发现的问题，按严重程度排序（🔴严重 🟡注意 🟢小建议）。
如果一切正常，请明确说"未发现一致性问题"。
直接输出检查结果，不要加开场白。'''},
        {'role': 'user', 'content': f'{char_summary}\n\n章节内容摘要：\n' + '\n'.join(chapter_summaries) + '\n\n请检查一致性。'}
    ]
    result = call_ai(messages, model=data.get('model'), temperature=0.3, max_tokens=1500)
    return ai_result(result, data.get('model'))

# ===== AI味检测 =====
@app.route('/api/projects/<project_id>/ai/detect-ai-flavor', methods=['POST'])
def ai_detect_ai_flavor(project_id):
    """检测文本中的AI味：过度工整、重复模式、缺乏个性等"""
    data = request.json
    text = data.get('content', '')
    if not text or not text.strip():
        return jsonify({'error': '请提供需要检测的文本'}), 400
    
    # 只取前3000字检测
    sample = text[:3000]
    
    messages = [
        {'role': 'system', 'content': '''你是一位专业的内容审校，擅长识别AI生成文本的特征。请从以下维度分析文本：

1. **句式多样性**：是否存在过度工整、排比过多、句式重复？
2. **词汇丰富度**：是否频繁使用"然而""此外""总而言之"等AI高频过渡词？
3. **情感真实性**：情感描写是否空洞、套路化？
4. **细节具体性**：是否缺少具体感官细节（视觉、听觉、触觉）？
5. **对话自然度**：人物对话是否过于书面化？
6. **节奏变化**：是否有自然的短句、断句变化？

输出格式（JSON）：
{
  "score": 0-100 (AI味浓度分数，越高越像AI写的),
  "issues": ["问题1", "问题2"],
  "highlights": [{"text": "问题文本片段", "reason": "原因"}],
  "verdict": "综合评价一句话",
  "suggestions": ["改进建议1", "改进建议2"]
}

直接输出JSON，不要加任何解释。'''},
        {'role': 'user', 'content': f'请分析以下文本的AI写作痕迹：\n\n{sample}'}
    ]
    result = call_ai(messages, model=data.get('model'), temperature=0.3, max_tokens=1500)
    return ai_result(result, data.get('model'))

# ===== AI 灵感/题材生成 =====
@app.route('/api/projects/<project_id>/ai/idea-generator', methods=['POST'])
def ai_idea_generator(project_id):
    """根据题材生成创作灵感和想法"""
    data = request.json
    genre = data.get('genre', '')
    mode = data.get('mode', 'ideas')  # ideas | opening | twist | character
    existing_context = data.get('context', '')
    
    genre_prompts = {
        'ideas': f'请为{genre or "小说"}题材生成5个创新且有趣的故事灵感。每个灵感包含：一句话梗概、核心冲突、独特卖点。格式清晰，用编号列出。',
        'opening': f'请为{genre or "小说"}题材创作3个精彩的开头段落（每个150-300字）。开头要抓人眼球，建立悬念或氛围。',
        'twist': f'请为{genre or "小说"}题材设计3个令人意想不到的情节反转。每个反转包含：当前状态、反转事件、反转后的影响。',
        'character': f'请为{genre or "小说"}题材设计3个独特魅力的角色概念。每个包含：姓名、年龄、核心特质、背景故事梗概、在故事中的作用。',
        'world': f'请为{genre or "小说"}题材构建一个世界观的3个独特设定。每个包括：核心规则、独特元素、对故事的影响。',
        'conflict': f'请为{genre or "小说"}题材设计3个层层递进的核心冲突。每个包含：冲突来源、各方立场、升级方式、解决方向。',
    }
    
    prompt = genre_prompts.get(mode, genre_prompts['ideas'])
    if existing_context:
        prompt += f'\n\n当前故事背景：{existing_context[:500]}\n\n请确保灵感与当前故事背景协调。'
    
    messages = [
        {'role': 'system', 'content': '你是一位极具创意的畅销小说策划师。核心要求：所有灵感必须是原创的，绝不能复制任何已知作品的核心设定、角色关系或情节框架。避免网文模板（系统/重生/穿越/退婚/废柴逆袭等）。思考"如果我来写这个题材，怎么写出完全不同的东西？"。用独特的人物动机、新颖的世界观规则、出人意料的情节结构来创造真正有辨识度的作品。直接输出内容，格式清晰。'},
        {'role': 'user', 'content': prompt}
    ]
    result = call_ai(messages, model=data.get('model'), temperature=0.95, max_tokens=2500)
    return ai_result(result, data.get('model'))

# ===== AI 引导式写作（渐进提问） =====
@app.route('/api/projects/<project_id>/ai/guided-write', methods=['POST'])
def ai_guided_write(project_id):
    """AI 引导式写作：通过渐进提问帮助没有想法的用户"""
    data = request.json
    step = data.get('step', 0)  # 0=开始, 1=继续对话
    answer = data.get('answer', '')
    history = data.get('history', [])
    
    guidance_system = '''你是一位耐心温柔的小说创作导师。你的任务是帮助完全没有想法的用户逐步找到创作方向。

引导规则：
1. 从不预设用户想法，始终通过提问引导
2. 每次只问1-2个问题，不要一次问太多
3. 根据用户回答，逐步缩小范围
4. 适时给出2-3个具体选项让用户选择
5. 当信息足够时，主动总结并给出创作方案

原创性原则（重要）：
- 给出的故事构想必须是原创的，不能是已知作品的翻版
- 避免推荐网文套路（如系统流、退婚流、重生复仇等模板化设定）
- 鼓励独特的角色关系、新颖的世界观和出人意料的情节设计
- 如果某个构思与已有作品太像，主动建议调整方向

引导流程（渐进式）：
- 第0步：欢迎，询问是否有想写的题材/类型
- 根据回答深入：氛围偏好、目标读者、想表达的主题
- 当信息足够：总结并给出3个具体故事构想
- 最终：帮助制定大纲框架

保持自然对话语气，不要机械化。用emoji增加亲和力。'''

    if step == 0:
        # 欢迎引导
        messages = [
            {'role': 'system', 'content': guidance_system},
            {'role': 'user', 'content': '我是第一次使用，完全没有想法，请引导我开始创作。'}
        ]
    else:
        # 继续引导对话
        msgs = [{'role': 'system', 'content': guidance_system}]
        for h in history:
            msgs.append({'role': h.get('role', 'user'), 'content': h.get('content', '')})
        if answer:
            msgs.append({'role': 'user', 'content': answer})
        messages = msgs
    
    result = call_ai(messages, model=data.get('model'), temperature=0.8, max_tokens=2500)
    return ai_result(result, data.get('model'))

# ===== AI 续写增强版（方向控制） =====
@app.route('/api/projects/<project_id>/ai/continue-v2', methods=['POST'])
def ai_continue_writing_v2(project_id):
    """AI 续写增强版 - 支持方向控制"""
    data = request.json
    content = data.get('content', '')
    prompt = data.get('prompt', '请继续创作，保持文风一致')
    direction = data.get('direction', 'auto')  # auto|mainline|subplot|climax|foreshadow|character|suspense
    chapter_id = data.get('chapter_id', '')

    # 方向控制指令
    direction_guides = {
        'mainline': '严格沿主线情节推进，不要引入无关支线。保持故事聚焦在核心矛盾上。',
        'subplot': '发展一条有趣的支线情节，可以引入新角色或新线索，但最终要与主线呼应。',
        'climax': '推动情节走向高潮，增加张力和冲突。人物面临重大抉择或危机。',
        'foreshadow': '埋下伏笔和悬念，为后续情节做铺垫。注意细节的连贯性，让读者产生好奇。',
        'character': '聚焦角色内心成长，深化人物性格。通过具体事件展现角色的变化和抉择。',
        'suspense': '制造悬念和期待，让读者迫切想知道接下来会发生什么。节奏紧凑，信息逐步释放。',
    }

    dir_instruction = direction_guides.get(direction, '')

    # 使用精简上下文锚定大纲
    context = content[-3000:] if len(content) > 3000 else content
    condensed = build_condensed_context(project_id, chapter_id)

    context_block = ''
    if condensed:
        context_block = f'【大纲锚定信息】请严格参照以下信息续写，确保不偏离主线：\n{condensed}\n\n'

    direction_text = f'【写作方向要求】{dir_instruction}\n\n' if dir_instruction else ''

    messages = [
        {'role': 'system', 'content': SURVIVAL_REALISM_HARDENING + '\n' + BAIMIAO_HARDENING + '\n' + f'''你是一位优秀的小说作家，擅长续写故事。

【人物一致性铁律】
1. 每个角色的性格、说话方式、行为动机必须与设定严格一致，绝不越界
2. 角色不能突然改变立场或能力级别，除非有明确的剧情转折铺垫
3. 对话必须体现角色的独特性格和出身背景，不同角色语气应有明显差异
4. 已死亡/已离场的角色不能再次出现

【情节一致性铁律】
1. 严格对照大纲锚定信息，续写内容必须推进大纲指定的主线
2. 不要引入大纲中不存在的新势力/新角色/新设定
3. 前文埋下的伏笔和线索需要延续，不能无故消失
4. 每一个情节转折都应该让读者感到新鲜、不可预测

【写作质量铁律】
1. 避免使用AI常见的套话和过度工整的句式，让文字自然
2. 用具体而独特的感官细节代替抽象和通用的描写
3. 句子长短要有变化，节奏要有起伏
4. 直接输出续写内容，不要加任何解释和标注'''},
        {'role': 'user', 'content': f'{context_block}当前章节末尾内容：\n{context}\n\n{direction_text}{prompt}'}
    ]
    result = call_ai(messages, model=data.get('model'), temperature=0.75, max_tokens=2000)

    # Post-write quality scan
    if isinstance(result, str) and not result.startswith(('错误', '__')):
        quality_warnings = _quick_quality_scan(result, project_id, chapter_id)
        if quality_warnings:
            result = {'content': result, 'quality_warnings': quality_warnings}

    return ai_result(result, data.get('model'))

# ===== 导出 API =====
@app.route('/api/projects/<project_id>/export/docx', methods=['GET'])
def export_docx(project_id):
    """导出为 Word 文档"""
    if not check_feature('export', 'docx'):
        return jsonify({'error': '免费版不支持 Word 导出，请升级会员'}), 403
    conn = get_db(project_id)
    project_name = conn.execute('SELECT value FROM project_info WHERE key = "name"').fetchone()
    project_name = project_name['value'] if project_name else '未命名'
    chapters = conn.execute('SELECT title, content FROM chapters ORDER BY sort_order').fetchall()
    conn.close()
    
    doc = Document()
    doc.add_heading(project_name, 0)
    for ch in chapters:
        doc.add_heading(ch['title'], level=1)
        # 简单处理换行
        for para in (ch['content'] or '').split('\n'):
            if para.strip():
                doc.add_paragraph(para)
            else:
                doc.add_paragraph()
    
    output_path = PROJECTS_DIR / project_id / f'{project_name}.docx'
    doc.save(str(output_path))
    return send_from_directory(str(PROJECTS_DIR / project_id), f'{project_name}.docx', as_attachment=True)

@app.route('/api/projects/<project_id>/export/txt', methods=['GET'])
def export_txt(project_id):
    """导出为纯文本"""
    conn = get_db(project_id)
    project_name = conn.execute('SELECT value FROM project_info WHERE key = "name"').fetchone()
    project_name = project_name['value'] if project_name else '未命名'
    chapters = conn.execute('SELECT title, content FROM chapters ORDER BY sort_order').fetchall()
    conn.close()
    
    content_lines = [project_name, '=' * len(project_name), '']
    for ch in chapters:
        content_lines.append(f'## {ch["title"]}')
        content_lines.append('')
        content_lines.append(ch['content'] or '')
        content_lines.append('')
        content_lines.append('-' * 40)
        content_lines.append('')
    
    output_path = PROJECTS_DIR / project_id / f'{project_name}.txt'
    output_path.write_text('\n'.join(content_lines), encoding='utf-8')
    return send_from_directory(str(PROJECTS_DIR / project_id), f'{project_name}.txt', as_attachment=True)

@app.route('/api/projects/<project_id>/export/pdf', methods=['GET'])
def export_pdf(project_id):
    """导出为 PDF（支持中文）"""
    if not check_feature('export', 'pdf'):
        return jsonify({'error': '免费版不支持 PDF 导出，请升级会员'}), 403
    conn = get_db(project_id)
    project_name = conn.execute('SELECT value FROM project_info WHERE key = "name"').fetchone()
    project_name = project_name['value'] if project_name else '未命名'
    chapters = conn.execute('SELECT title, content FROM chapters ORDER BY sort_order').fetchall()
    conn.close()
    
    output_path = PROJECTS_DIR / project_id / f'{project_name}.pdf'
    
    # 注册中文字体
    font_path = '/System/Library/Fonts/STHeiti Medium.ttc'
    try:
        pdfmetrics.registerFont(TTFont('STHeitiCN', font_path))
        cn_font = 'STHeitiCN'
    except Exception:
        cn_font = 'Helvetica'
    
    doc = SimpleDocTemplate(
        str(output_path), pagesize=A4,
        leftMargin=72, rightMargin=72, topMargin=72, bottomMargin=72
    )
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CNTitle', parent=styles['Title'],
        fontName=cn_font, fontSize=22, spaceAfter=24,
        textColor='#1a1a2e'
    )
    heading_style = ParagraphStyle(
        'CNHeading', parent=styles['Heading2'],
        fontName=cn_font, fontSize=16, spaceBefore=20, spaceAfter=10,
        textColor='#16213e'
    )
    body_style = ParagraphStyle(
        'CNBody', parent=styles['Normal'],
        fontName=cn_font, fontSize=11, leading=20,
        spaceAfter=8, firstLineIndent=22
    )
    
    story = []
    story.append(Paragraph(project_name, title_style))
    story.append(Spacer(1, 12))
    
    for ch in chapters:
        story.append(Paragraph(ch['title'], heading_style))
        for para in (ch['content'] or '').split('\n'):
            if para.strip():
                # 转义 XML 特殊字符
                safe_text = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_text, body_style))
            else:
                story.append(Spacer(1, 8))
        if ch != chapters[-1]:
            story.append(PageBreak())
    
    doc.build(story)
    return send_from_directory(str(PROJECTS_DIR / project_id), f'{project_name}.pdf', as_attachment=True)

# ===== 智能大纲 AI 生成 API =====
@app.route('/api/projects/<project_id>/outline/generate', methods=['POST'])
def ai_generate_outline(project_id):
    """AI 智能生成情节大纲"""
    data = request.json
    premise = data.get('premise', '')
    genre = data.get('genre', '')
    style = data.get('style', '')
    
    if not premise:
        return jsonify({'error': '请提供故事梗概'}), 400
    
    # 获取角色信息作为上下文
    conn = get_db(project_id)
    chars = conn.execute('SELECT name, personality, background FROM characters').fetchall()
    conn.close()
    
    char_context = ''
    if chars:
        char_parts = []
        for c in chars:
            parts = [c['name']]
            if c['personality']: parts.append(f'性格{c["personality"]}')
            if c['background']: parts.append(f'背景：{c["background"]}')
            char_parts.append('，'.join(parts))
        char_context = '已有角色：\n' + '\n'.join(f'- {p}' for p in char_parts)
    
    messages = [
        {'role': 'system', 'content': SURVIVAL_REALISM_HARDENING + '\n\n' + '''你是一位专业的小说大纲策划师。你的任务是根据故事梗概，生成结构清晰的多层级情节大纲。

输出格式要求（严格遵守）：
- 用JSON数组格式输出，每个节点包含 level（层级0-3）、title（标题）、children（子节点数组）
- level 0: 卷/部（如"第一卷"）
- level 1: 章（如"第一章：开端"）  
- level 2: 节（如"第1节：相遇"）
- level 3: 关键情节点

示例输出格式：
[
  {"level":0,"title":"第一卷：启程","children":[
    {"level":1,"title":"第一章：命运之门","children":[
      {"level":2,"title":"第1节：平凡的开始","children":[
        {"level":3,"title":"主角日常生活"},
        {"level":3,"title":"神秘信件到来"}
      ]},
      {"level":2,"title":"第2节：意外发现","children":[]}
    ]}
  ]}
]

直接输出JSON数组，不要加任何解释或markdown标记。'''},
        {'role': 'user', 'content': f'''请根据以下信息生成情节大纲：

故事梗概：{premise}
{'类型/风格：' + genre if genre else ''}
{'写作风格：' + style if style else ''}
{char_context}

请生成一个完整的多层级大纲结构。'''}
    ]
    
    result = call_ai(messages, model=data.get('model'), temperature=0.7, max_tokens=3000)
    return ai_result(result, data.get('model'))

@app.route('/api/projects/<project_id>/outline/import', methods=['POST'])
def import_outline(project_id):
    """导入 AI 生成的大纲节点"""
    data = request.json
    nodes = data.get('nodes', [])
    
    conn = get_db(project_id)
    now = datetime.now().isoformat()
    created_ids = []
    
    def import_node(node, parent_id=None, sort=0):
        node_id = secrets.token_hex(8)
        level = node.get('level', 0)
        title = node.get('title', '未命名')
        conn.execute('''
            INSERT INTO outline (id, parent_id, title, content, level, sort_order, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ''', (node_id, parent_id, title, node.get('content', ''), level, sort, now, now))
        created_ids.append(node_id)
        
        children = node.get('children', [])
        for idx, child in enumerate(children):
            import_node(child, node_id, idx)
    
    for idx, node in enumerate(nodes):
        import_node(node, None, idx)
    
    conn.commit()
    conn.close()
    return jsonify({'success': True, 'count': len(created_ids), 'ids': created_ids})

# ===== AI 世界观与角色自动生成 =====
@app.route('/api/projects/<project_id>/ai/generate-worldbuilding', methods=['POST'])
def ai_generate_worldbuilding(project_id):
    """AI 自动生成世界观设定条目，覆盖全部 8 个分类"""
    if not check_feature('ai_call'):
        return jsonify({'error': '余额不足，请先充值'}), 403

    data = request.json
    premise = data.get('premise', '')
    genre = data.get('genre', '')
    style = data.get('style', '')

    if not premise:
        return jsonify({'error': '请提供故事梗概'}), 400

    cat_list = '\n'.join([f'- {k}: {v}' for k, v in WORLDBUILDING_CATEGORIES.items()])

    messages = [
        {'role': 'system', 'content': f'''你是一位专业的小说世界观构建师。根据故事梗概，为以下每个分类生成 1-3 个设定条目：

{cat_list}

输出格式（JSON 数组）：
[
  {{"category": "geography", "name": "条目名称", "description": "详细描述（100-300字），包括其独特性、对故事的影响"}},
  ...
]

要求：
1. 每个分类至少 1 个条目，核心分类（geography, faction, magic）至少 2 个
2. 条目必须与故事梗概紧密相关，不是通用设定
3. 描述要具体可落地，包含独特细节
4. 条目名称简洁有力（2-8字）

直接输出 JSON 数组，不要加解释或 markdown 代码块。'''},
        {'role': 'user', 'content': f'故事梗概：{premise}\n类型：{genre or "未指定"}\n风格：{style or "未指定"}\n\n请生成完整的世界观设定条目。'}
    ]

    result = call_ai(messages, model=data.get('model'), temperature=0.8, max_tokens=2500)

    if isinstance(result, str) and (result.startswith('__RATE_LIMIT__') or result.startswith('__INSUFFICIENT_BALANCE__')):
        return ai_result(result)

    # 解析 AI 返回的 JSON
    items = safe_json_extract(result, default=[])
    if not isinstance(items, list):
        items = []

    if not items:
        billing = _get_billing()
        return jsonify({'content': result, 'model': data.get('model', 'deepseek'),
                        'billing': billing, 'imported': 0, 'error': '无法解析 AI 生成结果'})

    # 批量插入
    conn = get_db(project_id)
    now = datetime.now().isoformat()
    imported = 0
    for item in items:
        cat = item.get('category', 'other')
        if cat not in WORLDBUILDING_CATEGORIES:
            cat = 'other'
        name = item.get('name', '').strip()
        if not name:
            continue
        desc = item.get('description', '')
        wid = secrets.token_hex(8)
        max_sort = conn.execute(
            'SELECT COALESCE(MAX(sort_order), 0) FROM worldbuilding WHERE category = ?', (cat,)
        ).fetchone()[0]
        conn.execute('''
            INSERT INTO worldbuilding (id, category, name, description, details, sort_order, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ''', (wid, cat, name, desc, '{}', max_sort + 1, now, now))
        imported += 1

    conn.commit()
    conn.close()

    billing = _get_billing()
    return jsonify({
        'content': result,
        'model': data.get('model', 'deepseek'),
        'billing': billing,
        'imported': imported,
        'success': True
    })


@app.route('/api/projects/<project_id>/ai/generate-character', methods=['POST'])
def ai_generate_character(project_id):
    """AI 自动生成角色"""
    if not check_feature('ai_call'):
        return jsonify({'error': '余额不足，请先充值'}), 403

    data = request.json
    name = data.get('name', '主角').strip()
    premise = data.get('premise', '')
    genre = data.get('genre', '')
    style = data.get('style', '')

    # 获取世界观上下文
    world_ctx = _build_world_context(project_id)

    # 获取已有角色避免重复
    conn = get_db(project_id)
    existing = conn.execute('SELECT name, personality, background, goal FROM characters').fetchall()
    conn.close()
    existing_str = ''
    if existing:
        existing_str = '已有角色：\n' + '\n'.join(
            [f'- {c["name"]}：{c.get("personality","")[:80]}' for c in existing])

    messages = [
        {'role': 'system', 'content': '''你是一位专业的小说角色设计师。根据故事信息，生成一个完整的角色设定。

输出格式（JSON）：
{
  "name": "角色名",
  "gender": "男/女/其他",
  "age": "年龄（数字或描述如"外表20岁"）",
  "personality": "性格描述（80-200字），包括核心性格特质、行为模式、优缺点",
  "background": "角色背景（80-200字），包括出身、经历、关键转折事件",
  "goal": "角色目标/动机（30-80字），他们想要什么、为什么",
  "appearance": "外貌描述（30-100字），关键特征",
  "notes": "补充说明（可选），与其他角色的关系、在故事中的作用等"
}

要求：
1. 角色的性格必须有层次，不能扁平
2. 角色背景必须与故事世界观一致
3. 角色目标必须驱动剧情
4. 避免与已有角色完全重复

直接输出 JSON，不要加解释或 markdown 代码块。'''},
        {'role': 'user', 'content': f'''请为以下故事生成角色「{name}」的完整设定：

故事梗概：{premise or "未指定"}
类型：{genre or "未指定"}
风格：{style or "未指定"}
{world_ctx}
{existing_str}'''}
    ]

    result = call_ai(messages, model=data.get('model'), temperature=0.8, max_tokens=1500)
    return ai_result(result, data.get('model'))


@app.route('/api/projects/<project_id>/ai/generate-characters-batch', methods=['POST'])
def ai_generate_characters_batch(project_id):
    """AI 批量生成多个角色（3-5个核心角色）"""
    if not check_feature('ai_call'):
        return jsonify({'error': '余额不足，请先充值'}), 403

    data = request.json
    premise = data.get('premise', '')
    genre = data.get('genre', '')
    style = data.get('style', '')

    if not premise:
        return jsonify({'error': '请提供故事梗概'}), 400

    world_ctx = _build_world_context(project_id)

    messages = [
        {'role': 'system', 'content': '''你是一位专业的小说角色设计师。根据故事梗概，生成 3-5 个核心角色。

输出格式（JSON 数组）：
[
  {
    "name": "角色名",
    "gender": "男/女/其他",
    "age": "年龄",
    "personality": "性格描述（50-150字）",
    "background": "角色背景（50-150字）",
    "goal": "角色目标/动机（20-60字）",
    "appearance": "外貌描述（20-80字）",
    "notes": "在故事中的作用"
  },
  ...
]

要求：
1. 必须包含主角、主要配角、可能的反派
2. 角色之间有关系和冲突
3. 每个角色的目标应与故事主线相关
4. 性格要有区分度，不能雷同

直接输出 JSON 数组，不要加解释或 markdown 代码块。'''},
        {'role': 'user', 'content': f'故事梗概：{premise}\n类型：{genre or "未指定"}\n风格：{style or "未指定"}\n{world_ctx}\n\n请生成核心角色群。'}
    ]

    result = call_ai(messages, model=data.get('model'), temperature=0.8, max_tokens=2500)

    if isinstance(result, str) and (result.startswith('__RATE_LIMIT__') or result.startswith('__INSUFFICIENT_BALANCE__')):
        return ai_result(result)

    # 解析并批量创建
    chars = safe_json_extract(result, default=[])
    if not isinstance(chars, list):
        chars = []

    if not chars:
        billing = _get_billing()
        return jsonify({'content': result, 'model': data.get('model', 'deepseek'),
                        'billing': billing, 'imported': 0})

    conn = get_db(project_id)
    now = datetime.now().isoformat()
    imported = 0
    for ch in chars:
        name = ch.get('name', '').strip()
        if not name:
            continue
        char_id = secrets.token_hex(8)
        conn.execute('''
            INSERT INTO characters (id, name, gender, age, personality, background, goal, appearance, notes, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            char_id, name, ch.get('gender', ''), str(ch.get('age', '')),
            ch.get('personality', ''), ch.get('background', ''), ch.get('goal', ''),
            ch.get('appearance', ''), ch.get('notes', ''), now, now
        ))
        imported += 1

    conn.commit()
    conn.close()

    billing = _get_billing()
    return jsonify({
        'content': result,
        'model': data.get('model', 'deepseek'),
        'billing': billing,
        'imported': imported,
        'success': True
    })

# ===== 一键全书生成 =====
@app.route('/api/projects/<project_id>/ai/generate-book-plan', methods=['POST'])
def ai_generate_book_plan(project_id):
    """AI 生成章节计划"""
    if not check_feature('ai_call'):
        return jsonify({'error': '余额不足，请先充值'}), 403

    data = request.json
    premise = data.get('premise', '')
    genre = data.get('genre', '')
    style = data.get('style', '')
    chapter_count = int(data.get('chapter_count', 10))
    chapter_count = max(3, min(30, chapter_count))

    if not premise.strip():
        return jsonify({'error': '请提供故事构思'}), 400

    conn = get_db(project_id)
    chars = conn.execute('SELECT name, personality, background, goal FROM characters').fetchall()
    conn.close()

    char_ctx = ''
    if chars:
        char_parts = []
        for c in chars:
            parts = [c['name']]
            if c['personality']: parts.append(c['personality'][:50])
            if c['goal']: parts.append(c['goal'][:40])
            char_parts.append('，'.join(parts))
        char_ctx = '已有角色：\n' + '\n'.join(f'- {p}' for p in char_parts)

    messages = [
        {'role': 'system', 'content': SURVIVAL_REALISM_HARDENING + '\n\n' + f'''你是一位专业的小说策划师。根据故事构思，生成{chapter_count}章小说章节计划。

输出格式（JSON 数组）：
[
  {{"title": "第1章 具体标题", "summary": "本章核心内容概要（50-80字），包括主要情节、出场角色、关键转折"}},
  ...
]

要求：
1. 标题格式为"第X章 具体标题"，有吸引力
2. 章节间情节递进，不能孤立
3. 前3章快速建立世界观和核心矛盾
4. 中间章节展开冲突、引入新线索
5. 最后3章推向高潮并收尾
6. 每章概要具体可执行
7. {"严格遵循已有角色设定" if chars else "可自由设计角色"}

直接输出 JSON 数组，不要加解释或 markdown 代码块。'''},
        {'role': 'user', 'content': f'请为以下故事构思生成{chapter_count}章章节计划：\n\n故事构思：{premise}\n类型：{genre or "未指定"}\n写作风格：{style or "未指定"}\n{char_ctx}'}
    ]

    result = call_ai(messages, model=data.get('model'), temperature=0.7, max_tokens=3000)
    return ai_result(result, data.get('model'))


@app.route('/api/projects/<project_id>/ai/generate-book-chapter', methods=['POST'])
def ai_generate_book_chapter(project_id):
    """AI 生成单章正文"""
    if not check_feature('ai_call'):
        return jsonify({'error': '余额不足，请先充值'}), 403

    data = request.json
    chapter_index = data.get('chapter_index', 1)
    chapter_title = data.get('chapter_title', '')
    chapter_summary = data.get('chapter_summary', '')
    previous_summary = data.get('previous_summary', '')
    chapter_count = data.get('chapter_count', 10)
    chapter_id = data.get('chapter_id', '')

    if not chapter_title:
        return jsonify({'error': '缺少章节标题'}), 400

    condensed = build_condensed_context(project_id, chapter_id)
    world_ctx = _build_world_context(project_id)

    previous_block = ''
    if previous_summary:
        previous_block = f'【前文摘要】\n{previous_summary}\n\n'

    pace_guide = '开篇建立世界观和悬念，快速抓住读者' if chapter_index <= 3 else '逐步推进情节，展开冲突，深化角色关系' if chapter_index <= chapter_count - 3 else '推向高潮，解决核心矛盾，准备收尾'

    messages = [
        {'role': 'system', 'content': SURVIVAL_REALISM_HARDENING + '\n' + BAIMIAO_HARDENING + '\n' + f'''你是一位专业小说作家，正在创作一部{chapter_count}章小说的第{chapter_index}章。

{world_ctx}

{condensed}

写作要求：
1. 本章标题为"【{chapter_title}】"，严格围绕概要创作
2. 本章概要：{chapter_summary}
3. 直接输出正文，不要写标题
4. 保持与【前文摘要】的连贯性
5. 字数 1500-3500 字，有起承转合
6. 避免AI味：用自然短句，消除"然而""此外""总之"等高频词
7. 角色对话有各自口语特点
8. 用具体感官细节替代抽象形容
9. 节奏指导：{pace_guide}
10. 保持原创性，不用网文常见套话

直接输出正文，不用 markdown，不加开场白。'''},
        {'role': 'user', 'content': f'{previous_block}请创作第{chapter_index}章：{chapter_title}\n\n本章概要：{chapter_summary}'}
    ]

    result = call_ai(messages, model=data.get('model'), temperature=0.75, max_tokens=4000)
    return ai_result(result, data.get('model'))


# ===== 小说转剧本 =====
SCRIPT_PROMPTS = {
    'movie': '''你是一位资深影视编剧。请将以下小说内容改编为电影剧本格式。

要求：
1. 使用标准电影剧本格式：场景标题（INT./EXT. 地点 - 时间）、动作描写、角色对白
2. 角色名在对白前，顶格；对白另起一行缩进
3. 动作描写用现在时态，简洁有力
4. 保留原小说的核心情节和对话，删除冗长的内心描写
5. 加入必要的镜头提示（特写/全景/切至等），用【】标记
6. 整体节奏紧凑，适合电影时长（90-120分钟）
7. 场景之间用空行分隔

格式示例：
INT. 咖啡厅 - 白天

【全景】咖啡厅内客人稀少，角落坐着一男一女。

李明
（放下咖啡杯）
你到底想说什么？

王芳
（直视对方）
我怀孕了。

李明的手停在半空，表情凝固。''',

    'tv': '''你是一位电视剧编剧。请将以下小说内容改编为电视剧剧本格式。

要求：
1. 使用标准电视剧剧本格式，按"集"划分（每集约5000-8000字内容为一集）
2. 每集开头标注"第X集"，集末留悬念钩子【下集预告】
3. 场景标题：INT./EXT. 地点 - 时间
4. 角色对白格式：角色名顶格，对白在下一行
5. 动作描写用现在时态，节奏比电影稍缓
6. 保留角色间的情感交流和次要情节线
7. 商业广告插播点用【※】标记（每集约2-3处）
8. 对白要口语化，符合电视剧观众的听觉习惯

格式示例：
第一集

INT. 公司办公室 - 上午

【全景】开放式办公区，员工们各自忙碌。

张伟
（盯着电脑屏幕，皱眉）
这个数据不对——谁做的报表？
''',

    'stage': '''你是一位舞台剧编剧。请将以下小说内容改编为舞台剧本格式。

要求：
1. 使用标准舞台剧本格式：分"幕"和"场"（第一幕/第二幕）
2. 每幕开头写【舞台布景】描述舞台设计和氛围
3. 角色动作和表情用（括号）标注在角色名后
4. 对白自然流畅，适合演员朗读
5. 独白用【独白】标记，旁白用【旁白】标记
6. 舞台指示用【】标记：灯光变化、音效、道具移动等
7. 简化场景转换，将多个地点浓缩为代表性场景
8. 保留戏剧冲突和高潮，删减不必要的过渡

格式示例：
第一幕

【舞台布景】
舞台中央是一张旧木桌。左侧是书架，右侧是沙发。暖黄色灯光笼罩整个舞台。

第一场

（王芳独自坐在沙发上，翻看旧相册。灯光渐暗，追光打在她身上。）

王芳
（轻声翻页，目光停留在一张照片上）
十年了……你还会回来吗？
（合上相册，起身走向书架）

【灯光渐变——转为冷白色】

（李明从舞台左侧入场，手里提着公文包）''',
}

@app.route('/api/projects/<project_id>/ai/convert-to-script', methods=['POST'])
def ai_convert_to_script(project_id):
    """将小说转换为剧本格式"""
    if not check_feature('ai_call'):
        return jsonify({'error': '余额不足，请先充值'}), 403

    data = request.json
    chapter_ids = data.get('chapter_ids', [])
    script_type = data.get('script_type', 'movie')
    if script_type not in SCRIPT_PROMPTS:
        script_type = 'movie'

    conn = get_db(project_id)
    if chapter_ids and len(chapter_ids) > 0:
        placeholders = ','.join('?' * len(chapter_ids))
        rows = conn.execute(
            f'SELECT title, content FROM chapters WHERE id IN ({placeholders}) ORDER BY sort_order',
            chapter_ids
        ).fetchall()
    else:
        rows = conn.execute('SELECT title, content FROM chapters ORDER BY sort_order').fetchall()
    conn.close()

    if not rows:
        return jsonify({'error': '项目中没有章节内容'}), 400

    # 构建小说文本（限制总长度）
    novel_text_parts = []
    total_chars = 0
    max_chars = 12000
    for r in rows:
        content = (r['content'] or '').strip()
        if not content:
            continue
        chapter_block = f'\n## {r["title"]}\n\n{content[:3000]}'
        if total_chars + len(chapter_block) > max_chars:
            remaining = max_chars - total_chars
            if remaining > 200:
                chapter_block = chapter_block[:remaining] + '\n...'
            else:
                break
        novel_text_parts.append(chapter_block)
        total_chars += len(chapter_block)

    novel_text = '\n'.join(novel_text_parts)

    messages = [
        {'role': 'system', 'content': SCRIPT_PROMPTS[script_type]},
        {'role': 'user', 'content': f'请将以下小说内容改编为剧本：\n\n{novel_text}'}
    ]

    result = call_ai(messages, model=data.get('model'), temperature=0.5, max_tokens=4000)
    return ai_result(result, data.get('model'))


# ===== AI 生成章节自动保存 =====
@app.route('/api/projects/<project_id>/chapters/from-ai', methods=['POST'])
def create_chapter_from_ai(project_id):
    """从 AI 生成内容创建章节"""
    data = request.json
    title = data.get('title', 'AI 生成章节')
    content = data.get('content', '')
    source = data.get('source', 'ai')
    
    now = datetime.now().isoformat()
    chapter_id = secrets.token_hex(8)
    word_count = len(content.replace(' ', '').replace('\n', ''))
    
    conn = get_db(project_id)
    
    # 确保 source 列存在
    try:
        conn.execute('ALTER TABLE chapters ADD COLUMN source TEXT DEFAULT "manual"')
        conn.commit()
    except sqlite3.OperationalError:
        pass  # 列已存在
    
    max_sort = conn.execute('SELECT COALESCE(MAX(sort_order), 0) as mx FROM chapters').fetchone()['mx']
    conn.execute('''
        INSERT INTO chapters (id, title, content, sort_order, created_at, updated_at, word_count, source)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    ''', (chapter_id, title, content, max_sort + 1, now, now, word_count, source))
    conn.commit()
    conn.close()
    
    return jsonify({'id': chapter_id, 'title': title, 'word_count': word_count, 'source': source})

# ===== 智能导入 API =====
MAX_UPLOAD_SIZE = 500 * 1024  # 500KB

@app.route('/api/projects/<project_id>/import/upload', methods=['POST'])
def import_upload(project_id):
    """接收上传文件，返回纯文本内容"""
    if 'file' not in request.files:
        return jsonify({'error': '请选择文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400
    
    filename = file.filename.lower()
    content = ''
    
    try:
        if filename.endswith('.txt') or filename.endswith('.md'):
            raw = file.read(MAX_UPLOAD_SIZE + 1)
            if len(raw) > MAX_UPLOAD_SIZE:
                return jsonify({'error': '文件过大，请限制在500KB以内'}), 400
            content = raw.decode('utf-8', errors='replace')
        elif filename.endswith('.docx'):
            from io import BytesIO
            raw = file.read(MAX_UPLOAD_SIZE * 2)
            if len(raw) > MAX_UPLOAD_SIZE * 2:
                return jsonify({'error': '文件过大，请限制在500KB以内'}), 400
            doc = Document(BytesIO(raw))
            content = '\n'.join(p.text for p in doc.paragraphs)
        else:
            return jsonify({'error': '仅支持 .txt / .md / .docx 格式'}), 400
    except Exception as e:
        return jsonify({'error': f'文件读取失败: {str(e)}'}), 400
    
    if not content.strip():
        return jsonify({'error': '文件内容为空'}), 400
    
    return jsonify({
        'content': content,
        'size': len(content),
        'filename': file.filename
    })

@app.route('/api/projects/<project_id>/import/analyze', methods=['POST'])
def import_analyze(project_id):
    """AI 分析文本内容，识别类型并提取结构化数据"""
    data = request.json
    text = data.get('text', '')
    filename = data.get('filename', '')
    
    if not text.strip():
        return jsonify({'error': '请提供文本内容'}), 400
    
    # 截取前 4000 字符用于 AI 分析
    sample = text[:4000]
    
    messages = [
        {'role': 'system', 'content': '''你是一位专业的小说编辑AI。请分析以下文本，判断它属于什么类型的小说素材，并尽可能提取出结构化数据。

输出格式（严格遵守JSON格式，不要加任何解释或markdown标记）：
{
  "type": "outline|characters|chapters|mixed",
  "confidence": 0.0-1.0,
  "outline": [
    {"level": 0-3, "title": "节点标题", "children": []}
  ],
  "characters": [
    {"name": "角色名", "gender": "男/女", "personality": "性格", "background": "背景", "role": "主角/配角/反派"}
  ],
  "chapters": [
    {"title": "章节标题", "content": "章节内容摘要或完整文本"}
  ],
  "summary": "一段简要说明（20字以内）"
}

识别规则：
- outline: 有明显的层级结构（卷/章/节），如"第一卷"、"第一章"等
- characters: 有角色名称、性格、背景描述等人物介绍
- chapters: 有连续叙事内容，可能包含"第一章"等章节标记
- mixed: 同时包含以上多种类型'''},
        {'role': 'user', 'content': f'文件名: {filename}\n\n文本内容（前4000字）:\n{sample[:4000]}'}
    ]
    
    result = call_ai(messages, model=data.get('model'), temperature=0.3, max_tokens=3000)
    return jsonify({'analysis': result, 'model': data.get('model', get_setting('default_model', 'deepseek'))})

@app.route('/api/projects/<project_id>/import/apply', methods=['POST'])
def import_apply(project_id):
    """将 AI 分析结果导入数据库"""
    data = request.json
    items = data.get('items', {})  # {outline: [...], characters: [...], chapters: [...]}
    
    conn = get_db(project_id)
    now = datetime.now().isoformat()
    result = {'outline': 0, 'characters': 0, 'chapters': 0}
    
    # 导入大纲
    for node_data in items.get('outline', []):
        def import_node(node, parent_id=None, sort=0):
            node_id = secrets.token_hex(8)
            level = node.get('level', 0)
            title = node.get('title', '未命名')
            conn.execute('''
                INSERT INTO outline (id, parent_id, title, content, level, sort_order, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (node_id, parent_id, title, node.get('content', ''), level, sort, now, now))
            result['outline'] += 1
            for idx, child in enumerate(node.get('children', [])):
                import_node(child, node_id, idx)
        import_node(node_data)
    
    # 导入角色
    for char_data in items.get('characters', []):
        char_id = secrets.token_hex(8)
        conn.execute('''
            INSERT INTO characters (id, name, gender, age, personality, background, goal, appearance, notes, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            char_id, char_data.get('name', ''), char_data.get('gender', ''),
            char_data.get('age', ''), char_data.get('personality', ''),
            char_data.get('background', ''), char_data.get('goal', ''),
            char_data.get('appearance', ''), char_data.get('notes', ''), now, now
        ))
        result['characters'] += 1
    
    # 导入章节
    max_sort = conn.execute('SELECT COALESCE(MAX(sort_order), 0) as mx FROM chapters').fetchone()['mx']
    for i, ch_data in enumerate(items.get('chapters', [])):
        ch_id = secrets.token_hex(8)
        content = ch_data.get('content', '')
        word_count = len(content.replace(' ', '').replace('\n', ''))
        conn.execute('''
            INSERT INTO chapters (id, title, content, sort_order, created_at, updated_at, word_count, source)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ''', (ch_id, ch_data.get('title', '新章节'), content, max_sort + i + 1, now, now, word_count, 'import'))
        result['chapters'] += 1
    
    conn.commit()
    conn.close()
    return jsonify({'success': True, 'imported': result})

# ===== 会员系统 =====
import hashlib
import hmac

LICENSE_DB = BASE_DIR / 'license.db'

def init_license_db():
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    conn.execute('PRAGMA journal_mode=WAL')
    conn.execute('PRAGMA busy_timeout=5000')
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS licenses (
            id TEXT PRIMARY KEY,
            tier TEXT NOT NULL,
            status TEXT DEFAULT 'active',
            activated_at TEXT,
            expires_at TEXT,
            device_id TEXT,
            customer_email TEXT,
            customer_name TEXT,
            created_at TEXT
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS activation_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            license_id TEXT,
            action TEXT,
            timestamp TEXT,
            ip TEXT,
            detail TEXT
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS license_config (
            key TEXT PRIMARY KEY,
            value TEXT
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS settings (
            key TEXT PRIMARY KEY,
            value TEXT
        )
    ''')
    # 默认设置：启用许可系统
    c.execute('INSERT OR IGNORE INTO license_config (key, value) VALUES (?, ?)', ('enabled', 'true'))
    # 默认 API Key（DeepSeek 已有，Claude 留空）
    # API Key 不再存储于数据库，统一从环境变量读取
    c.execute('INSERT OR IGNORE INTO settings (key, value) VALUES (?, ?)', ('default_model', 'deepseek'))
    # ===== 充值模式新表 =====
    c.execute('''
        CREATE TABLE IF NOT EXISTS user_balances (
            user_id TEXT PRIMARY KEY,
            balance REAL DEFAULT 0.0,
            free_daily_credits REAL DEFAULT 0.02,
            free_daily_used REAL DEFAULT 0.0,
            free_daily_date TEXT,
            total_recharged REAL DEFAULT 0.0,
            total_cost REAL DEFAULT 0.0,
            created_at TEXT,
            updated_at TEXT
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS ai_transactions (
            id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL,
            model TEXT NOT NULL,
            input_tokens INTEGER DEFAULT 0,
            output_tokens INTEGER DEFAULT 0,
            total_tokens INTEGER DEFAULT 0,
            cost REAL NOT NULL,
            balance_before REAL,
            balance_after REAL,
            free_credit_used REAL DEFAULT 0,
            endpoint TEXT,
            created_at TEXT
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS feature_usage (
            user_id TEXT NOT NULL,
            feature TEXT NOT NULL,
            usage_month TEXT NOT NULL,
            call_count INTEGER DEFAULT 0,
            PRIMARY KEY (user_id, feature, usage_month)
        )
    ''')
    # 迁移标记
    c.execute('INSERT OR IGNORE INTO license_config (key, value) VALUES (?, ?)', ('system_version', '2.0'))
    conn.commit()
    conn.close()

init_license_db()

# ===== 充值模式：用户标识 =====

def get_user_id():
    """获取当前用户的唯一标识。
    Supabase 模式：使用 JWT 中的 user_id（多租户隔离）
    SQLite 模式：使用设备级随机 ID（本地单机）
    """
    # 优先使用 Supabase JWT 用户
    if _HAS_SUPABASE_REPO and is_supabase_mode():
        uid = get_current_user_id()
        if uid:
            return uid
        # 降级：Supabase 配置了但 JWT 缺失 → 创建匿名 ID
        uid = _get_or_create_anonymous_id()
        return uid

    # SQLite 本地模式
    return _get_or_create_anonymous_id()

def _get_or_create_anonymous_id():
    """SQLite 本地模式：创建/获取设备级匿名 ID"""
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    row = conn.execute("SELECT value FROM license_config WHERE key = 'user_id'").fetchone()
    if row:
        conn.close()
        return row['value']
    user_id = secrets.token_hex(16)
    conn.execute('INSERT OR IGNORE INTO license_config (key, value) VALUES (?, ?)', ('user_id', user_id))
    conn.commit()
    conn.close()
    _ensure_balance_record(user_id)
    return user_id

def _ensure_balance_record(user_id, conn=None):
    own_conn = conn is None
    if own_conn:
        conn = sqlite3.connect(str(LICENSE_DB))
    now = datetime.now().isoformat()
    conn.execute('''
        INSERT OR IGNORE INTO user_balances (user_id, balance, free_daily_credits, free_daily_used, free_daily_date, created_at, updated_at)
        VALUES (?, 0.0, 0.02, 0.0, ?, ?, ?)
    ''', (user_id, datetime.now().strftime('%Y-%m-%d'), now, now))
    if own_conn:
        conn.commit()
        conn.close()

# ===== RLS 多租户隔离（占位） =====

def _get_user_project_path(user_id, project_id):
    """获取用户项目目录路径。
    当前为全局模式（所有用户共享同一项目空间）。
    接入云端数据库后，取消注释第二行即可激活租户隔离。
    """
    # RLS ACTIVATION: 取消下面这行注释，同时注释当前的 global 路径
    # return PROJECTS_DIR / user_id / project_id
    return PROJECTS_DIR / project_id  # global: 无租户隔离

def _get_user_price(model):
    """获取模型单价（元/千tokens）"""
    prices = {'deepseek': 0.01, 'gemini': 0.04, 'claude': 0.12}
    return prices.get(model, 0.01)

# ===== 模型套餐系统 =====

MODEL_TIERS = {
    'free': {
        'name': '免费版',
        'models': ['deepseek'],
        'desc': 'DeepSeek Chat · 每日0.02元免费额度',
        'min_recharge': 0
    },
    'basic': {
        'name': '基础版',
        'models': ['deepseek', 'gemini'],
        'desc': 'DeepSeek + Gemini 2.5 Flash · 更多选择',
        'min_recharge': 1
    },
    'pro': {
        'name': '专业版',
        'models': ['deepseek', 'gemini', 'claude'],
        'desc': '全部模型 · DeepSeek + Gemini + Claude Sonnet 4',
        'min_recharge': 50
    },
    'premium': {
        'name': '旗舰版',
        'models': ['deepseek', 'gemini', 'claude'],
        'desc': '全部模型 + 优先支持 · 专业创作体验',
        'min_recharge': 200
    }
}

MODEL_META = {
    'deepseek': {'name': 'DeepSeek Chat', 'icon': '🤖', 'price': '¥0.01/千tokens', 'badge': '高性价比'},
    'gemini': {'name': 'Gemini 2.5 Flash', 'icon': '⚡', 'price': '¥0.04/千tokens', 'badge': '均衡之选'},
    'claude': {'name': 'Claude Sonnet 4', 'icon': '🧠', 'price': '¥0.12/千tokens', 'badge': '顶级品质'}
}

def _get_user_tier(user_id):
    """根据用户累计充值金额确定套餐等级"""
    # Supabase 模式：从 users 表读取
    if _HAS_SUPABASE_REPO and is_supabase_mode():
        repo = get_repo()
        if repo:
            tier = repo.get_user_tier()
            return tier

    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    row = conn.execute('SELECT COALESCE(total_recharged, 0) as total FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()
    conn.close()
    total = row['total'] if row else 0
    if total >= 200:
        return 'premium'
    elif total >= 50:
        return 'pro'
    elif total >= 1:
        return 'basic'
    return 'free'

def _get_available_models(user_id):
    """获取用户当前可用的模型列表"""
    tier = _get_user_tier(user_id)
    return MODEL_TIERS.get(tier, MODEL_TIERS['free'])['models']

@app.route('/api/user/tier', methods=['GET'])
def api_user_tier():
    """获取当前用户的套餐等级和可用模型"""
    user_id = get_user_id()
    tier = _get_user_tier(user_id)
    tier_info = dict(MODEL_TIERS[tier])
    tier_info['id'] = tier
    available_models = tier_info['models']
    tier_info['models_detail'] = [dict(MODEL_META[m], id=m) for m in available_models]
    # 充值信息
    # Supabase 模式：余额系统尚未集成，返回默认值
    total_recharged = 0.0
    balance = 0.0
    if _HAS_SUPABASE_REPO and is_supabase_mode():
        # TODO: 将 user_balances 表迁移到 Supabase 后启用真实余额查询
        pass
    else:
        conn = sqlite3.connect(str(LICENSE_DB))
        conn.row_factory = sqlite3.Row
        row = conn.execute('SELECT total_recharged, balance FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()
        conn.close()
        total_recharged = round(row['total_recharged'] or 0, 2) if row else 0
        balance = round(row['balance'] or 0, 2) if row else 0

    tier_info['total_recharged'] = total_recharged
    tier_info['balance'] = balance
    tier_info['next_tier'] = None
    if tier == 'free':
        tier_info['next_tier'] = 'basic'
        tier_info['next_tier_name'] = '基础版'
        tier_info['next_tier_need'] = '充值任意金额即可升级'
    elif tier == 'basic':
        tier_info['next_tier'] = 'pro'
        tier_info['next_tier_name'] = '专业版'
        tier_info['next_tier_need'] = f'还需充值 ¥{50 - tier_info["total_recharged"]:.2f}'
    elif tier == 'pro':
        tier_info['next_tier'] = 'premium'
        tier_info['next_tier_name'] = '旗舰版'
        tier_info['next_tier_need'] = f'还需充值 ¥{200 - tier_info["total_recharged"]:.2f}'
    # SaaS 会员阶梯：后端 4 级 → 前端 3 级映射
    tier_map = {'free': 'free', 'basic': 'pro', 'pro': 'pro', 'premium': 'platinum'}
    tier_info['frontend_tier'] = tier_map.get(tier, 'free')
    # Pro 层级月度预测配额
    if tier in ('basic', 'pro', 'premium'):
        tier_info['predict_retention_limit'] = None if tier == 'premium' else 50
        tier_info['predict_retention_used'] = _get_feature_usage(user_id, 'predict_retention') if tier != 'premium' else 0
    else:
        tier_info['predict_retention_limit'] = 0
        tier_info['predict_retention_used'] = 0
    return jsonify(tier_info)

def _is_admin():
    """检查当前请求是否来自管理员（本地请求或设置了 admin token）"""
    # 安全方式：本地请求 + admin_token 验证
    admin_token = get_setting('admin_token', '')
    if admin_token:
        auth = request.headers.get('X-Admin-Token', '')
        if auth == admin_token:
            return True
    # 本地回环地址也视为管理员
    return request.remote_addr in ('127.0.0.1', '::1', 'localhost')

@app.route('/api/admin/save-api-key', methods=['POST'])
def api_admin_save_key():
    """API Key 管理已迁移至环境变量（.env 文件）"""
    return jsonify({'error': 'API Key 现在通过环境变量配置，请编辑 .env 文件后重启服务', 'env_vars': ['DEEPSEEK_API_KEY', 'CLAUDE_API_KEY', 'GEMINI_API_KEY']}), 410

def _has_paid(user_id):
    """用户是否充过值"""
    # Supabase 模式：检查 user_tier（非 free 即为付费用户）
    if _HAS_SUPABASE_REPO and is_supabase_mode():
        repo = get_repo()
        if repo:
            tier = repo.get_user_tier()
            return tier != 'free'

    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    row = conn.execute('SELECT total_recharged FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()
    conn.close()
    return row and row['total_recharged'] > 0

def _increment_feature_usage(user_id, feature):
    """原子递增功能月调用次数，用于 Pro 层级配额控制"""
    month = datetime.now().strftime('%Y-%m')
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.execute('PRAGMA busy_timeout=5000')
    conn.execute('''
        INSERT INTO feature_usage (user_id, feature, usage_month, call_count)
        VALUES (?, ?, ?, 1)
        ON CONFLICT(user_id, feature, usage_month)
        DO UPDATE SET call_count = call_count + 1
    ''', (user_id, feature, month))
    conn.commit()
    conn.close()

def _get_feature_usage(user_id, feature):
    """查询当前月功能已用次数"""
    month = datetime.now().strftime('%Y-%m')
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    row = conn.execute(
        'SELECT call_count FROM feature_usage WHERE user_id=? AND feature=? AND usage_month=?',
        (user_id, feature, month)
    ).fetchone()
    conn.close()
    return row['call_count'] if row else 0

# ===== 旧版许可证系统（保留向后兼容，新流程不再调用） =====

def generate_license_key(tier):
    """[已废弃] 生成许可证密钥，充值模式不再使用"""
    tier_codes = {'day': 'DAY', 'week': 'WEK', 'month': 'MON', 'year': 'YER'}
    code = tier_codes.get(tier, 'FR')
    random_hex = secrets.token_hex(4).upper()
    checksum = hmac.new(LICENSE_SECRET.encode(), f'{code}{random_hex}'.encode(), hashlib.sha256).hexdigest()[:4].upper()
    return f'NVS-{code}-{random_hex}-{checksum}'

def verify_license_key(key):
    """验证许可证密钥格式和校验和"""
    parts = key.strip().upper().split('-')
    if len(parts) != 4 or parts[0] != 'NVS':
        return False, '格式错误'
    tier_code, random_hex, checksum = parts[1], parts[2], parts[3]
    expected = hmac.new(LICENSE_SECRET.encode(), f'{tier_code}{random_hex}'.encode(), hashlib.sha256).hexdigest()[:4].upper()
    if checksum != expected:
        return False, '密钥无效'
    return True, tier_code

def get_tier_expiry(tier):
    """获取等级对应的过期时间"""
    durations = {'day': 86400, 'week': 604800, 'month': 2592000, 'year': 31536000}
    return int(time.time()) + durations.get(tier, 0)

def get_tier_limits(tier):
    """[已废弃] 旧版等级限制，新系统使用余额模式"""
    limits = {
        'free': {'projects': 1, 'ai_per_day': 3, 'exports': ['txt']},
        'day': {'projects': 10, 'ai_per_day': 999, 'exports': ['txt', 'docx', 'pdf']},
        'week': {'projects': 20, 'ai_per_day': 999, 'exports': ['txt', 'docx', 'pdf']},
        'month': {'projects': 50, 'ai_per_day': 9999, 'exports': ['txt', 'docx', 'pdf']},
        'year': {'projects': 9999, 'ai_per_day': 999999, 'exports': ['txt', 'docx', 'pdf']},
    }
    return limits.get(tier, limits['free'])

def get_current_license():
    """[已废弃] 获取旧版许可证，新系统使用 get_user_id() + 余额"""
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    lic = conn.execute(
        "SELECT * FROM licenses WHERE status = 'active' AND (expires_at IS NULL OR expires_at > ?) ORDER BY created_at DESC LIMIT 1",
        (datetime.now().isoformat(),)
    ).fetchone()
    conn.close()
    return dict(lic) if lic else None

def check_feature(feature, value=None):
    """余额模式下的功能门控（Supabase 模式从 users 表读取等级）"""
    user_id = get_user_id()
    paid = _has_paid(user_id)

    if feature == 'create_project':
        # Supabase 模式：RLS 自动限制，不在此处计数
        if _HAS_SUPABASE_REPO and is_supabase_mode():
            return True
        proj_count = len([p for p in PROJECTS_DIR.iterdir() if p.is_dir() and (p / 'novel.db').exists()])
        limit = 50 if paid else 1
        return proj_count < limit

    if feature == 'ai_call':
        # Supabase 模式：余额系统尚未集成到 Supabase，暂允许所有调用
        # TODO: 将余额/充值系统迁移到 Supabase 后启用真实校验
        if _HAS_SUPABASE_REPO and is_supabase_mode():
            return True
        # 余额模式：只要余额 > 0 或有免费额度即可
        conn = sqlite3.connect(str(LICENSE_DB))
        conn.row_factory = sqlite3.Row
        row = conn.execute('SELECT balance, free_daily_credits, free_daily_used, free_daily_date FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()
        conn.close()
        if not row:
            return False
        today = datetime.now().strftime('%Y-%m-%d')
        free_remaining = (row['free_daily_credits'] or 0.02) if (row['free_daily_date'] or '') != today else max(0, (row['free_daily_credits'] or 0.02) - (row['free_daily_used'] or 0.0))
        return (row['balance'] or 0.0) + free_remaining > 0

    if feature == 'export':
        allowed = ['txt'] if not paid else ['txt', 'docx', 'pdf']
        return value in allowed

    if feature == 'premium':
        return paid

    # ===== SaaS 会员阶梯功能门控 =====
    if feature == 'predict_retention':
        tier = _get_user_tier(user_id)
        if tier == 'premium':
            return True
        if tier in ('basic', 'pro'):
            return _get_feature_usage(user_id, 'predict_retention') < 50
        return False  # free 不可用

    if feature == 'plot_deviation':
        return paid  # 充过值即可用

    return True

def log_license_action(license_id, action, detail='', conn=None):
    """[已废弃] 旧版操作日志，保留向后兼容"""
    close_conn = False
    if conn is None:
        conn = sqlite3.connect(str(LICENSE_DB))
        close_conn = True
    conn.execute(
        'INSERT INTO activation_log (license_id, action, timestamp, ip, detail) VALUES (?, ?, ?, ?, ?)',
        (license_id, action, datetime.now().isoformat(), request.remote_addr or '127.0.0.1', detail)
    )
    if close_conn:
        conn.commit()
        conn.close()

@app.route('/api/license/status', methods=['GET'])
def license_status():
    """余额模式下的用户状态"""
    user_id = get_user_id()

    # Supabase 模式：余额系统尚未集成，返回基本默认值
    if _HAS_SUPABASE_REPO and is_supabase_mode():
        paid = _has_paid(user_id)
        # TODO: 将 user_balances 表迁移到 Supabase 后启用真实余额
        return jsonify({
            'mode': 'supabase',
            'balance': 0.0,
            'free_daily_credits': 0.02,
            'free_daily_used': 0.0,
            'free_daily_remaining': 0.02,
            'total_recharged': 0.0,
            'is_free_user': not paid,
            'has_recharged': paid,
            'projects': {'used': 0, 'limit': 50 if paid else 1},
            'exports': ['txt', 'docx', 'pdf'] if paid else ['txt'],
            'today_stats': {'calls': 0, 'cost': 0}
        })

    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    row = conn.execute('SELECT * FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()
    if not row:
        conn.close()
        _ensure_balance_record(user_id)
        conn = sqlite3.connect(str(LICENSE_DB))
        conn.row_factory = sqlite3.Row
        row = conn.execute('SELECT * FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()
    today = datetime.now().strftime('%Y-%m-%d')
    free_date = row['free_daily_date'] or ''
    if free_date != today:
        free_used = 0.0
    else:
        free_used = row['free_daily_used'] or 0.0
    free_total = row['free_daily_credits'] or 0.02
    free_remaining = max(0, free_total - free_used)
    balance = round(row['balance'] or 0.0, 4)
    paid = (row['total_recharged'] or 0.0) > 0
    # 今日调用统计
    today_calls = conn.execute(
        "SELECT COUNT(*) as cnt, COALESCE(SUM(cost),0) as total_cost FROM ai_transactions WHERE user_id = ? AND created_at LIKE ?",
        (user_id, today + '%')
    ).fetchone()
    proj_count = len([p for p in PROJECTS_DIR.iterdir() if p.is_dir() and (p / 'novel.db').exists()])
    conn.close()
    return jsonify({
        'mode': 'balance',
        'balance': balance,
        'free_daily_credits': free_total,
        'free_daily_used': round(free_used, 4),
        'free_daily_remaining': round(free_remaining, 4),
        'total_recharged': round(row['total_recharged'] or 0.0, 2),
        'is_free_user': not paid,
        'has_recharged': paid,
        'projects': {'used': proj_count, 'limit': 50 if paid else 1},
        'exports': ['txt', 'docx', 'pdf'] if paid else ['txt'],
        'today_stats': {
            'calls': today_calls['cnt'] if today_calls else 0,
            'cost': round(today_calls['total_cost'] if today_calls else 0, 4)
        }
    })

@app.route('/api/license/activate', methods=['POST'])
def license_activate():
    """激活许可证"""
    data = request.json
    key = data.get('key', '').strip()
    email = data.get('email', '').strip()
    name = data.get('name', '').strip()
    
    if not key:
        return jsonify({'error': '请输入许可证密钥'}), 400
    
    valid, tier_code = verify_license_key(key)
    if not valid:
        return jsonify({'error': tier_code or '密钥无效'}), 400
    
    tier_map = {'DAY': 'day', 'WEK': 'week', 'MON': 'month', 'YER': 'year'}
    tier = tier_map.get(tier_code, 'free')
    
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    
    # 检查是否已存在此密钥
    existing = conn.execute('SELECT * FROM licenses WHERE id = ?', (key,)).fetchone()
    if existing:
        if existing['status'] == 'active':
            conn.close()
            return jsonify({'error': '此密钥已激活'}), 400
        # 重新激活
        now = datetime.now().isoformat()
        expires_at = datetime.fromtimestamp(get_tier_expiry(tier)).isoformat() if tier != 'free' else None
        conn.execute(
            'UPDATE licenses SET status = ?, activated_at = ?, expires_at = ?, customer_email = ?, customer_name = ? WHERE id = ?',
            ('active', now, expires_at, email, name, key)
        )
        log_license_action(key, 'reactivate', f'tier={tier}', conn=conn)
    else:
        now = datetime.now().isoformat()
        expires_at = datetime.fromtimestamp(get_tier_expiry(tier)).isoformat() if tier != 'free' else None
        conn.execute(
            'INSERT INTO licenses (id, tier, status, activated_at, expires_at, customer_email, customer_name, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
            (key, tier, 'active', now, expires_at, email, name, now)
        )
        log_license_action(key, 'activate', f'tier={tier}', conn=conn)

    # 创建/更新客户档案
    if email or name:
        existing_cp = conn.execute('SELECT license_id FROM customer_profiles WHERE license_id = ?', (key,)).fetchone()
        if existing_cp:
            conn.execute('UPDATE customer_profiles SET email = ?, name = ?, updated_at = ? WHERE license_id = ?',
                        (email, name, now, key))
        else:
            conn.execute('''
                INSERT OR IGNORE INTO customer_profiles (license_id, email, name, last_active_at, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (key, email, name, now, now, now))

    conn.commit()
    conn.close()

    return jsonify({'success': True, 'tier': tier, 'message': f'已激活{tier_map.get(tier_code, tier)}会员'})

@app.route('/api/license/check', methods=['GET'])
def license_check():
    """检查特定功能是否可用"""
    feature = request.args.get('feature', '')
    value = request.args.get('value', '')
    allowed = check_feature(feature, value)
    return jsonify({'allowed': allowed, 'feature': feature})

@app.route('/api/license/consume_ai', methods=['POST'])
def license_consume_ai():
    """消耗一次 AI 调用次数"""
    if not check_feature('ai_call'):
        return jsonify({'error': '今日 AI 次数已用完，请升级会员'}), 403
    
    lic = get_current_license()
    if lic:
        log_license_action(lic['id'], 'ai_call')
    
    return jsonify({'success': True})

# ===== 充值模式：余额 API =====

@app.route('/api/balance', methods=['GET'])
def api_get_balance():
    """获取当前余额和免费额度"""
    user_id = get_user_id()
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    row = conn.execute('SELECT * FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()
    if not row:
        conn.close()
        _ensure_balance_record(user_id)
        conn = sqlite3.connect(str(LICENSE_DB))
        conn.row_factory = sqlite3.Row
        row = conn.execute('SELECT * FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()
    # 重置每日免费额度
    today = datetime.now().strftime('%Y-%m-%d')
    free_date = row['free_daily_date'] or ''
    if free_date != today:
        conn.execute('UPDATE user_balances SET free_daily_used = 0.0, free_daily_date = ?, updated_at = ? WHERE user_id = ?',
                     (today, datetime.now().isoformat(), user_id))
        conn.commit()
        free_used = 0.0
    else:
        free_used = row['free_daily_used'] or 0.0
    free_total = row['free_daily_credits'] or 0.02
    # 今日 AI 调用次数
    today_calls = conn.execute(
        "SELECT COUNT(*) as cnt FROM ai_transactions WHERE user_id = ? AND created_at LIKE ?",
        (user_id, today + '%')
    ).fetchone()['cnt']
    conn.close()
    return jsonify({
        'balance': round(row['balance'] or 0.0, 4),
        'free_daily_credits': free_total,
        'free_daily_used': round(free_used, 4),
        'free_daily_remaining': round(max(0, free_total - free_used), 4),
        'total_recharged': round(row['total_recharged'] or 0.0, 2),
        'has_recharged': (row['total_recharged'] or 0) > 0,
        'today_calls': today_calls
    })

@app.route('/api/transactions', methods=['GET'])
def api_get_transactions():
    """获取 AI 调用交易记录"""
    user_id = get_user_id()
    limit = request.args.get('limit', 50, type=int)
    offset = request.args.get('offset', 0, type=int)
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    rows = conn.execute(
        'SELECT * FROM ai_transactions WHERE user_id = ? ORDER BY created_at DESC LIMIT ? OFFSET ?',
        (user_id, limit, offset)
    ).fetchall()
    total = conn.execute('SELECT COUNT(*) as cnt FROM ai_transactions WHERE user_id = ?', (user_id,)).fetchone()['cnt']
    conn.close()
    return jsonify({
        'transactions': [dict(r) for r in rows],
        'total': total, 'limit': limit, 'offset': offset
    })

@app.route('/api/recharge/packages', methods=['GET'])
def api_recharge_packages():
    """获取充值套餐和模型定价"""
    return jsonify({
        'packages': [
            {'id': '5', 'amount': 5, 'label': '¥5', 'desc': '试用体验 · 约500次基础调用'},
            {'id': '10', 'amount': 10, 'label': '¥10', 'desc': '约1,000次DeepSeek基础调用'},
            {'id': '20', 'amount': 20, 'label': '¥20', 'desc': '约2,000次DeepSeek基础调用'},
            {'id': '50', 'amount': 50, 'label': '¥50', 'desc': '约5,000次DeepSeek调用', 'popular': True},
            {'id': '100', 'amount': 100, 'label': '¥100', 'desc': '约10,000次DeepSeek调用'},
            {'id': '200', 'amount': 200, 'label': '¥200', 'desc': '约20,000次DeepSeek调用'},
            {'id': 'custom', 'amount': None, 'label': '自定义', 'desc': '输入任意金额'}
        ],
        'model_prices': {
            'deepseek': {'price_per_1k': 0.01, 'name': 'DeepSeek', 'desc': '¥0.01/千tokens · 最具性价比'},
            'gemini': {'price_per_1k': 0.04, 'name': 'Gemini 2.5 Flash', 'desc': '¥0.04/千tokens · 均衡之选'},
            'claude': {'price_per_1k': 0.12, 'name': 'Claude Sonnet 4', 'desc': '¥0.12/千tokens · 顶级品质'}
        }
    })

# ===== 购买 & 支付流程 =====

# 在 license.db 中添加 purchases 表
def init_purchase_table():
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.execute('''
        CREATE TABLE IF NOT EXISTS purchases (
            id TEXT PRIMARY KEY,
            tier TEXT NOT NULL,
            amount REAL NOT NULL,
            status TEXT DEFAULT 'pending',
            customer_email TEXT,
            customer_name TEXT,
            user_id TEXT,
            license_key TEXT,
            created_at TEXT,
            paid_at TEXT
        )
    ''')
    conn.commit()
    conn.close()

init_purchase_table()

def init_customer_platform_tables():
    """初始化客户档案和资金流水表"""
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.execute('''
        CREATE TABLE IF NOT EXISTS customer_profiles (
            license_id TEXT PRIMARY KEY,
            email TEXT DEFAULT '',
            name TEXT DEFAULT '',
            phone TEXT DEFAULT '',
            wechat_id TEXT DEFAULT '',
            total_purchases REAL DEFAULT 0,
            total_projects INTEGER DEFAULT 0,
            total_ai_calls INTEGER DEFAULT 0,
            last_active_at TEXT,
            notes TEXT DEFAULT '',
            created_at TEXT,
            updated_at TEXT
        )
    ''')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS platform_accounts (
            id TEXT PRIMARY KEY,
            type TEXT NOT NULL,
            amount REAL NOT NULL,
            source TEXT DEFAULT '',
            reference_id TEXT DEFAULT '',
            status TEXT DEFAULT 'pending',
            note TEXT DEFAULT '',
            created_at TEXT,
            updated_at TEXT
        )
    ''')
    try:
        conn.execute("ALTER TABLE purchases ADD COLUMN license_key TEXT DEFAULT ''")
    except sqlite3.OperationalError:
        pass
    conn.commit()
    conn.close()

init_customer_platform_tables()

# ===== 管理后台认证 =====
ADMIN_PASSWORD_HASH = os.environ.get('ADMIN_PASSWORD_HASH', '')
_ADMIN_TOKENS = {}  # {token: expires_at_timestamp}
_ADMIN_TOKEN_TTL = 86400  # 24 小时过期

if not ADMIN_PASSWORD_HASH:
    raise RuntimeError('[Novel Studio] 启动失败：ADMIN_PASSWORD_HASH 环境变量未设置。'
                       '请运行 python3 gen_password_hash.py 并将结果配置为环境变量。')

def verify_admin_password(password):
    """验证管理员密码，支持 SHA-256（旧格式）和 PBKDF2（新格式）"""
    if ADMIN_PASSWORD_HASH.startswith('pbkdf2:'):
        # 新格式: pbkdf2:salt_hex:key_hex
        _, salt_hex, key_hex = ADMIN_PASSWORD_HASH.split(':')
        salt = bytes.fromhex(salt_hex)
        key = hashlib.pbkdf2_hmac('sha256', password.encode(), salt, 100000)
        return key.hex() == key_hex
    # 旧格式: 纯 SHA-256 hex（向后兼容）
    return hashlib.sha256(password.encode()).hexdigest() == ADMIN_PASSWORD_HASH

def _cleanup_expired_tokens():
    """移除过期的 admin token"""
    now = time.time()
    expired = [t for t, exp in _ADMIN_TOKENS.items() if now > exp]
    for t in expired:
        del _ADMIN_TOKENS[t]

def admin_required(f):
    """装饰器：验证管理后台 token，24 小时过期"""
    @functools.wraps(f)
    def wrapper(*args, **kwargs):
        token = request.headers.get('X-Admin-Token', '')
        _cleanup_expired_tokens()
        if not token or token not in _ADMIN_TOKENS:
            return jsonify({'error': '未授权访问，请先登录'}), 401
        return f(*args, **kwargs)
    return wrapper

# ===== QR 码图片上传 =====
UPLOAD_DIR = BASE_DIR / 'static' / 'uploads'
UPLOAD_DIR.mkdir(exist_ok=True)

@app.route('/api/admin/upload-qr', methods=['POST'])
@admin_required
def admin_upload_qr():
    if 'file' not in request.files:
        return jsonify({'error': '请选择文件'}), 400
    file = request.files['file']
    qr_type = request.form.get('type', 'wechat')
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else 'png'
    if ext not in ('png', 'jpg', 'jpeg', 'gif', 'webp'):
        return jsonify({'error': '仅支持图片格式（PNG/JPG/GIF/WEBP）'}), 400
    filename = f'qr_{qr_type}_{secrets.token_hex(8)}.{ext}'
    filepath = UPLOAD_DIR / filename
    file.save(str(filepath))
    relative_url = f'/static/uploads/{filename}'
    save_setting(f'{qr_type}_qr', relative_url)
    return jsonify({'success': True, 'url': relative_url})

# ===== 客户管理 API =====
@app.route('/api/admin/customers', methods=['GET'])
@admin_required
def admin_list_customers():
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    conn.execute('''
        CREATE TABLE IF NOT EXISTS customer_profiles (
            license_id TEXT PRIMARY KEY, email TEXT DEFAULT '', name TEXT DEFAULT '',
            phone TEXT DEFAULT '', wechat_id TEXT DEFAULT '', total_purchases REAL DEFAULT 0,
            total_projects INTEGER DEFAULT 0, total_ai_calls INTEGER DEFAULT 0,
            last_active_at TEXT, notes TEXT DEFAULT '', created_at TEXT, updated_at TEXT
        )
    ''')
    rows = conn.execute('SELECT * FROM customer_profiles ORDER BY last_active_at DESC').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/admin/customers/<license_id>', methods=['GET'])
@admin_required
def admin_get_customer(license_id):
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    profile = conn.execute('SELECT * FROM customer_profiles WHERE license_id = ?', (license_id,)).fetchone()
    if not profile:
        conn.close()
        return jsonify({'error': '客户不存在'}), 404
    purchases = conn.execute(
        "SELECT * FROM purchases WHERE customer_email = ? OR customer_name = ? ORDER BY created_at DESC",
        (profile['email'] or '', profile['name'] or '')
    ).fetchall() if (profile['email'] or profile['name']) else []
    conn.close()
    return jsonify({'profile': dict(profile), 'purchases': [dict(p) for p in purchases]})

@app.route('/api/admin/customers/<license_id>', methods=['PUT'])
@admin_required
def admin_update_customer(license_id):
    data = request.json
    conn = sqlite3.connect(str(LICENSE_DB))
    now = datetime.now().isoformat()
    for key in ['email', 'name', 'phone', 'wechat_id', 'notes']:
        if key in data:
            conn.execute(f'UPDATE customer_profiles SET {key} = ?, updated_at = ? WHERE license_id = ?',
                        (data[key], now, license_id))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

# ===== 资金流水 API =====
@app.route('/api/admin/platform-accounts', methods=['GET'])
@admin_required
def admin_list_platform_accounts():
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    conn.execute('''
        CREATE TABLE IF NOT EXISTS platform_accounts (
            id TEXT PRIMARY KEY, type TEXT NOT NULL, amount REAL NOT NULL,
            source TEXT DEFAULT '', reference_id TEXT DEFAULT '', status TEXT DEFAULT 'pending',
            note TEXT DEFAULT '', created_at TEXT, updated_at TEXT
        )
    ''')
    rows = conn.execute('SELECT * FROM platform_accounts ORDER BY created_at DESC').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/admin/platform-accounts', methods=['POST'])
@admin_required
def admin_create_platform_account():
    data = request.json
    now = datetime.now().isoformat()
    rec_id = secrets.token_hex(8)
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.execute('''
        INSERT INTO platform_accounts (id, type, amount, source, reference_id, status, note, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (rec_id, data.get('type', 'withdrawal'), data.get('amount', 0),
          data.get('source', 'manual'), data.get('reference_id', ''),
          data.get('status', 'pending'), data.get('note', ''), now, now))
    conn.commit()
    conn.close()
    return jsonify({'id': rec_id, 'success': True})

@app.route('/api/admin/platform-accounts/<rec_id>', methods=['PUT'])
@admin_required
def admin_update_platform_account(rec_id):
    data = request.json
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.execute('UPDATE platform_accounts SET status = ?, updated_at = ? WHERE id = ?',
                 (data.get('status', 'completed'), datetime.now().isoformat(), rec_id))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

# ===== 管理后台：余额管理 =====

@app.route('/api/admin/balances', methods=['GET'])
@admin_required
def admin_list_balances():
    """列出所有用户余额"""
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    rows = conn.execute('''
        SELECT ub.*,
               (SELECT COUNT(*) FROM ai_transactions WHERE user_id = ub.user_id) as txn_count,
               (SELECT SUM(cost) FROM ai_transactions WHERE user_id = ub.user_id) as total_cost
        FROM user_balances ub ORDER BY ub.total_recharged DESC
    ''').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/admin/balances/adjust', methods=['POST'])
@admin_required
def admin_adjust_balance():
    """调整用户余额（增减）"""
    data = request.json
    user_id = data.get('user_id', '').strip()
    try:
        delta = float(data.get('delta', 0))
    except (TypeError, ValueError):
        return jsonify({'error': '无效金额'}), 400

    if not user_id or delta == 0:
        return jsonify({'error': '缺少参数'}), 400

    conn = sqlite3.connect(str(LICENSE_DB))
    _ensure_balance_record(user_id, conn)
    now = datetime.now().isoformat()
    conn.execute('UPDATE user_balances SET balance = MAX(0, balance + ?), updated_at = ? WHERE user_id = ?',
                 (delta, now, user_id))
    row = conn.execute('SELECT balance FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()

    # 记录调整
    note = f'管理员{"充值" if delta > 0 else "扣除"} ¥{abs(delta):.2f}'
    conn.execute('''
        INSERT INTO platform_accounts (id, type, amount, source, reference_id, status, note, created_at, updated_at)
        VALUES (?, ?, ?, 'admin', ?, 'completed', ?, ?, ?)
    ''', (secrets.token_hex(8), 'deposit' if delta > 0 else 'withdrawal', abs(delta), user_id, note, now, now))
    conn.commit()
    conn.close()

    return jsonify({'success': True, 'new_balance': round(row[0], 2), 'delta': delta})

@app.route('/api/purchase/create', methods=['POST'])
def purchase_create():
    """创建充值订单"""
    data = request.json
    package_id = data.get('package', '50')
    custom_amount = data.get('amount', 0)
    email = data.get('email', '').strip()
    name = data.get('name', '').strip()

    packages = {'5': 5, '10': 10, '20': 20, '50': 50, '100': 100, '200': 200}
    if package_id == 'custom' and custom_amount > 0:
        amount = float(custom_amount)
    else:
        amount = packages.get(package_id, 50)

    order_id = 'NVO-' + secrets.token_hex(6).upper()
    now = datetime.now().isoformat()

    conn = sqlite3.connect(str(LICENSE_DB))
    user_id = get_user_id()
    conn.execute(
        'INSERT INTO purchases (id, tier, amount, status, customer_email, customer_name, user_id, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
        (order_id, 'recharge', amount, 'pending', email, name, user_id, now)
    )
    conn.commit()
    conn.close()

    return jsonify({
        'order_id': order_id,
        'package': package_id,
        'amount': amount,
        'label': f'充值 ¥{amount:.0f}' if amount == int(amount) else f'充值 ¥{amount:.2f}',
        'qr_url': f'/api/purchase/qr/{order_id}'
    })

@app.route('/api/purchase/confirm', methods=['POST'])
@admin_required
def purchase_confirm():
    """确认支付成功 → 余额充值"""
    data = request.json
    order_id = data.get('order_id', '').strip()
    if not order_id:
        return jsonify({'error': '缺少订单号'}), 400

    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row

    order = conn.execute('SELECT * FROM purchases WHERE id = ?', (order_id,)).fetchone()
    if not order:
        conn.close()
        return jsonify({'error': '订单不存在'}), 404

    if order['status'] == 'paid':
        conn.close()
        return jsonify({
            'success': True,
            'amount': order['amount'],
            'already_paid': True,
            'message': '此订单已支付，余额已到账'
        })

    now = datetime.now().isoformat()

    # 更新订单状态
    conn.execute(
        'UPDATE purchases SET status = ?, paid_at = ? WHERE id = ?',
        ('paid', now, order_id)
    )

    # 充值到用户余额
    user_id = order['user_id'] or get_user_id()
    _ensure_balance_record(user_id, conn)
    old_balance = conn.execute('SELECT balance FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()
    old_bal = old_balance['balance'] if old_balance else 0.0
    new_bal = old_bal + order['amount']
    conn.execute('UPDATE user_balances SET balance = ?, total_recharged = total_recharged + ?, updated_at = ? WHERE user_id = ?',
                 (new_bal, order['amount'], now, user_id))
    conn.commit()
    # 读取最新余额
    final = conn.execute('SELECT balance, total_recharged FROM user_balances WHERE user_id = ?', (user_id,)).fetchone()

    cust_name = order['customer_name'] or ''
    # 资金流水记录
    conn.execute('''
        INSERT INTO platform_accounts (id, type, amount, source, reference_id, status, note, created_at, updated_at)
        VALUES (?, 'deposit', ?, 'recharge', ?, 'completed', ?, ?, ?)
    ''', (order_id, order['amount'], order_id,
          f'充值 {order["amount"]:.0f}元 - {cust_name}', now, now))

    # 更新客户档案
    cust_email = order['customer_email'] or ''
    if cust_email or cust_name:
        existing_cust = conn.execute('SELECT license_id FROM customer_profiles WHERE license_id = ?', (user_id,)).fetchone()
        if existing_cust:
            conn.execute('UPDATE customer_profiles SET total_purchases = total_purchases + ?, last_active_at = ?, updated_at = ? WHERE license_id = ?',
                        (order['amount'], now, now, user_id))
        else:
            conn.execute('''
                INSERT INTO customer_profiles (license_id, email, name, total_purchases, last_active_at, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (user_id, cust_email, cust_name, order['amount'], now, now, now))

    conn.commit()
    conn.close()

    return jsonify({
        'success': True,
        'amount': order['amount'],
        'new_balance': round(final['balance'] if final else new_bal, 2),
        'total_recharged': round(final['total_recharged'] if final else order['amount'], 2),
        'message': f'充值成功！¥{order["amount"]:.0f} 已到账，当前余额 ¥{final["balance"] if final else new_bal:.2f}'
    })

@app.route('/api/purchase/history', methods=['GET'])
def purchase_history():
    """购买记录"""
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    rows = conn.execute("SELECT * FROM purchases WHERE status = 'paid' ORDER BY paid_at DESC LIMIT 20").fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

# ===== AI 原创性保障 =====

@app.route('/api/projects/<project_id>/ai/originality-check', methods=['POST'])
def ai_originality_check(project_id):
    """AI 原创性检查：检测内容是否过于套路化、模板化、或存在潜在雷同"""
    if not check_feature('premium'):
        return jsonify({'error': '原创性检查为高级功能，请升级会员'}), 403
    data = request.json
    content = data.get('content', '')
    if not content or len(content.strip()) < 30:
        return jsonify({'error': '请提供至少30字的内容进行检查'}), 400
    
    sample = content[:3000]
    
    messages = [
        {'role': 'system', 'content': '''你是一位严格的文学原创性审稿人。请从以下维度分析文本：

1. **套路化程度**：是否使用了网文常见套话、万能开头、固定模板？
2. **人物对话雷同**：不同角色语气是否千篇一律？
3. **情节公式化**：是否有明显的"打脸—逆袭—收后宫"等公式？
4. **描写陈词滥调**：是否有"剑眉星目""倾国倾城"等用烂的形容词？
5. **独特视角**：叙述角度和世界观是否具有独特性？
6. **风格辨识度**：是否有鲜明的个人风格，还是像AI随机生成的？

输出JSON格式：
{
  "originality_score": 0-100 (原创性分数，越高越原创),
  "risk_level": "low|medium|high",
  "template_patterns": ["发现的套路1", "套路2"],
  "unique_elements": ["独特之处1", "独特之处2"],
  "cliche_list": [{"phrase": "陈词滥调", "suggestion": "替换建议"}],
  "verdict": "综合评价一句话",
  "improvement_tips": ["提升原创性建议1", "建议2"]
}
直接输出JSON，不要任何解释。'''},
        {'role': 'user', 'content': f'请检查以下文本的原创性：\n\n{sample}'}
    ]
    result = call_ai(messages, model=data.get('model'), temperature=0.3, max_tokens=1500)
    return ai_result(result, data.get('model'))

@app.route('/api/projects/<project_id>/ai/style-fingerprint', methods=['POST'])
def ai_style_fingerprint(project_id):
    """生成写作风格指纹 — 提取用户独特写作特征，后续AI写作时参照"""
    if not check_feature('premium'):
        return jsonify({'error': '风格指纹为高级功能，请升级会员'}), 403
    data = request.json
    samples = data.get('samples', '')  # 用户手动写的章节内容
    
    if not samples or len(samples.strip()) < 100:
        return jsonify({'error': '需要至少100字的手写样本'}), 400
    
    # 取多个章节的样本（最多5000字）
    sample = samples[:5000]
    
    messages = [
        {'role': 'system', 'content': '''你是写作风格分析师。请从提供的文本中提取作者的独特风格特征，生成"风格指纹"。
        
分析维度：
1. 句式特征：长短句比例、常用句式结构
2. 词汇偏好：高频词汇、独特用词
3. 修辞习惯：常用修辞手法
4. 对话风格：角色语言特点
5. 节奏感：段落长度、叙述节奏
6. 情感色彩：常用情感表达方式

输出JSON格式：
{
  "style_summary": "一句话风格概括",
  "sentence_pattern": "句式特征描述",
  "vocabulary_traits": ["词汇特征1", "特征2"],
  "rhetoric_style": "修辞风格描述",
  "dialogue_traits": "对话特点",
  "pacing": "节奏特征",
  "emotional_tone": "情感基调",
  "writing_commands": ["写作指令1（用于AI提示词）", "指令2"],
  "avoid_patterns": ["应避免的写作模式1", "模式2"]
}
直接输出JSON，不含其他内容。'''},
        {'role': 'user', 'content': f'请分析以下文本的写作风格：\n\n{sample}'}
    ]
    result = call_ai(messages, model=data.get('model'), temperature=0.3, max_tokens=1200)
    return ai_result(result, data.get('model'))

# ===== 内容审阅官 =====
@app.route('/api/projects/<project_id>/ai/review', methods=['POST'])
def ai_review(project_id):
    """AI 内容审阅官 — 语法、标点、内容、情节审查"""
    if not check_feature('premium'):
        return jsonify({'error': '内容审阅为高级功能，请升级会员'}), 403
    data = request.json
    content = data.get('content', '')
    review_type = data.get('review_type', 'full')

    if not content or len(content.strip()) < 30:
        return jsonify({'error': '请提供至少30字的内容进行审阅'}), 400

    sample = content[:4000]

    review_prompts = {
        'full': '全面审阅：语法标点、内容质量、情节逻辑、文笔风格',
        'grammar': '重点检查：错别字、标点符号误用、语法错误、的地得用法、语句通顺度',
        'plot': '重点检查：情节连贯性、逻辑漏洞、前后矛盾、伏笔遗漏、节奏问题',
        'style': '重点检查：文笔流畅度、句式变化、修辞恰当性、节奏控制、读者体验',
    }

    messages = [{
        'role': 'system',
        'content': '''你是一位资深文学编辑，拥有20年审稿经验。请审阅以下小说内容，给出专业、具体的反馈。

输出格式（严格遵守JSON）：
{
  "overall_score": 0-100,
  "grammar_issues": [
    {"text": "原文片段", "correction": "修改建议", "severity": "high|medium|low", "reason": "原因"}
  ],
  "plot_issues": [
    {"type": "逻辑漏洞|前后矛盾|节奏问题|伏笔缺失", "description": "描述", "suggestion": "建议", "severity": "high|medium|low"}
  ],
  "style_feedback": {
    "strengths": ["优点1", "优点2"],
    "weaknesses": ["问题1", "问题2"],
    "suggestions": ["建议1", "建议2"]
  },
  "professional_advice": "100-200字的综合写作建议，要具体、可操作",
  "highlight_phrases": ["写得好的句子1", "句子2"]
}

直接输出JSON，不含其他内容。'''
    }, {
        'role': 'user',
        'content': f'审阅类型：{review_prompts.get(review_type, review_prompts["full"])}\n\n待审阅内容：\n{sample}'
    }]

    result = call_ai(messages, model=data.get('model'), temperature=0.3, max_tokens=2500)
    return ai_result(result, data.get('model'))

# ===== 公开支付信息 API（供用户购买页面使用）=====
@app.route('/api/payment-info', methods=['GET'])
def api_payment_info():
    """返回当前收款配置（公开接口）"""
    return jsonify({
        'bank_name': get_setting('bank_name', ''),
        'bank_account': get_setting('bank_account', ''),
        'bank_holder': get_setting('bank_holder', ''),
        'wechat_qr': get_setting('wechat_qr', ''),
        'alipay_qr': get_setting('alipay_qr', ''),
        'contact_wechat': get_setting('contact_wechat', ''),
        'contact_email': get_setting('contact_email', ''),
    })

@app.route('/api/admin/verify', methods=['POST'])
def admin_verify():
    data = request.json
    if verify_admin_password(data.get('password', '')):
        token = 'admin-' + secrets.token_hex(32)
        _ADMIN_TOKENS[token] = time.time() + _ADMIN_TOKEN_TTL
        resp = {'token': token, 'expires_in': _ADMIN_TOKEN_TTL}
        # 旧格式密码哈希提示升级
        if not ADMIN_PASSWORD_HASH.startswith('pbkdf2:'):
            resp['security_notice'] = '建议升级密码哈希为 PBKDF2 格式，运行 python3 gen_password_hash.py'
        return jsonify(resp)
    return jsonify({'error': '密码错误'}), 403

@app.route('/api/admin/licenses', methods=['GET'])
@admin_required
def admin_list_licenses():
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    rows = conn.execute('SELECT * FROM licenses ORDER BY created_at DESC').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/admin/licenses/generate', methods=['POST'])
@admin_required
def admin_generate_license():
    data = request.json
    tier = data.get('tier', 'month')
    count = data.get('count', 1)
    email = data.get('email', '')
    name = data.get('name', '')

    keys = []
    conn = sqlite3.connect(str(LICENSE_DB))
    now = datetime.now().isoformat()

    for _ in range(count):
        key = generate_license_key(tier)
        expires_at = datetime.fromtimestamp(get_tier_expiry(tier)).isoformat()
        conn.execute(
            'INSERT INTO licenses (id, tier, status, activated_at, expires_at, customer_email, customer_name, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
            (key, tier, 'pending', None, expires_at, email, name, now)
        )
        keys.append({'key': key, 'tier': tier})

    conn.commit()
    conn.close()

    return jsonify({'success': True, 'keys': keys})

@app.route('/api/admin/licenses/<key>/revoke', methods=['POST'])
@admin_required
def admin_revoke_license(key):
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.execute("UPDATE licenses SET status = 'revoked' WHERE id = ?", (key,))
    conn.commit()
    conn.close()
    log_license_action(key, 'revoke')
    return jsonify({'success': True})

@app.route('/api/admin/stats', methods=['GET'])
@admin_required
def admin_stats():
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row

    total = conn.execute('SELECT COUNT(*) as cnt FROM licenses').fetchone()['cnt']
    active = conn.execute("SELECT COUNT(*) as cnt FROM licenses WHERE status = 'active'").fetchone()['cnt']
    today_activations = conn.execute(
        "SELECT COUNT(*) as cnt FROM activation_log WHERE action = 'activate' AND timestamp LIKE ?",
        (datetime.now().strftime('%Y-%m-%d') + '%',)
    ).fetchone()['cnt']
    conn.close()

    proj_count = len([p for p in PROJECTS_DIR.iterdir() if p.is_dir() and (p / 'novel.db').exists()])

    return jsonify({
        'total_licenses': total,
        'active_licenses': active,
        'today_activations': today_activations,
        'total_projects': proj_count
    })

# ===== 资金管理 API =====

@app.route('/api/admin/purchases', methods=['GET'])
@admin_required
def admin_purchases():
    """获取购买记录列表"""
    status_filter = request.args.get('status', '')
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    if status_filter:
        rows = conn.execute('SELECT * FROM purchases WHERE status = ? ORDER BY created_at DESC', (status_filter,)).fetchall()
    else:
        rows = conn.execute('SELECT * FROM purchases ORDER BY created_at DESC').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/admin/purchases/<purchase_id>/status', methods=['POST'])
@admin_required
def admin_update_purchase_status(purchase_id):
    """更新购买记录状态"""
    data = request.json
    new_status = data.get('status', '')
    if new_status not in ('pending', 'paid', 'cancelled', 'refunded'):
        return jsonify({'error': '无效状态'}), 400
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.execute("UPDATE purchases SET status = ?, paid_at = ? WHERE id = ?",
                 (new_status, datetime.now().isoformat() if new_status == 'paid' else None, purchase_id))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

@app.route('/api/admin/finance-summary', methods=['GET'])
@admin_required
def admin_finance_summary():
    """资金概览"""
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row

    total_revenue = conn.execute(
        "SELECT COALESCE(SUM(amount), 0) as total FROM purchases WHERE status = 'paid'"
    ).fetchone()['total']
    pending_amount = conn.execute(
        "SELECT COALESCE(SUM(amount), 0) as total FROM purchases WHERE status = 'pending'"
    ).fetchone()['total']
    total_orders = conn.execute('SELECT COUNT(*) as cnt FROM purchases').fetchone()['cnt']
    paid_orders = conn.execute("SELECT COUNT(*) as cnt FROM purchases WHERE status = 'paid'").fetchone()['cnt']

    # 按套餐统计
    by_tier = conn.execute(
        "SELECT tier, COUNT(*) as cnt, COALESCE(SUM(amount),0) as total FROM purchases WHERE status='paid' GROUP BY tier"
    ).fetchall()

    # 本月收入
    this_month = datetime.now().strftime('%Y-%m')
    month_revenue = conn.execute(
        "SELECT COALESCE(SUM(amount), 0) as total FROM purchases WHERE status='paid' AND paid_at LIKE ?",
        (this_month + '%',)
    ).fetchone()['total']

    conn.close()

    return jsonify({
        'total_revenue': total_revenue,
        'pending_amount': pending_amount,
        'total_orders': total_orders,
        'paid_orders': paid_orders,
        'month_revenue': month_revenue,
        'by_tier': [{'tier': r['tier'], 'count': r['cnt'], 'total': r['total']} for r in by_tier]
    })

# ===== 收款配置 & 联系方式 API =====

@app.route('/api/admin/payment-config', methods=['GET'])
@admin_required
def admin_get_payment_config():
    """获取收款配置"""
    return jsonify({
        'bank_name': get_setting('bank_name', ''),
        'bank_account': get_setting('bank_account', ''),
        'bank_holder': get_setting('bank_holder', ''),
        'wechat_qr': get_setting('wechat_qr', ''),
        'alipay_qr': get_setting('alipay_qr', ''),
    })

@app.route('/api/admin/payment-config', methods=['POST'])
@admin_required
def admin_save_payment_config():
    """保存收款配置"""
    data = request.json
    for key in ('bank_name', 'bank_account', 'bank_holder', 'wechat_qr', 'alipay_qr'):
        if key in data:
            save_setting(key, data[key])
    return jsonify({'success': True})

@app.route('/api/admin/contact-config', methods=['GET'])
@admin_required
def admin_get_contact_config():
    """获取联系方式"""
    return jsonify({
        'wechat_id': get_setting('contact_wechat', ''),
        'email': get_setting('contact_email', ''),
        'phone': get_setting('contact_phone', ''),
        'note': get_setting('contact_note', ''),
    })

@app.route('/api/admin/contact-config', methods=['POST'])
@admin_required
def admin_save_contact_config():
    """保存联系方式"""
    data = request.json
    for key in ('contact_wechat', 'contact_email', 'contact_phone', 'contact_note'):
        if key in data:
            save_setting(key, data[key])
    return jsonify({'success': True})

# ===== 用户反馈 =====
@app.route('/api/feedback', methods=['POST'])
def api_submit_feedback():
    """接收用户反馈并持久化"""
    data = request.json
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.execute('''
        CREATE TABLE IF NOT EXISTS feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            rating INTEGER DEFAULT 0,
            type TEXT DEFAULT '',
            content TEXT NOT NULL,
            contact TEXT DEFAULT '',
            created_at TEXT
        )
    ''')
    conn.execute(
        'INSERT INTO feedback (rating, type, content, contact, created_at) VALUES (?, ?, ?, ?, ?)',
        (data.get('rating', 0), data.get('type', ''), data.get('content', ''), data.get('contact', ''), datetime.now().isoformat())
    )
    conn.commit()
    conn.close()
    return jsonify({'success': True})

@app.route('/api/feedback', methods=['GET'])
@admin_required
def api_get_feedback():
    """管理后台查看反馈列表"""
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    conn.execute('CREATE TABLE IF NOT EXISTS feedback (id INTEGER PRIMARY KEY AUTOINCREMENT, rating INTEGER DEFAULT 0, type TEXT DEFAULT "", content TEXT NOT NULL, contact TEXT DEFAULT "", created_at TEXT)')
    rows = conn.execute('SELECT * FROM feedback ORDER BY created_at DESC LIMIT 100').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

# ===== 推荐裂变系统 =====

def _ensure_referral_tables():
    """初始化推荐系统表"""
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.execute('''CREATE TABLE IF NOT EXISTS referral_codes (
        code TEXT PRIMARY KEY,
        owner_license_id TEXT DEFAULT '',
        created_at TEXT,
        total_clicks INTEGER DEFAULT 0,
        total_signups INTEGER DEFAULT 0,
        rewards_claimed INTEGER DEFAULT 0
    )''')
    conn.execute('''CREATE TABLE IF NOT EXISTS referral_rewards (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        license_id TEXT NOT NULL,
        reward_type TEXT DEFAULT 'ai_calls',
        reward_amount INTEGER DEFAULT 10,
        reason TEXT DEFAULT '',
        claimed INTEGER DEFAULT 0,
        created_at TEXT
    )''')
    conn.commit()
    conn.close()

@app.route('/api/invite/generate', methods=['POST'])
def api_generate_invite():
    """生成专属推荐码，创建推荐关系"""
    _ensure_referral_tables()
    user_id = get_user_id()

    # 检查是否已有推荐码
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    existing = conn.execute('SELECT code FROM referral_codes WHERE owner_license_id = ?', (user_id,)).fetchone()
    if existing:
        conn.close()
        return jsonify({'code': existing['code'], 'url': f'/invite/{existing["code"]}'})

    # 生成新推荐码
    code = 'NVS-' + secrets.token_hex(4).upper()
    conn.execute(
        'INSERT INTO referral_codes (code, owner_license_id, created_at) VALUES (?, ?, ?)',
        (code, user_id, datetime.now().isoformat())
    )
    conn.commit()
    conn.close()
    return jsonify({'code': code, 'url': f'/invite/{code}'})

@app.route('/api/invite/<code>/stats', methods=['GET'])
def api_invite_stats(code):
    """查看推荐码的转化数据"""
    _ensure_referral_tables()
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    row = conn.execute('SELECT * FROM referral_codes WHERE code = ?', (code,)).fetchall()
    conn.close()
    if not row:
        return jsonify({'error': '推荐码不存在'}), 404
    r = dict(row[0])
    return jsonify({
        'code': r['code'],
        'created_at': r['created_at'],
        'clicks': r['total_clicks'],
        'signups': r['total_signups'],
        'rewards_claimed': r['rewards_claimed'],
    })

@app.route('/api/invite/claim-reward', methods=['POST'])
def api_claim_referral_reward():
    """被推荐人通过推荐链接进入后，双方获得余额奖励"""
    data = request.json
    invite_code = data.get('code', '')
    _ensure_referral_tables()
    user_id = get_user_id()

    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row

    # 查找推荐码
    ref = conn.execute('SELECT * FROM referral_codes WHERE code = ?', (invite_code,)).fetchone()
    if not ref:
        conn.close()
        return jsonify({'error': '推荐码无效'}), 400

    if ref['owner_license_id'] == user_id:
        conn.close()
        return jsonify({'error': '不能使用自己的推荐码'}), 400

    # 检查是否已领取过
    existing = conn.execute(
        'SELECT id FROM referral_rewards WHERE license_id = ? AND reason LIKE ?',
        (user_id, f'%{invite_code}%')
    ).fetchone()
    if existing:
        conn.close()
        return jsonify({'success': True, 'message': '奖励已领取过', 'already_claimed': True})

    # 更新推荐码统计
    conn.execute(
        'UPDATE referral_codes SET total_signups = total_signups + 1 WHERE code = ?',
        (invite_code,)
    )

    now = datetime.now().isoformat()
    # 给推荐人发 ¥2 余额
    _ensure_balance_record(ref['owner_license_id'], conn)
    conn.execute(
        'UPDATE user_balances SET balance = balance + 2.0, updated_at = ? WHERE user_id = ?',
        (now, ref['owner_license_id'])
    )
    conn.execute(
        'INSERT INTO referral_rewards (license_id, reward_type, reward_amount, reason, created_at) VALUES (?, ?, ?, ?, ?)',
        (ref['owner_license_id'], 'balance', 2.0, f'推荐用户入站奖励', now)
    )

    # 给被推荐人发 ¥1 余额
    _ensure_balance_record(user_id, conn)
    conn.execute(
        'UPDATE user_balances SET balance = balance + 1.0, updated_at = ? WHERE user_id = ?',
        (now, user_id)
    )
    conn.execute(
        'INSERT INTO referral_rewards (license_id, reward_type, reward_amount, reason, created_at) VALUES (?, ?, ?, ?, ?)',
        (user_id, 'balance', 1.0, f'通过推荐码 {invite_code} 注册', now)
    )

    conn.commit()
    conn.close()
    return jsonify({'success': True, 'message': '奖励已发放！推荐人获得 ¥2.00，你获得 ¥1.00 余额奖励'})

# ===== 运营数据统计 =====
@app.route('/api/site/stats', methods=['GET'])
def api_site_stats():
    """公开运营数据（可展示在落地页）"""
    proj_count = len([p for p in PROJECTS_DIR.iterdir() if p.is_dir() and (p / 'novel.db').exists()])
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.row_factory = sqlite3.Row
    total_users = conn.execute('SELECT COUNT(*) as cnt FROM licenses WHERE status = "active"').fetchone()['cnt']
    conn.close()
    return jsonify({
        'total_projects': proj_count,
        'total_users': total_users,
    })

# ===== 新手引导数据 =====
@app.route('/api/onboarding', methods=['POST'])
def api_onboarding():
    """记录新用户引导完成状态"""
    data = request.json
    save_setting('onboarding_completed', 'true')
    save_setting('user_experience_level', data.get('experience', ''))
    save_setting('user_genre_preference', data.get('genre', ''))
    return jsonify({'success': True})

@app.route('/api/onboarding', methods=['GET'])
def api_get_onboarding():
    """获取引导状态"""
    return jsonify({
        'completed': get_setting('onboarding_completed', '') == 'true',
        'experience': get_setting('user_experience_level', ''),
        'genre': get_setting('user_genre_preference', ''),
    })

# ===== AI 智能客服 =====
SUPPORT_KNOWLEDGE = """
你是 Novel Studio 的智能客服助手，请用友好、专业的中文回答用户问题。

## 产品信息
- 产品名：Novel Studio，AI 驱动的小说写作工具
- 核心功能：AI 智能续写（DeepSeek + Claude 双引擎）、文笔润色、AI味检测与去除、头脑风暴、灵感构思、风格指纹分析、角色档案管理、情节大纲规划、多格式导出（TXT/Word/PDF）
- 特色：支持 DeepSeek 和 Claude 双模型自由切换，9种AI写作功能，结构化角色和大纲管理

## 套餐与价格
- 免费版：每天3次AI调用、1个项目、TXT导出
- 日卡 ¥9.9：无限AI调用、无限项目、全格式导出
- 月卡 ¥49.9（推荐）：无限AI调用、无限项目、全格式导出、风格指纹分析、优先客服
- 年卡 ¥299：月卡全部权益，相当于¥25/月，节省50%

## 使用指南
- 开始写作：创建项目 → 新建章节 → 在编辑区写作 → 使用侧边栏AI面板续写/润色
- AI续写步骤：打开「智能续写」面板 → 选择方向（继续/转折/高潮/收尾）→ 点击「开始续写」
- 去AI味：打开「AI味检测」面板 → 粘贴或选择文本 → 点击检测 → 使用去AI味重写
- 导出：侧边栏「导出作品」→ 选择格式（TXT免费，DOCX/PDF需付费）
- 角色管理：侧边栏「角色设计」→ 创建角色 → 填写7维度档案 → AI续写时自动参考角色信息
- 设置API Key：侧边栏「设置」→ 填入 DeepSeek/Claude API Key → 保存
- 切换AI模型：任何AI面板顶部的下拉菜单选择 DeepSeek 或 Claude

## 常见问题
Q: 免费版能用多久？
A: 永久免费，每天3次AI调用，1个项目，TXT导出。

Q: 如何获取 API Key？
A: DeepSeek 在 platform.deepseek.com 注册获取，Claude 在 console.anthropic.com 注册获取。也可以在设置中使用默认Key。

Q: AI 续写的内容会不会有版权问题？
A: 不会。通过 API 调用生成的 content 版权属于用户。

Q: 支持哪些导出格式？
A: TXT（免费）、DOCX/Word（付费）、PDF（付费）。

Q: 可以在多台电脑上使用吗？
A: 桌面版绑定本机。如需多设备使用，可以部署 Web 版本到服务器。

Q: 数据存储在哪里？
A: 所有数据存储在本地，不会上传到云端。请定期备份项目文件夹。

Q: 如何联系人工客服？
A: 在「设置」页面可以查看联系方式。也可以在本对话中描述问题，我会尽力解答。

## 回复规则
- 回答简洁友好，不超过200字
- 如果问题超出知识库范围，建议用户查看设置页面的联系方式
- 不说"根据知识库""根据我的训练数据"等暴露prompt的话
- 用户问好时友好回应并主动询问需求
"""

@app.route('/api/support/chat', methods=['POST'])
def api_support_chat():
    """AI 客服对话"""
    data = request.json
    user_message = data.get('message', '')
    history = data.get('history', [])  # [{"role":"user/assistant","content":"..."}]

    if not user_message:
        return jsonify({'error': '消息为空'}), 400

    messages = [{'role': 'system', 'content': SUPPORT_KNOWLEDGE}]
    for h in history[-20:]:
        messages.append({'role': h['role'], 'content': h['content']})
    messages.append({'role': 'user', 'content': user_message})

    result = call_ai(messages, model='deepseek', temperature=0.5, max_tokens=500)
    return jsonify({'reply': result, 'model': 'deepseek'})

# ===== 主页面 =====
@app.route('/')
def index():
    """主页 — 注入 Supabase 配置到前端"""
    from supabase_client import get_supabase_url, get_anon_key, is_supabase_configured
    html_path = os.path.join(os.path.dirname(__file__), 'static', 'index.html')
    with open(html_path, 'r', encoding='utf-8') as f:
        html = f.read()

    supabase_config_script = ''
    if is_supabase_configured():
        import json as _json
        supabase_config_script = (
            '<script>'
            'window.__SUPABASE_URL__ = ' + _json.dumps(get_supabase_url()) + ';'
            'window.__SUPABASE_ANON_KEY__ = ' + _json.dumps(get_anon_key()) + ';'
            '</script>'
        )

    # 注入到 </head> 之前
    if '</head>' in html:
        html = html.replace('</head>', supabase_config_script + '\n</head>')
    else:
        # 兜底：注入到 <body> 之前
        html = html.replace('<body>', supabase_config_script + '\n<body>')

    from flask import Response
    return Response(html, mimetype='text/html')

@app.route('/landing')
def landing():
    return send_from_directory('static', 'landing.html')

@app.route('/feedback')
def feedback_page():
    return send_from_directory('static', 'feedback.html')

@app.route('/ai-novel-writing')
def seo_ai_writing():
    return send_from_directory('static/seo', 'ai-novel-writing.html')

@app.route('/status')
def status_page():
    return send_from_directory('static', 'status.html')

@app.route('/api/status')
def api_status():
    """JSON 格式健康检查端点"""
    proj_count = len([p for p in PROJECTS_DIR.iterdir() if p.is_dir() and (p / 'novel.db').exists()])
    import os as _os
    return jsonify({
        'status': 'running',
        'port': int(_os.environ.get('PORT', 80)),
        'total_projects': proj_count,
        'timestamp': datetime.now().isoformat()
    })

@app.route('/invite/<code>')
def invite_landing(code):
    """邀请链接落地页 —— 记录点击并展示带奖励提示的产品页"""
    _ensure_referral_tables()
    conn = sqlite3.connect(str(LICENSE_DB))
    conn.execute('UPDATE referral_codes SET total_clicks = total_clicks + 1 WHERE code = ?', (code,))
    conn.commit()
    conn.close()

    # 渲染带推荐码的落地页
    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Novel Studio — 你的好友邀请你来体验 AI 写作工具</title>
    <meta name="description" content="AI 驱动的小说写作工具，DeepSeek + Claude 双引擎，智能续写/润色/去AI味，免费试用">
    <style>
        * {{ margin:0; padding:0; box-sizing:border-box; }}
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'PingFang SC', 'Microsoft YaHei', sans-serif; background:#0a0a0f; color:#f5f5f7; display:flex; align-items:center; justify-content:center; min-height:100vh; padding:24px; }}
        .card {{ max-width:480px; width:100%; background:rgba(255,255,255,0.03); border:1px solid rgba(255,255,255,0.08); border-radius:16px; padding:40px 32px; text-align:center; }}
        .gift {{ font-size:56px; margin-bottom:16px; }}
        h1 {{ font-size:24px; font-weight:700; margin-bottom:8px; }}
        .sub {{ color:#98989d; font-size:14px; margin-bottom:24px; line-height:1.7; }}
        .reward-box {{ background:linear-gradient(135deg,rgba(10,132,255,0.08),rgba(94,92,230,0.06)); border:1px solid rgba(10,132,255,0.2); border-radius:12px; padding:16px; margin-bottom:24px; }}
        .reward-box .title {{ font-size:13px; font-weight:600; color:var(--accent-hover); margin-bottom:8px; }}
        .reward-box .items {{ font-size:12px; color:#98989d; line-height:2; }}
        .btn {{ display:inline-block; padding:14px 40px; border-radius:10px; font-size:15px; font-weight:600; cursor:pointer; border:none; text-decoration:none; transition:all 0.2s; background:linear-gradient(135deg,#0A84FF,#5e5ce6); color:#fff; }}
        .btn:hover {{ transform:translateY(-1px); box-shadow:0 8px 24px rgba(10,132,255,0.35); }}
        .code {{ font-family:monospace; font-size:12px; color:#55555a; margin-top:20px; }}
    </style>
</head>
<body>
<div class="card">
    <div class="gift">🎁</div>
    <h1>你的好友邀请你体验<br>Novel Studio</h1>
    <p class="sub">AI 驱动的小说写作工具<br>DeepSeek + Claude 双引擎 · 智能续写 · 去 AI 味</p>
    <div class="reward-box">
        <div class="title" style="color:#409cff;">🎉 通过邀请链接注册，双方都有奖励</div>
        <div class="items">
            👉 你：获得 <strong style="color:#30d158;">5 次</strong> 额外 AI 调用<br>
            👉 邀请人：获得 <strong style="color:#30d158;">10 次</strong> 额外 AI 调用<br>
            💡 激活后在设置页点击「领取推荐奖励」
        </div>
    </div>
    <a href="/?invite={code}" class="btn">开始免费使用</a>
    <div class="code">推荐码: {code}</div>
</div>
</body>
</html>'''
    return html

@app.route('/admin')
def admin_panel():
    return send_from_directory('static', 'admin.html')

@app.route('/manifest.json')
def manifest_json():
    return send_from_directory('static', 'manifest.json')

@app.route('/sw.js')
def service_worker():
    return send_from_directory('static', 'sw.js', mimetype='application/javascript')

@app.route('/offline.html')
def offline_page():
    return send_from_directory('static', 'offline.html')

# ===== Supabase 多租户 API（仅当 Supabase 环境变量配置时启用）=====
from supabase_client import (
    get_supabase, is_supabase_configured, require_auth,
    get_current_user, get_user_tier as supabase_get_user_tier,
    get_or_create_user_profile
)

if is_supabase_configured():

    @app.route('/api/auth/me', methods=['GET'])
    @require_auth
    def api_auth_me():
        """获取当前登录用户的 profile 和 tier"""
        user = get_current_user()
        profile = get_or_create_user_profile(user['id'], user.get('email', ''))
        tier = supabase_get_user_tier(user['id'])
        return jsonify({
            'user_id': user['id'],
            'email': user.get('email', ''),
            'tier': tier,
            'profile': profile
        })

    @app.route('/api/auth/books', methods=['GET'])
    @require_auth
    def api_auth_list_books():
        """列出当前用户的所有云书籍"""
        user = get_current_user()
        supabase = get_supabase()
        result = supabase.table('books').select('*')\
            .eq('user_id', user['id'])\
            .order('updated_at', desc=True)\
            .execute()
        return jsonify(result.data)

    @app.route('/api/auth/books', methods=['POST'])
    @require_auth
    def api_auth_create_book():
        """创建云书籍"""
        user = get_current_user()
        data = request.json or {}
        supabase = get_supabase()
        book = {
            'user_id': user['id'],
            'title': data.get('title', '未命名项目'),
            'description': data.get('description', ''),
            'genre': data.get('genre', ''),
        }
        result = supabase.table('books').insert(book).execute()
        if result.data:
            return jsonify(result.data[0]), 201
        return jsonify({'error': '创建失败'}), 500

    @app.route('/api/auth/books/<book_id>', methods=['DELETE'])
    @require_auth
    def api_auth_delete_book(book_id):
        """删除云书籍（RLS 保证只能删自己的）"""
        user = get_current_user()
        supabase = get_supabase()
        result = supabase.table('books').delete()\
            .eq('id', book_id).eq('user_id', user['id']).execute()
        return jsonify({'success': bool(result.data)})

    @app.route('/api/auth/books/<book_id>/chapters', methods=['GET'])
    @require_auth
    def api_auth_list_chapters(book_id):
        """列出云书籍的章节"""
        user = get_current_user()
        supabase = get_supabase()
        result = supabase.table('chapters').select('*')\
            .eq('book_id', book_id)\
            .eq('user_id', user['id'])\
            .order('sort_order', desc=False)\
            .execute()
        return jsonify(result.data)

    @app.route('/api/auth/books/<book_id>/chapters', methods=['POST'])
    @require_auth
    def api_auth_create_chapter(book_id):
        """创建云章节"""
        user = get_current_user()
        data = request.json or {}
        supabase = get_supabase()
        chapter = {
            'book_id': book_id,
            'user_id': user['id'],
            'title': data.get('title', '新章节'),
            'content': data.get('content', ''),
            'word_count': data.get('word_count', 0),
            'source': data.get('source', 'manual'),
        }
        result = supabase.table('chapters').insert(chapter).execute()
        if result.data:
            return jsonify(result.data[0]), 201
        return jsonify({'error': '创建失败'}), 500

    @app.route('/api/auth/books/<book_id>/chapters/<chapter_id>', methods=['PUT'])
    @require_auth
    def api_auth_update_chapter(book_id, chapter_id):
        """更新云章节"""
        user = get_current_user()
        data = request.json or {}
        supabase = get_supabase()
        update_data = {}
        for field in ['title', 'content', 'word_count']:
            if field in data:
                update_data[field] = data[field]
        if not update_data:
            return jsonify({'error': '无更新字段'}), 400
        result = supabase.table('chapters').update(update_data)\
            .eq('id', chapter_id).eq('user_id', user['id']).execute()
        if result.data:
            return jsonify(result.data[0])
        return jsonify({'error': '更新失败'}), 404

    @app.route('/api/auth/books/<book_id>/chapters/<chapter_id>', methods=['DELETE'])
    @require_auth
    def api_auth_delete_chapter(book_id, chapter_id):
        """删除云章节"""
        user = get_current_user()
        supabase = get_supabase()
        result = supabase.table('chapters').delete()\
            .eq('id', chapter_id).eq('user_id', user['id']).execute()
        return jsonify({'success': bool(result.data)})

if __name__ == '__main__':
    import sys
    port = int(sys.argv[1]) if len(sys.argv) > 1 else int(os.environ.get('PORT', 5050))
    print('=' * 60)
    print('  📖 Novel Studio — 专业小说写作工具')
    print(f'  运行地址: http://localhost:{port}')
    print(f'  管理后台: http://localhost:{port}/admin')
    print('=' * 60)
    app.run(host='0.0.0.0', port=port, debug=False)
